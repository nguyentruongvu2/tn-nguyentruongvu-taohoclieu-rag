"""
Convert Route - Stage 1: File to Markdown Conversion

This route handles the conversion of uploaded documents (PDF/DOCX)
to cleaned Markdown format without chunking.

POST /documents/convert
- Input: FormData with file + optional query parameters
- Output: JSON with markdown text and metadata
"""

import os
import logging
import re
import uuid
import json
import hashlib
import unicodedata
from datetime import datetime
from typing import Any, Dict, Literal
from fastapi import APIRouter, UploadFile, File, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel
import pdfplumber
from docx import Document
import concurrent.futures

try:
    import pytesseract
    from pytesseract import Output as TesseractOutput
except Exception:  # pragma: no cover - optional runtime dependency
    pytesseract = None
    TesseractOutput = None

try:
    from paddleocr import PaddleOCR
except Exception:  # pragma: no cover - optional runtime dependency
    PaddleOCR = None

try:
    import numpy as np
except Exception:  # pragma: no cover - optional runtime dependency
    np = None

try:
    import cv2
except Exception:  # pragma: no cover - optional runtime dependency
    cv2 = None

try:
    import pypdfium2 as pdfium
except Exception:  # pragma: no cover - optional runtime dependency
    pdfium = None

try:
    from PIL import Image, ImageOps, ImageFilter
except Exception:  # pragma: no cover - optional runtime dependency
    Image = None
    ImageOps = None
    ImageFilter = None

from ..markdown_cleaner import clean_markdown, clean_markdown_advanced
from ..markdown_advanced_divider_cleaner import AdvancedMarkdownCleaner

# Configure logging
logger = logging.getLogger(__name__)

# Constants
UPLOAD_DIR_PATH = os.path.join(os.path.dirname(__file__), '../../uploads')
OCR_CACHE_DIR_PATH = os.getenv("OCR_CACHE_DIR", os.path.join(UPLOAD_DIR_PATH, "ocr_page_cache"))
OCR_CACHE_ENABLED = os.getenv("OCR_CACHE_ENABLED", "true").strip().lower() in {"1", "true", "yes", "on"}
ALLOWED_EXTENSIONS = {'.pdf', '.docx'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR_PATH, exist_ok=True)
if OCR_CACHE_ENABLED:
    os.makedirs(OCR_CACHE_DIR_PATH, exist_ok=True)

# Response Model
class ConvertResponse(BaseModel):
    """Response from conversion endpoint"""
    success: bool
    markdown: str
    file_name: str
    file_size: int
    extraction_method: str
    cleaning_method: str = "standard"  # "standard" or "advanced"
    document_id: str
    preview_url: str | None = None
    pages: int = 0
    conversion_time_ms: float
    noise_removed_ratio: float = 0.0  # Percentage of lines removed
    quality: Literal["good", "medium", "bad"] = "medium"
    ocr_used: bool = False
    message: str


class ConvertSessionResponse(BaseModel):
    """Response for restoring an existing converted document session"""
    success: bool
    document_id: str
    markdown: str
    file_name: str
    cleaning_method: str
    pages: int = 0
    preview_url: str | None = None
    message: str


router = APIRouter(prefix="/documents", tags=["documents"])


CONVERSION_SESSIONS: Dict[str, Dict[str, Any]] = {}
OCR_MODE_VALUES = {"auto", "on", "off"}
OCR_UNCLEAR_MESSAGE = "Không thể khôi phục nội dung từ OCR"
_PADDLE_OCR_ENGINE = None

_VIETNAMESE_OCR_REGEX_FIXES: list[tuple[str, str]] = [
    (r"\bchu\s+ki\b", "chu kì"),
    (r"\btan\s+s6\b", "tần số"),
    (r"\btan\s+so\b", "tần số"),
    (r"\bging\b", "giống"),
    (r"\bnhu\b", "như"),
    (r"\bchuy[eé]n\b", "chuyển"),
    (r"\bdng\b", "động"),
    (r"\btrn\b", "tròn"),
    (r"\bd[eé]u\b", "đều"),
    (r"\bdao\s+dng\b", "dao động"),
    (r"\bdi[eé]u\b", "điều"),
    (r"\bc[uü]ng\b", "cũng"),
    (r"\bc6\b", "có"),
    (r"\btinh\s+chat\b", "tính chất"),
    (r"\bthat\s+vay\b", "thật vậy"),
    (r"\btudn\s+hoan\b", "tuần hoàn"),
    (r"\bkho[aá]ng\b", "khoảng"),
    (r"\bthi\s+gian\b", "thời gian"),
    (r"\bgoi\s+la\b", "gọi là"),
    (r"\bdi[eé]m\b", "điểm"),
    (r"\bduoc\b", "được"),
    (r"\bcn\b", "còn"),
    (r"\bv[oö]ng\b", "vòng"),
    (r"\blai\s+tr[oö]\s+v[eé]\b", "lại trở về"),
    (r"\bvi\s+tri\b", "vị trí"),
    (r"\btheo\s+h[uư]6ng\b", "theo hướng"),
    (r"\btu\s+d6\b", "từ đó"),
    (r"\bcac\b", "các"),
    (r"\bdinh\s+nghia\b", "định nghĩa"),
    (r"\bcua\b", "của"),
    (r"\bthuc\s+hien\b", "thực hiện"),
    (r"\btoan\s+phan\b", "toàn phần"),
    (r"\bvat\b", "vật"),
    (r"\bdon\s+vi\b", "đơn vị"),
    (r"\bgiay\b", "giây"),
    (r"\bki\s+hi[eé]u\b", "kí hiệu"),
    (r"\btr[eé]n\b", "trên"),
    (r"\bh[eé]c\b", "héc"),
    (r"\bki\s+hieu\b", "kí hiệu"),
    (r"\bm[oö]t\b", "một"),
]

_VIETNAMESE_OCR_WORD_FIXES: list[tuple[str, str]] = [
    ("Ging", "Giống"),
    ("nhu", "như"),
    ("chuyén", "chuyển"),
    ("dng", "động"),
    ("trn", "tròn"),
    ("déu", "đều"),
    ("dóng", "động"),
    ("döng", "động"),
    ("diéu", "điều"),
    ("cüng", "cũng"),
    ("c6", "có"),
    ("That", "Thật"),
    ("vay", "vậy"),
    ("cu", "cứ"),
    ("khoáng", "khoảng"),
    ("goi", "gọi"),
    ("la", "là"),
    ("thi", "thì"),
    ("diém", "điểm"),
    ("duoc", "được"),
    ("mt", "một"),
    ("mót", "một"),
    ("vöng", "vòng"),
    ("cn", "còn"),
    ("thuc", "thực"),
    ("hién", "hiện"),
    ("trö", "trở"),
    ("vé", "về"),
    ("hu6ng", "hướng"),
    ("Tu", "Từ"),
    ("d6", "đó"),
    ("cac", "các"),
    ("dinh", "định"),
    ("nghia", "nghĩa"),
    ("cua", "của"),
    ("laT", "là T"),
    ("laf", "là f"),
    ("dé", "để"),
    ("vat", "vật"),
    ("Don", "Đơn"),
    ("giay", "giây"),
    ("Tan", "Tần"),
    ("só", "số"),
    ("ki", "kí"),
    ("hiéu", "hiệu"),
    ("trén", "trên"),
]


def _sanitize_markdown_for_output(markdown: str) -> str:
    """Remove page-marker artifacts like 'Page 1/11' from user-facing markdown."""
    if not markdown:
        return ""

    page_line_pattern = re.compile(
        r'^\s{0,3}#{0,6}\s*(?:page|trang)\s*\d+\s*(?:/\s*\d+)?\s*[:.]?\s*$',
        re.IGNORECASE,
    )

    filtered_lines: list[str] = []
    for line in markdown.splitlines():
        if page_line_pattern.match(line):
            continue
        filtered_lines.append(line)

    cleaned = "\n".join(filtered_lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned


def _promote_markdown_headings(markdown: str) -> str:
    """Promote numbered/chapter lines to explicit markdown headings (#, ##, ###)."""
    if not markdown:
        return ""

    lines = markdown.splitlines()
    normalized: list[str] = []

    chapter_re = re.compile(r"^((?:ch[uư]ơng|bài|phần)\s+\d+[\.:\-]?\s*.*)$", re.IGNORECASE)
    level3_re = re.compile(r"^\d+\.\d+\.\d+(?:\.\d+)?\s*[:.)-]?\s+.+")
    level2_re = re.compile(r"^\d+\.\d+\s*[:.)-]?\s+.+")
    level1_re = re.compile(r"^\d+\s*[:.)-]\s+.+")

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            normalized.append("")
            continue

        # Keep existing markdown headings unchanged.
        if stripped.startswith("#"):
            normalized.append(stripped)
            continue

        if chapter_re.match(stripped):
            normalized.append(f"# {stripped}")
            continue
        if level3_re.match(stripped):
            normalized.append(f"### {stripped}")
            continue
        if level2_re.match(stripped):
            normalized.append(f"## {stripped}")
            continue
        if level1_re.match(stripped):
            normalized.append(f"# {stripped}")
            continue

        normalized.append(line)

    text = "\n".join(normalized)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _normalize_paragraph_line_breaks(markdown: str) -> str:
    """Merge hard-wrapped paragraph lines while preserving markdown structure."""
    if not markdown:
        return ""

    def _is_structural_line(line: str) -> bool:
        stripped = line.strip()
        if not stripped:
            return True
        if stripped.startswith("#"):
            return True
        if stripped.startswith(">"):
            return True
        if re.match(r"^\s*(?:[-*+]\s+|\d+[.)]\s+)", line):
            return True
        if stripped.startswith("|"):
            return True
        if re.match(r"^\s*\|?\s*:?[-]{3,}:?\s*(?:\|\s*:?[-]{3,}:?\s*)+\|?\s*$", stripped):
            return True
        return False

    def _merge_paragraph(lines: list[str]) -> str:
        merged = lines[0].strip()
        for raw_next in lines[1:]:
            nxt = raw_next.strip()
            if not nxt:
                continue

            # Keep hyphenated words natural when OCR/PDF wrapped at line end.
            if merged.endswith("-") and nxt and re.match(r"^[A-Za-zÀ-ỹ]", nxt):
                merged = merged[:-1] + nxt
                continue

            if re.search(r"[.!?][\"')\]]*$", merged):
                merged = f"{merged} {nxt}"
            else:
                merged = f"{merged} {nxt}"

        merged = re.sub(r"\s+([,.;:!?])", r"\1", merged)
        return re.sub(r"\s{2,}", " ", merged).strip()

    output: list[str] = []
    paragraph_buffer: list[str] = []
    in_code_block = False

    def _flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        output.append(_merge_paragraph(paragraph_buffer))
        paragraph_buffer = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            _flush_paragraph()
            output.append(line)
            in_code_block = not in_code_block
            continue

        if in_code_block:
            output.append(line)
            continue

        if not stripped:
            _flush_paragraph()
            if output and output[-1] != "":
                output.append("")
            continue

        if _is_structural_line(line):
            _flush_paragraph()
            output.append(line)
            continue

        paragraph_buffer.append(line)

    _flush_paragraph()
    return "\n".join(output).strip()


def _normalize_line_for_compare(value: str) -> str:
    """Normalize a line for robust duplicate comparison."""
    normalized = (value or "").strip().lower()
    normalized = re.sub(r"\s+", " ", normalized)
    normalized = re.sub(r"[^\w\s]", "", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def _normalize_math_formulas(text: str) -> str:
    """Normalize OCR math expressions into inline LaTeX without changing meaning."""
    if not text:
        return ""

    # Shield URLs to prevent them from being corrupted by math/fraction regexes
    url_pattern = re.compile(r'\bhttps?://\S+|\bwww\.\S+|\b[A-Za-z0-9.-]+\.[A-Za-z]{2,6}/\S+')
    urls = []
    def url_shield(match):
        urls.append(match.group(0))
        return f"__URL_PLACEHOLDER_{len(urls)-1}__"

    shielded_text = url_pattern.sub(url_shield, text)

    repaired = _repair_math_symbol_glyphs(shielded_text)
    repaired = _normalize_set_builder_notation(repaired)

    lines = repaired.splitlines()
    normalized_lines = [_normalize_math_line(line) for line in lines]
    normalized_text = "\n".join(normalized_lines)

    # Restore URLs
    for idx, url in enumerate(urls):
        normalized_text = normalized_text.replace(f"__URL_PLACEHOLDER_{idx}__", url)

    return normalized_text


def _count_private_use_chars(text: str) -> int:
    """Count Unicode private-use glyphs often seen in broken PDF math extraction."""
    if not text:
        return 0
    return sum(1 for ch in text if "\uE000" <= ch <= "\uF8FF")


def _repair_math_symbol_glyphs(text: str) -> str:
    """Repair common OCR-corrupted math symbols (especially private-use glyphs)."""
    if not text:
        return ""

    translation = str.maketrans(
        {
            "": "=",
            "": "{",
            "": "}",
            "": "-",
            "": "+",
            "": "∈",
            "": "≠",
            "": "⊂",
            "": "{",
            "": "{",
            "": "}",
            "": "}",
            "": " ",
        }
    )
    repaired = text.translate(translation)

    # Normalize relation symbols to LaTeX-safe tokens.
    repaired = repaired.replace("∈", r"\in")
    repaired = repaired.replace("≠", r"\ne")
    repaired = repaired.replace("⊂", r"\subset")

    return repaired


def _normalize_set_builder_notation(text: str) -> str:
    """Normalize OCR-broken set notation while preserving original meaning."""
    if not text:
        return ""

    normalized = text

    # Q set-builder often breaks into 3 lines with decorative braces in OCR.
    normalized = re.sub(
        r"(?is)\{\s*a\s*}\s*\n\s*Q\s*=\s*,\s*a\s*\\in\s*Z\s*,\s*b\s*\\in\s*Z\s*,\s*b\s*\\ne\s*0\s*:\s*([^\n]+)\n\s*\{\s*b\s*}",
        r"Q = {a/b, a \\in Z, b \\in Z, b \\ne 0}: \1",
        normalized,
    )

    # Canonicalize simple set literals: N = { ... } / Z = { ... } / etc.
    normalized = re.sub(r"(?m)^\s*([NZQRC])\s*=\s*\{\s*", r"\1 = \\{", normalized)
    normalized = re.sub(r"(?m)\s*}\s*:", r"\\}:", normalized)
    normalized = re.sub(r"(?m)^\s*R\s*:\s*", r"$R$: ", normalized)

    return normalized


def _normalize_math_line(line: str) -> str:
    if not line or not line.strip():
        return line

    parts = re.split(r"(\$[^$]+\$)", line)
    processed: list[str] = []

    for part in parts:
        if not part:
            continue
        if part.startswith("$") and part.endswith("$"):
            processed.append(part)
            continue

        segment = _replace_math_ocr_patterns(part)
        segment = _wrap_inline_math_segments(segment)
        processed.append(segment)

    normalized = "".join(processed)
    normalized = re.sub(r"\s+([,.;:])", r"\1", normalized)
    return normalized


def _replace_math_ocr_patterns(segment: str) -> str:
    superscript_map = {
        "⁰": "0",
        "¹": "1",
        "²": "2",
        "³": "3",
        "⁴": "4",
        "⁵": "5",
        "⁶": "6",
        "⁷": "7",
        "⁸": "8",
        "⁹": "9",
    }

    def superscript_repl(match: re.Match[str]) -> str:
        base = match.group(1)
        sup = match.group(2)
        exp = "".join(superscript_map[ch] for ch in sup if ch in superscript_map)
        return f"{base}^{exp}" if exp else match.group(0)

    segment = re.sub(r"([A-Za-z0-9\)])([⁰¹²³⁴⁵⁶⁷⁸⁹]+)", superscript_repl, segment)
    segment = re.sub(r"\b([A-Za-z])\s*([23])\b", r"\1^\2", segment)
    segment = re.sub(r"√\s*\(([^()]{1,80})\)", r"\\sqrt{\1}", segment)
    segment = re.sub(r"√\s*([A-Za-z0-9]+)", r"\\sqrt{\1}", segment)

    def frac_repl(match: re.Match[str]) -> str:
        num = (match.group(1) or "").strip()
        den = (match.group(2) or "").strip()
        if not num or not den:
            return match.group(0)

        # Avoid converting date-like numeric tokens (e.g., 12/2024).
        if num.isdigit() and den.isdigit() and (len(num) > 2 or len(den) > 2):
            return match.group(0)

        # Avoid converting text slashes (e.g., Git/GitHub, yes/no, input/output, TCP/IP, x-axis/y-axis)
        def is_math_expr(expr: str) -> bool:
            words = re.findall(r"[A-Za-z]+", expr)
            math_words = {
                "dx", "dy", "dt", "dz", "dr", "df", "dg", "dp", "dq", "ds", "dtheta", "du", "dv", "dw",
                "sin", "cos", "tan", "cot", "sec", "csc", "log", "ln", "exp", "lim", "max", "min",
                "pi", "theta", "phi", "psi", "omega", "alpha", "beta", "gamma", "delta", "lambda", "sigma", "sqrt"
            }
            for w in words:
                if len(w) > 1 and w.lower() not in math_words:
                    return False
            return True

        if not is_math_expr(num) or not is_math_expr(den):
            return match.group(0)

        return f"\\frac{{{num}}}{{{den}}}"

    segment = re.sub(
        r"(?<![\\\w])([A-Za-z0-9()+\-^]{1,20})\s*/\s*([A-Za-z0-9()+\-^]{1,20})(?![\w/])",
        frac_repl,
        segment,
    )

    return segment


def _wrap_inline_math_segments(segment: str) -> str:
    if not segment or "$" in segment:
        return segment

    wrapped = segment

    # Wrap equation-like expressions first (e.g., i^2 = -1, x^2 + y^2 = z^2).
    def equation_repl(match: re.Match[str]) -> str:
        expr = (match.group(1) or "").strip()
        letters = len(re.findall(r"[A-Za-z]", expr))
        if letters == 0:
            return match.group(0)
        if letters > 20:
            return match.group(0)
        return f"${expr}$"

    wrapped = re.sub(
        r"(?<!\w)([A-Za-z0-9\\{}^+\-*/().,= ]{1,160}\s*(?:=|≤|≥|<|>|≠)\s*[A-Za-z0-9\\{}^+\-*/().,= ]{1,160})(?!\w)",
        equation_repl,
        wrapped,
    )

    # When a whole equation is already wrapped, avoid nested wrapping of subparts.
    if "$" in wrapped:
        return wrapped

    # Wrap standalone formula tokens.
    wrapped = re.sub(r"(?<![$\\])(\\frac\{[^{}]+\}\{[^{}]+\})(?![$\w])", r"$\1$", wrapped)
    wrapped = re.sub(r"(?<![$\\])(\\sqrt\{[^{}]+\})(?![$\w])", r"$\1$", wrapped)
    wrapped = re.sub(r"(?<![$\\])\b([A-Za-z]\^\d+)\b(?![$\w])", r"$\1$", wrapped)

    # Wrap set and relation expressions using LaTeX relation operators.
    wrapped = re.sub(
        r"(?<![$\\])\b([A-Za-z](?:\s*\\(?:in|subset|ne)\s*[A-Za-z0-9])+[A-Za-z0-9\s]*)\b(?![$\\])",
        lambda m: f"${m.group(1).strip()}$",
        wrapped,
    )

    # Wrap simple set definitions like N = \{...\}.
    wrapped = re.sub(
        r"(?<![$\\])\b([NZQRC]\s*=\s*\\\{[^\n{}]{1,120}\\\})(?![$\\])",
        lambda m: f"${m.group(1).strip()}$",
        wrapped,
    )

    return wrapped


def _looks_like_math_line(line: str) -> bool:
    """Detect whether a line primarily represents a mathematical expression."""
    if not line:
        return False

    if re.search(r"\$(?:[^$]+)\$", line):
        return True
    if re.search(r"\\frac\{|\\sqrt\{|\^", line):
        return True

    has_operator = bool(re.search(r"(?:=|≤|≥|<|>|≠|\+|\-|/|\*)", line))
    has_symbolic_token = bool(re.search(r"\b[A-Za-z]\b|\d", line))
    return has_operator and has_symbolic_token


def _fix_common_vietnamese_ocr_errors(text: str) -> str:
    """Apply conservative OCR error corrections for common Vietnamese patterns."""
    fixed = text
    for pattern, replacement in _VIETNAMESE_OCR_REGEX_FIXES:
        fixed = re.sub(pattern, replacement, fixed, flags=re.IGNORECASE)

    # Token-level fixes for frequent OCR corruption in Vietnamese educational texts.
    for wrong, right in _VIETNAMESE_OCR_WORD_FIXES:
        fixed = re.sub(
            rf"(?<!\w){re.escape(wrong)}(?!\w)",
            right,
            fixed,
            flags=re.IGNORECASE,
        )

    # Single-token replacements that are common in OCR from teaching slides.
    fixed = re.sub(r"\bs6\b", "số", fixed, flags=re.IGNORECASE)
    fixed = re.sub(r"\bva\b", "và", fixed, flags=re.IGNORECASE)
    fixed = re.sub(r"\bki\s*hi[eé]u\s+la\s*([A-Za-z])\b", r"kí hiệu là \1", fixed, flags=re.IGNORECASE)

    # Normalize accidental OCR heading markers like '## Theo hu6ng ...' inside body text.
    fixed = re.sub(r"(?m)^\s*##\s+(theo\b.+)$", r"\1", fixed, flags=re.IGNORECASE)

    # Phrase-level corrections for frequent educational OCR mistakes.
    fixed = re.sub(r"\bđiều\s+hoa\b", "điều hòa", fixed, flags=re.IGNORECASE)
    fixed = re.sub(r"\bvị\s+trí\s+ca\s+theo\s+hướng\s+cü\b", "vị trí cũ theo hướng cũ", fixed, flags=re.IGNORECASE)
    fixed = re.sub(r"\blá\s+khoảng\b", "là khoảng", fixed, flags=re.IGNORECASE)
    fixed = re.sub(r"\blà\s+s\s+dao\s+động\b", "là số dao động", fixed, flags=re.IGNORECASE)
    fixed = re.sub(r"\btần\s+s\s+là\b", "tần số là", fixed, flags=re.IGNORECASE)

    # Normalize compact heading numbering: "1.chu kì" -> "1. chu kì".
    fixed = re.sub(r"(?m)^\s*(\d+)\.(?=\S)", r"\1. ", fixed)

    return fixed


def _merge_broken_ocr_lines(lines: list[str]) -> list[str]:
    """Merge hard-wrapped OCR lines into paragraphs when safe."""
    if not lines:
        return []

    merged: list[str] = []
    idx = 0
    while idx < len(lines):
        current = lines[idx].strip()
        if not current:
            if merged and merged[-1] != "":
                merged.append("")
            idx += 1
            continue

        # Keep markdown headings and list items as standalone lines.
        if current.startswith("#") or current.startswith("- "):
            merged.append(current)
            idx += 1
            continue

        while idx + 1 < len(lines):
            nxt = lines[idx + 1].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("- "):
                break

            current_ends_sentence = bool(re.search(r"[.!?:;)]$", current))
            next_looks_title = bool(re.match(r"^\d+(?:\.\d+){0,3}[\s.:)]", nxt))
            next_starts_lower = bool(re.match(r"^[a-zà-ỹ]", nxt))

            if current_ends_sentence or next_looks_title or not next_starts_lower:
                break

            current = f"{current} {nxt}"
            idx += 1

        merged.append(current)
        idx += 1

    return merged


def _deduplicate_ocr_lines(text: str) -> str:
    """Remove repeated OCR lines while preserving original order and structure."""
    output: list[str] = []
    seen_keys: set[str] = set()
    previous_key = ""

    for raw in text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            if output and output[-1] != "":
                output.append("")
            continue

        key = _normalize_line_for_compare(stripped)
        if key and key == previous_key:
            continue

        # Keep very short lines even if repeated (list markers, section counters).
        if key and len(key) >= 12:
            if key in seen_keys:
                continue
            seen_keys.add(key)

        output.append(stripped)
        previous_key = key

    return "\n".join(output).strip()


def _normalize_heading_case(text: str) -> str:
    """Capitalize heading lines while keeping content untouched."""
    normalized_lines: list[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("#"):
            marker, _, rest = line.partition(" ")
            heading = rest.strip()
            if heading:
                heading = heading[0].upper() + heading[1:]
                normalized_lines.append(f"{marker} {heading}")
            else:
                normalized_lines.append(line)
            continue
        normalized_lines.append(line)
    return "\n".join(normalized_lines).strip()


def _looks_unrecoverable_ocr_text(text: str) -> bool:
    """Detect heavily corrupted OCR text that cannot be reliably recovered."""
    if not text or len(text.strip()) < 12:
        return True

    tokens = re.findall(r"[A-Za-zÀ-ỹ0-9]+", text)
    if not tokens:
        return True

    alpha_tokens = [tk for tk in tokens if re.search(r"[A-Za-zÀ-ỹ]", tk)]
    if not alpha_tokens:
        return True

    meaningful_tokens = [tk for tk in alpha_tokens if len(tk) >= 2]
    if len(meaningful_tokens) < 3 and len(text) < 80:
        return True

    weird_chars = len(re.findall(r"[^\w\sÀ-ỹ,.;:!?()\-#|]", text))
    weird_ratio = weird_chars / max(1, len(text))
    alphabetic_chars = len(re.findall(r"[A-Za-zÀ-ỹ]", text))
    alphabetic_ratio = alphabetic_chars / max(1, len(text))

    return weird_ratio > 0.52 and alphabetic_ratio < 0.45


@router.post("/convert")
@router.post("/upload")

async def convert_document(
    file: UploadFile = File(...),
    advanced: bool = Query(False, description="Use advanced header/footer detection"),
    cleaning_mode: str = Query("standard", description="Cleaning mode: standard, advanced, or divider"),
    ocr_mode: str = Query(
        "auto",
        description="Deprecated. OCR mode is always automatic.",
    ),
) -> ConvertResponse:
    """
    Stage 1: Convert uploaded document to cleaned Markdown
    
    This endpoint:
    1. Validates the uploaded file
    2. Extracts text based on file type (PDF or DOCX)
    3. Cleans the extracted text using markdown_cleaner
    4. Returns the cleaned Markdown
    
    Args:
        file: Uploaded file (PDF or DOCX)
        advanced: DEPRECATED - Use cleaning_mode instead
        cleaning_mode: Cleaning strategy (default: "standard")
            - "standard": Aggressive noise removal
            - "advanced": Header/footer detection with frequency analysis
            - "divider": Page structure-based extraction (for divider-separated pages)
    
    Returns:
        ConvertResponse with markdown text and metadata
        
    Raises:
        HTTPException: If file validation or conversion fails
    
    Example:
        # Standard cleaning
        POST /documents/convert?file=document.pdf
        
        # Divider-based cleaning for special PDF format
        POST /documents/convert?file=document.pdf&cleaning_mode=divider
        
        # Advanced cleaning with header/footer detection
        POST /documents/convert?file=document.pdf&cleaning_mode=advanced
    """
    import time
    start_time = time.time()
    
    try:
        # Validate file type
        file_ext = os.path.splitext(file.filename or '')[1].lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {ALLOWED_EXTENSIONS}"
            )

        # Read file content
        file_content = await file.read()
        file_size = len(file_content)

        # Validate file size
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE / 1024 / 1024:.0f}MB"
            )

        document_id = uuid.uuid4().hex
        stored_file_name = f"{document_id}{file_ext}"
        stored_file_path = os.path.join(UPLOAD_DIR_PATH, stored_file_name)
        with open(stored_file_path, "wb") as stored_file:
            stored_file.write(file_content)

        # Extract text based on file type
        ocr_quality: Literal["good", "medium", "bad"] = "good"
        ocr_used = False
        if file_ext == '.pdf':
            extracted_text, pages, ocr_quality, ocr_used = _extract_from_pdf_with_meta(
                file_content,
                ocr_mode="auto",
            )
            extraction_method = "pdfplumber"
        elif file_ext == '.docx':
            extracted_text = _extract_from_docx(file_content)
            pages = 0
            extraction_method = "python-docx"
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type"
            )

        # Validate extraction
        if not extracted_text or len(extracted_text.strip()) < 10:
            detail = "Could not extract text from file. File may be empty or corrupted."
            if file_ext == ".pdf":
                detail = OCR_UNCLEAR_MESSAGE
            raise HTTPException(
                status_code=422,
                detail=detail,
            )

        # Repair broken math glyphs early so cleaner does not drop formula-heavy lines.
        extracted_text = _repair_math_symbol_glyphs(extracted_text)

        # Clean markdown using selected method
        # Support both old 'advanced' param and new 'cleaning_mode' param
        if cleaning_mode not in ["standard", "advanced", "divider"]:
            cleaning_mode = "advanced" if advanced else "standard"
        
        logger.info(f"Cleaning markdown for {file.filename} (method: {cleaning_mode})")
        
        if cleaning_mode == "divider":
            # Use divider-based cleaner
            try:
                cleaner = AdvancedMarkdownCleaner()
                cleaned_text, page_contents = cleaner.clean_markdown_divider_based(extracted_text)
                cleaning_method = "divider"
                noise_removed_ratio = 0.0  # Will calculate if needed
                logger.info(f"Divider cleaner: extracted {len(page_contents)} pages, output {len(cleaned_text)} chars")
            except Exception as e:
                logger.error(f"Divider cleaner error: {str(e)}", exc_info=True)
                raise HTTPException(
                    status_code=500,
                    detail=f"Divider cleaner error: {str(e)}"
                )
        elif cleaning_mode == "advanced":
            cleaned_text, metadata = clean_markdown_advanced(extracted_text, debug=False)
            cleaning_method = "advanced"
            noise_removed_ratio = metadata.cleaning_ratio if metadata.total_lines_removed > 0 else 0.0
        else:
            cleaned_text, metadata = clean_markdown(extracted_text, debug=False)
            cleaning_method = "standard"
            noise_removed_ratio = metadata.cleaning_ratio if metadata.total_lines_removed > 0 else 0.0

        cleaned_text = _sanitize_markdown_for_output(cleaned_text)
        cleaned_text = _normalize_math_formulas(cleaned_text)
        cleaned_text = _promote_markdown_headings(cleaned_text)
        cleaned_text = _normalize_paragraph_line_breaks(cleaned_text)

        preview_url = (
            f"/documents/convert/preview/{document_id}"
            if file_ext == ".pdf"
            else None
        )

        CONVERSION_SESSIONS[document_id] = {
            "document_id": document_id,
            "markdown": cleaned_text,
            "file_name": file.filename or "document",
            "cleaning_method": cleaning_method,
            "pages": pages,
            "preview_url": preview_url,
            "stored_file_path": stored_file_path,
        }

        conversion_time_ms = (time.time() - start_time) * 1000

        logger.info(
            f"Successfully converted {file.filename} "
            f"({file_size} bytes, {pages} pages, "
            f"{len(cleaned_text)} chars output, "
            f"{cleaning_method} cleaning, "
            f"{noise_removed_ratio:.1f}% noise removed)"
        )

        return ConvertResponse(
            success=True,
            markdown=cleaned_text,
            file_name=file.filename or "document",
            file_size=file_size,
            extraction_method=extraction_method,
            cleaning_method=cleaning_method,
            document_id=document_id,
            preview_url=preview_url,
            pages=pages,
            noise_removed_ratio=noise_removed_ratio,
            quality=ocr_quality,
            ocr_used=ocr_used,
            conversion_time_ms=conversion_time_ms,
            message=f"Successfully converted {file.filename} to Markdown"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error converting document: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file: {str(e)}"
        )


@router.get("/convert/session/{document_id}")
async def get_convert_session(document_id: str) -> ConvertSessionResponse:
    """Restore a previously converted session to support browser reload/resume."""
    session = CONVERSION_SESSIONS.get(document_id)
    if not session:
        raise HTTPException(status_code=404, detail="Converted session not found")

    return ConvertSessionResponse(
        success=True,
        document_id=document_id,
        markdown=session["markdown"],
        file_name=session["file_name"],
        cleaning_method=session["cleaning_method"],
        pages=session.get("pages", 0),
        preview_url=session.get("preview_url"),
        message="Converted session restored",
    )


@router.get("/convert/preview/{document_id}")
async def get_convert_preview(document_id: str):
    """Return original uploaded file (PDF) for preview pane after reload."""
    session = CONVERSION_SESSIONS.get(document_id)
    if not session:
        raise HTTPException(status_code=404, detail="Preview session not found")

    file_path = session.get("stored_file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Preview file not found")

    if not str(file_path).lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Preview is only available for PDF files")

    return FileResponse(
        path=file_path,
        media_type="application/pdf",
        filename=session.get("file_name") or "document.pdf",
        content_disposition_type="inline",
    )


def _extract_from_pdf(file_content: bytes, ocr_mode: str = "auto") -> tuple[str, int]:
    extracted_text, page_count, _, _ = _extract_from_pdf_with_meta(file_content, ocr_mode=ocr_mode)
    return extracted_text, page_count


def _extract_from_pdf_with_meta(
    file_content: bytes,
    ocr_mode: str = "auto",
) -> tuple[str, int, Literal["good", "medium", "bad"], bool]:
    """
    Extract text from PDF content
    
    Args:
        file_content: Bytes of PDF file
    
    Returns:
        Tuple of (extracted_text, page_count, quality, ocr_used)
    """
    import io

    mode = (ocr_mode or "auto").strip().lower()
    if mode not in OCR_MODE_VALUES:
        mode = "auto"
    
    try:
        pdf_file = io.BytesIO(file_content)
        page_blocks: list[list[str]] = []
        page_native_chars: list[int] = []
        page_private_use_chars: list[int] = []
        pages_with_native_content = 0
        page_count = 0
        native_text_chars = 0

        with pdfplumber.open(pdf_file) as pdf:
            page_count = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                logger.debug(f"Extracting text from page {page_num}")

                # Explicit page marker helps downstream repeated header/footer detection.
                current_block: list[str] = [f"## Page {page_num}/{page_count}"]
                current_page_native_chars = 0
                current_page_private_use_chars = 0

                custom_table_settings = {
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "intersection_tolerance": 15,
                    "join_tolerance": 15
                }
                
                # Find tables with coordinates and sort them vertically (top-to-bottom)
                tables_found = page.find_tables(table_settings=custom_table_settings)
                if not tables_found:
                    tables_found = page.find_tables()
                
                tables_found = sorted(tables_found, key=lambda t: t.bbox[1]) if tables_found else []
                
                y_current = 0
                width = page.width
                height = page.height
                
                # We collect segments as (type, content) tuples
                segments = []
                
                for table_obj in tables_found:
                    x0, y0, x1, y1 = table_obj.bbox
                    
                    # Extract text before this table
                    if y0 > y_current:
                        text_area = page.crop((0, y_current, width, y0))
                        txt = text_area.extract_text() or ""
                        if txt.strip():
                            segments.append(("text", txt.strip()))
                    
                    # Extract table data
                    table_data = table_obj.extract()
                    if table_data:
                        segments.append(("table", table_data))
                        
                    y_current = y1
                
                # Extract remaining text after the last table
                if y_current < height:
                    text_area = page.crop((0, y_current, width, height))
                    txt = text_area.extract_text() or ""
                    if txt.strip():
                        segments.append(("text", txt.strip()))
                
                table_row_signatures: set[str] = set()
                
                # Pass 1: Build signatures of all table rows to filter text duplicates
                for seg_type, content in segments:
                    if seg_type == "table":
                        for row in content:
                            if row:
                                row_clean = [(cell or "").strip().replace("\n", " ") for cell in row]
                                row_signature = _normalize_line_for_compare(" ".join(cell for cell in row_clean if cell))
                                if row_signature:
                                    table_row_signatures.add(row_signature)
                
                # Pass 2: Process segments in order and append to current_block
                for seg_type, content in segments:
                    if seg_type == "table":
                        rows = [[(cell or "").strip().replace("\n", " ") for cell in row] for row in content if row]
                        if not rows:
                            continue
                        header = rows[0]
                        body = rows[1:] if len(rows) > 1 else []
                        col_count = max(1, len(header))
                        
                        header_line = "| " + " | ".join(header) + " |"
                        divider_line = "|" + "|".join(["---"] * col_count) + "|"
                        current_block.append(header_line)
                        current_block.append(divider_line)
                        
                        native_text_chars += len(header_line) + len(divider_line)
                        current_page_native_chars += len(header_line) + len(divider_line)
                        current_page_private_use_chars += _count_private_use_chars(header_line)
                        current_page_private_use_chars += _count_private_use_chars(divider_line)
                        
                        for row in body:
                            if len(row) < col_count:
                                row = row + [""] * (col_count - len(row))
                            current_row = row[:col_count]
                            row_line = "| " + " | ".join(current_row) + " |"
                            current_block.append(row_line)
                            native_text_chars += len(row_line)
                            current_page_native_chars += len(row_line)
                            current_page_private_use_chars += _count_private_use_chars(row_line)
                    else:
                        # Process text segment
                        text_lines = content.split('\n')
                        filtered_lines = []
                        for line in text_lines:
                            stripped = line.strip()
                            if stripped.count('|') >= 2:
                                continue
                            normalized_line = _normalize_line_for_compare(stripped)
                            if normalized_line and normalized_line in table_row_signatures:
                                continue
                            filtered_lines.append(line)
                            
                        filtered_text = '\n'.join(filtered_lines).strip()
                        if filtered_text:
                            if len(current_block) > 1 and current_block[-1] != "":
                                current_block.append("")
                            current_block.append(filtered_text)
                            native_text_chars += len(filtered_text)
                            current_page_native_chars += len(filtered_text)
                            current_page_private_use_chars += _count_private_use_chars(filtered_text)

                page_blocks.append(current_block)
                page_native_chars.append(current_page_native_chars)
                page_private_use_chars.append(current_page_private_use_chars)
                if current_page_native_chars >= 35:
                    pages_with_native_content += 1

        # OCR fallback for image-based/scan PDFs.
        # auto: run OCR only on candidate pages (sparse/corrupted)
        # on: force OCR pass on all pages
        target_ocr_pages: list[int] = []
        should_run_ocr = mode == "on"

        if mode == "on":
            target_ocr_pages = list(range(page_count))
        elif mode == "auto":
            sparse_threshold = max(80, page_count * 20)
            corrupted_symbol_threshold = max(3, page_count)
            has_corrupted_symbols = sum(page_private_use_chars) >= corrupted_symbol_threshold
            # Skip OCR for clearly text-native PDFs to keep conversion fast.
            clearly_text_native = (
                page_count > 0
                and pages_with_native_content == page_count
                and native_text_chars >= max(35 * page_count, 60)
                and not has_corrupted_symbols
            )

            if not clearly_text_native:
                for idx in range(page_count):
                    sparse_page = page_native_chars[idx] < 80
                    corrupted_page = page_private_use_chars[idx] >= 2
                    if sparse_page or corrupted_page:
                        target_ocr_pages.append(idx)

                # Safety net: if global text is too sparse, OCR all pages.
                if native_text_chars < sparse_threshold and not target_ocr_pages:
                    target_ocr_pages = list(range(page_count))

                should_run_ocr = len(target_ocr_pages) > 0

        ocr_quality: Literal["good", "medium", "bad"] = "good"
        ocr_used = False
        if mode != "off" and should_run_ocr:
            file_fingerprint = hashlib.sha256(file_content).hexdigest()
            ocr_pages, detected_ocr_quality, _ = _extract_text_from_pdf_ocr(
                file_content,
                page_count,
                target_page_indexes=target_ocr_pages,
                file_fingerprint=file_fingerprint,
            )
            ocr_applied_pages = 0
            for idx, ocr_text in enumerate(ocr_pages):
                if not ocr_text:
                    continue
                if idx >= len(page_blocks):
                    continue

                # Keep OCR only for pages where native extraction is effectively empty.
                has_native_content = any(
                    line.strip() and not line.startswith("## Page ")
                    for line in page_blocks[idx]
                )

                # In force mode, still avoid mixing OCR into pages that already have rich native text.
                if has_native_content and page_native_chars[idx] >= 80:
                    continue

                page_blocks[idx].append(ocr_text)
                native_text_chars += len(ocr_text)
                ocr_applied_pages += 1

            if ocr_applied_pages > 0:
                ocr_quality = detected_ocr_quality
                ocr_used = True
                logger.info(
                    "Applied OCR fallback for %s/%s PDF pages (requested=%s)",
                    ocr_applied_pages,
                    page_count,
                    len(target_ocr_pages) if target_ocr_pages else page_count,
                )

        all_parts: list[str] = []
        for block in page_blocks:
            all_parts.extend(block)
            all_parts.append("")

        extracted_text = "\n".join(all_parts).strip()
        logger.info(
            f"Extracted {len(extracted_text)} characters from {page_count} PDF pages "
            f"(with page markers and table preservation)"
        )
        return extracted_text, page_count, ocr_quality, ocr_used

    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise


def _extract_text_from_pdf_ocr(
    file_content: bytes,
    page_count: int,
    target_page_indexes: list[int] | None = None,
    file_fingerprint: str | None = None,
) -> tuple[list[str], Literal["good", "medium", "bad"], bool]:
    """OCR fallback for image-based PDF pages.

    Returns a list with one OCR text entry per page (may be empty strings).
    """
    if pdfium is None:
        logger.warning("OCR dependency pypdfium2 is not available")
        return [""] * max(page_count, 0), "bad", False

    try:
        ocr_results: list[str] = []
        page_qualities: list[Literal["good", "medium", "bad"]] = []
        any_ocr_used = False
        pdf_doc = pdfium.PdfDocument(file_content)

        target_set = set(target_page_indexes or [])
        cache_hits = 0
        cache_misses = 0

        # Process pages sequentially to guarantee thread-safety for pdfium and PaddleOCR,
        # and to prevent OOM/memory spikes from concurrent page rendering.
        results_map = {}
        for i in range(len(pdf_doc)):
            if target_set and i not in target_set:
                results_map[i] = ("", "good", False)
                continue

            cached = _load_cached_ocr_page(
                file_fingerprint=file_fingerprint,
                page_index=i,
            )
            if cached is not None:
                results_map[i] = (
                    cached.get("text", ""),
                    cached.get("quality", "good"),
                    bool(cached.get("ocr_used", True))
                )
                cache_hits += 1
                continue

            cache_misses += 1
            try:
                page = pdf_doc[i]
                # Scale 3.0 is a good balance between speed and quality
                bitmap = page.render(scale=3.0)
                image = bitmap.to_pil()
                normalized, quality, ocr_used = _extract_page_ocr_best(image)
                if quality == "bad":
                    normalized = ""
                
                _save_cached_ocr_page(
                    file_fingerprint=file_fingerprint,
                    page_index=i,
                    text=normalized,
                    quality=quality,
                    ocr_used=ocr_used,
                )
                results_map[i] = (normalized, quality, ocr_used)
            except Exception as page_err:
                logger.warning("OCR failed for page %s: %s", i + 1, page_err)
                results_map[i] = ("", "bad", False)

        ocr_results = []
        page_qualities = []
        for i in range(len(pdf_doc)):
            text, qual, used = results_map.get(i, ("", "bad", False))
            ocr_results.append(text)
            page_qualities.append(qual)
            any_ocr_used = any_ocr_used or used

        if cache_hits > 0:
            logger.info("OCR cache hit: %s pages (miss=%s)", cache_hits, cache_misses)

        return ocr_results, _merge_ocr_quality(page_qualities), any_ocr_used
    except Exception as err:
        logger.warning("OCR fallback failed: %s", err)
        return [""] * max(page_count, 0), "bad", False


def _ocr_cache_file_path(file_fingerprint: str, page_index: int) -> str:
    cache_key = f"v1:{file_fingerprint}:{page_index}"
    cache_name = hashlib.sha256(cache_key.encode("utf-8")).hexdigest() + ".json"
    return os.path.join(OCR_CACHE_DIR_PATH, cache_name)


def _load_cached_ocr_page(file_fingerprint: str | None, page_index: int) -> dict[str, Any] | None:
    if not OCR_CACHE_ENABLED or not file_fingerprint:
        return None

    cache_path = _ocr_cache_file_path(file_fingerprint, page_index)
    if not os.path.exists(cache_path):
        return None

    try:
        with open(cache_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        text = str(data.get("text", ""))
        quality = data.get("quality", "good")
        if quality not in {"good", "medium", "bad"}:
            quality = "good"
        return {
            "text": text,
            "quality": quality,
            "ocr_used": bool(data.get("ocr_used", bool(text.strip()))),
        }
    except Exception as cache_err:
        logger.debug("Cannot read OCR cache for page %s: %s", page_index + 1, cache_err)
        return None


def _save_cached_ocr_page(
    file_fingerprint: str | None,
    page_index: int,
    text: str,
    quality: Literal["good", "medium", "bad"],
    ocr_used: bool,
) -> None:
    if not OCR_CACHE_ENABLED or not file_fingerprint:
        return

    cache_path = _ocr_cache_file_path(file_fingerprint, page_index)
    payload = {
        "text": text,
        "quality": quality,
        "ocr_used": ocr_used,
        "updated_at": datetime.utcnow().isoformat() + "Z",
    }

    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False)
    except Exception as cache_err:
        logger.debug("Cannot write OCR cache for page %s: %s", page_index + 1, cache_err)


def _extract_page_ocr_best(image) -> tuple[str, Literal["good", "medium", "bad"], bool]:
    """Run OCR with adaptive strategy and keep the best-scoring candidate."""
    document_type = _classify_ocr_document_type(image)
    if document_type == "handwritten":
        return "", "bad", False

    variants = _build_ocr_image_variants(image, document_type=document_type)
    fallback_configs = [
        "--oem 1 --psm 6 -c preserve_interword_spaces=1",
        "--oem 1 --psm 4 -c preserve_interword_spaces=1",
        "--oem 1 --psm 3 -c preserve_interword_spaces=1",
    ]

    best_text = ""
    best_score = -1.0
    best_quality: Literal["good", "medium", "bad"] = "bad"
    any_ocr_used = False

    for variant in variants:
        for region in _build_ocr_regions(variant):
            raw_text, avg_conf, paddle_used = _run_paddle_ocr_with_confidence(region)
            quality = _evaluate_ocr_quality(raw_text, avg_conf)
            any_ocr_used = any_ocr_used or paddle_used

            # Fallback to tesseract only when PaddleOCR is unavailable.
            if not raw_text and not paddle_used:
                for config in fallback_configs:
                    fallback_text, fallback_conf = _run_ocr_with_confidence(region, config)
                    fallback_quality = _evaluate_ocr_quality(fallback_text, fallback_conf)
                    if _quality_rank(fallback_quality) >= _quality_rank(quality):
                        raw_text = fallback_text
                        avg_conf = fallback_conf
                        quality = fallback_quality

            cleaned = _post_process_ocr_text(raw_text)
            score = _score_ocr_candidate(cleaned, avg_conf) + (_quality_rank(quality) * 8.0)

            if score > best_score:
                best_score = score
                best_text = cleaned
                best_quality = quality
            
            # Early exit if we already found a high-quality candidate
            if best_quality == "good" and len(best_text.strip()) >= 50:
                return best_text, best_quality, any_ocr_used

    if best_quality == "bad":
        return "", "bad", any_ocr_used
    return best_text, best_quality, any_ocr_used


def _classify_ocr_document_type(image) -> Literal["printed", "noisy", "handwritten"]:
    """Classify image quality to select OCR strategy."""
    if np is None:
        return "printed"

    gray_arr = np.array(image.convert("L"))
    contrast = float(gray_arr.std()) if hasattr(gray_arr, "std") else 0.0
    blur_metric = 100.0
    if cv2 is not None:
        try:
            blur_metric = float(cv2.Laplacian(gray_arr, cv2.CV_64F).var())
        except Exception:
            blur_metric = 100.0

    if blur_metric < 10.0 and contrast < 20.0:
        return "handwritten"
    if blur_metric < 35.0 or contrast < 36.0:
        return "noisy"
    return "printed"


def _get_paddle_ocr_engine():
    global _PADDLE_OCR_ENGINE
    if PaddleOCR is None:
        return None
    if _PADDLE_OCR_ENGINE is None:
        _PADDLE_OCR_ENGINE = PaddleOCR(use_angle_cls=True, lang="vi", show_log=False)
    return _PADDLE_OCR_ENGINE


def _run_paddle_ocr_with_confidence(image) -> tuple[str, float, bool]:
    if np is None or PaddleOCR is None:
        return "", 0.0, False

    engine = _get_paddle_ocr_engine()
    if engine is None:
        return "", 0.0, False

    try:
        rgb_arr = np.array(image.convert("RGB"))
        result = engine.ocr(rgb_arr, cls=True)
        if not result:
            return "", 0.0, True

        lines = result[0] if isinstance(result, list) else result
        parsed: list[tuple[float, float, str, float]] = []
        confs: list[float] = []

        for line in lines or []:
            if not line or len(line) < 2:
                continue
            box = line[0]
            rec = line[1]
            text = str(rec[0] or "").strip()
            if not text:
                continue
            try:
                conf = float(rec[1])
            except Exception:
                conf = 0.0

            x_min = min(float(pt[0]) for pt in box) if box else 0.0
            y_min = min(float(pt[1]) for pt in box) if box else 0.0
            parsed.append((y_min, x_min, text, conf))
            if conf >= 0:
                confs.append(conf)

        if not parsed:
            return "", 0.0, True

        parsed.sort(key=lambda item: (item[0], item[1]))
        merged = "\n".join(item[2] for item in parsed).strip()
        avg_conf = (sum(confs) / len(confs)) if confs else 0.0
        return merged, avg_conf, True
    except Exception as err:
        logger.debug("PaddleOCR failed: %s", err)
        return "", 0.0, False


def _quality_rank(quality: Literal["good", "medium", "bad"]) -> int:
    if quality == "good":
        return 2
    if quality == "medium":
        return 1
    return 0


def _merge_ocr_quality(qualities: list[Literal["good", "medium", "bad"]]) -> Literal["good", "medium", "bad"]:
    if not qualities:
        return "bad"
    if any(item == "good" for item in qualities):
        return "good"
    if any(item == "medium" for item in qualities):
        return "medium"
    return "bad"


def _evaluate_ocr_quality(text: str, avg_conf: float) -> Literal["good", "medium", "bad"]:
    if not text:
        return "bad"

    words = re.findall(r"[A-Za-zÀ-ỹ0-9]{2,}", text)
    total_len = max(1, len(text))
    valid_chars = len(re.findall(r"[A-Za-zÀ-ỹ0-9\s,.;:!?()\-]", text))
    invalid_ratio = 1.0 - (valid_chars / total_len)
    single_tokens = len(re.findall(r"\b[A-Za-zÀ-ỹ0-9]\b", text))

    if avg_conf >= 0.75 and invalid_ratio <= 0.10 and len(words) >= 8 and single_tokens <= 6:
        return "good"
    if avg_conf >= 0.52 and invalid_ratio <= 0.22 and len(words) >= 4:
        return "medium"
    return "bad"


def _build_ocr_regions(image):
    """Return OCR region candidates to reduce noise from decorative slide zones."""
    regions = [image]

    try:
        width, height = image.size
    except Exception:
        return regions

    if width < 200 or height < 200:
        return regions

    # Many slide decks place text on center/right and illustrations on the left.
    right_crop = image.crop((int(width * 0.26), int(height * 0.06), int(width * 0.99), int(height * 0.97)))
    center_right_crop = image.crop((int(width * 0.38), int(height * 0.18), int(width * 0.99), int(height * 0.94)))
    regions.append(right_crop)
    regions.append(center_right_crop)

    return regions


def _build_ocr_image_variants(
    image,
    document_type: Literal["printed", "noisy", "handwritten"] = "printed",
):
    """Create image variants to improve OCR robustness on noisy slides/scans."""
    if document_type == "noisy":
        base = _preprocess_noisy_image(image)
    else:
        base = _preprocess_ocr_image(image)
    variants = [base]

    # Upscaled variant helps on low-resolution slide exports.
    if hasattr(base, "resize"):
        upscaled = base.resize((base.width * 2, base.height * 2))
        variants.append(upscaled)

    # Soft threshold variant improves low-contrast backgrounds.
    if hasattr(base, "point"):
        thresholded = base.point(lambda p: 255 if p > 165 else 0)
        variants.append(thresholded)

    return variants


def _preprocess_noisy_image(image):
    """Contrast + denoise preprocessing for noisy printed documents."""
    if Image is None or np is None or cv2 is None:
        return _preprocess_ocr_image(image)

    try:
        rgb = np.array(image.convert("RGB"))
        gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        denoised = cv2.fastNlMeansDenoising(enhanced, None, 11, 7, 21)
        return Image.fromarray(denoised)
    except Exception:
        return _preprocess_ocr_image(image)


def _run_ocr_with_confidence(image, config: str) -> tuple[str, float]:
    """Run OCR and return extracted text with average confidence."""
    if pytesseract is None:
        return "", 0.0

    langs = ["vie+eng", "eng"]
    last_err = None

    for lang in langs:
        try:
            text = pytesseract.image_to_string(image, lang=lang, config=config)
            avg_conf = _estimate_ocr_confidence(image, lang=lang, config=config)
            return text, avg_conf
        except Exception as err:
            last_err = err
            continue

    if last_err:
        logger.debug("OCR run failed for config '%s': %s", config, last_err)
    return "", 0.0


def _estimate_ocr_confidence(image, lang: str, config: str) -> float:
    """Estimate OCR confidence from tesseract word-level data."""
    if pytesseract is None or TesseractOutput is None:
        return 0.0

    try:
        data = pytesseract.image_to_data(
            image,
            lang=lang,
            config=config,
            output_type=TesseractOutput.DICT,
        )
        conf_values: list[float] = []
        for raw_conf in data.get("conf", []):
            try:
                conf = float(raw_conf)
            except Exception:
                continue
            if conf >= 0:
                conf_values.append(conf)
        if not conf_values:
            return 0.0
        return sum(conf_values) / len(conf_values)
    except Exception:
        return 0.0


def _score_ocr_candidate(text: str, avg_conf: float) -> float:
    """Score OCR output quality using confidence + text structure heuristics."""
    if not text:
        return 0.0

    words = re.findall(r"[A-Za-zÀ-ỹ0-9]{2,}", text)
    lines = [ln for ln in text.splitlines() if ln.strip()]
    alpha_chars = len(re.findall(r"[A-Za-zÀ-ỹ]", text))
    weird_chars = len(re.findall(r"[~`^_=*\\]{2,}|[{}<>]{2,}", text))
    non_text_chars = len(re.findall(r"[^\w\sÀ-ỹ,.;:!?()\-]", text))
    single_char_tokens = len(re.findall(r"\b[A-Za-zÀ-ỹ0-9]\b", text))
    gibberish_tokens = len(
        re.findall(r"\b(?![A-Za-zÀ-ỹ]*[AEIOUYaeiouyÀ-ỹ])[A-Za-zÀ-ỹ]{4,}\b", text)
    )
    vietnamese_chars = len(re.findall(r"[À-ỹ]", text))
    text_len = max(1, len(text))
    weird_ratio = non_text_chars / text_len

    # Weighted score favors readable text with good OCR confidence.
    return (
        (avg_conf * 1.8)
        + min(40.0, len(words) * 0.55)
        + min(20.0, len(lines) * 1.2)
        + min(30.0, alpha_chars * 0.03)
        + min(12.0, vietnamese_chars * 0.2)
        - (weird_chars * 2.2)
        - (weird_ratio * 120.0)
        - (single_char_tokens * 0.9)
        - (gibberish_tokens * 1.4)
    )


def _preprocess_ocr_image(image):
    """Improve OCR readability for slide-like pages before text recognition."""
    if ImageOps is None:
        return image

    # Grayscale + contrast normalization + sharpen improves slide OCR stability.
    processed = ImageOps.grayscale(image)
    processed = ImageOps.autocontrast(processed, cutoff=1)
    if ImageFilter is not None:
        processed = processed.filter(ImageFilter.SHARPEN)
    return processed


def _post_process_ocr_text(text: str | None) -> str:
    """Normalize noisy OCR output with lightweight Vietnamese-friendly cleanup."""
    if not text:
        return ""

    normalized = unicodedata.normalize("NFC", text)
    normalized = normalized.replace("\x0c", " ")
    normalized = re.sub(r"(?<=\w)\|(?=\w)", "I", normalized)
    normalized = re.sub(r"\bSOL\b", "SQL", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"\s*([,.;:])", r"\1", normalized)
    normalized = _normalize_math_formulas(normalized)
    normalized = _fix_common_vietnamese_ocr_errors(normalized)

    lines: list[str] = []
    for raw in normalized.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            continue

        if _looks_like_math_line(line):
            lines.append(line)
            continue

        line = re.sub(r"^[•●▪◦]\s*", "- ", line)
        if line.startswith("*") and len(line) > 1:
            line = "- " + line[1:].strip()

        # Drop obvious OCR noise lines.
        alpha_count = len(re.findall(r"[A-Za-zÀ-ỹ]", line))
        if alpha_count == 0 and not re.search(r"\d", line):
            continue

        punct_count = len(re.findall(r"[^\w\sÀ-ỹ]", line))
        if punct_count > 0 and punct_count >= max(4, int(len(line) * 0.55)):
            continue

        tokens = re.findall(r"[A-Za-zÀ-ỹ0-9]+", line)
        if tokens:
            single_token_count = sum(1 for token in tokens if len(token) == 1)
            if single_token_count >= 4 and single_token_count >= int(len(tokens) * 0.6):
                continue

        lines.append(line)

    if not lines:
        return ""

    # Promote slide-like title lines into headings.
    merged_lines = _merge_broken_ocr_lines(lines)

    promoted: list[str] = []
    for idx, line in enumerate(merged_lines):
        upper_ratio = _estimate_upper_ratio(line)
        is_numbered_title = bool(re.match(r"^\d+(?:\.\d+){0,3}\s+[A-Za-zÀ-ỹ].+", line))
        is_likely_title = (
            len(line) >= 10
            and len(line) <= 120
            and (upper_ratio >= 0.72 or is_numbered_title or line.endswith(":"))
        )

        if is_likely_title and not line.startswith("#"):
            heading_prefix = "#" if idx == 0 else "##"
            promoted.append(f"{heading_prefix} {line.rstrip(':')}")
        else:
            promoted.append(line)

    joined = "\n".join(promoted).strip()
    joined = _deduplicate_ocr_lines(joined)
    joined = _normalize_heading_case(joined)

    if _looks_unrecoverable_ocr_text(joined):
        return ""

    return joined


def _estimate_upper_ratio(line: str) -> float:
    letters = [ch for ch in line if ch.isalpha()]
    if not letters:
        return 0.0
    upper = [ch for ch in letters if ch.upper() == ch]
    return len(upper) / len(letters)


def _extract_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX content
    
    Args:
        file_content: Bytes of DOCX file
    
    Returns:
        Extracted text
    """
    import io
    
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract DOCX tables and preserve as markdown tables.
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(cells)

            if not rows:
                continue

            header = rows[0]
            body = rows[1:] if len(rows) > 1 else []
            col_count = max(1, len(header))

            parts.append("| " + " | ".join(header) + " |")
            parts.append("|" + "|".join(["---"] * col_count) + "|")

            for row in body:
                if len(row) < col_count:
                    row = row + [""] * (col_count - len(row))
                parts.append("| " + " | ".join(row[:col_count]) + " |")

        extracted_text = "\n".join(parts).strip()
        logger.info(
            f"Extracted {len(extracted_text)} characters from DOCX file "
            f"(with table preservation)"
        )
        return extracted_text

    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise
