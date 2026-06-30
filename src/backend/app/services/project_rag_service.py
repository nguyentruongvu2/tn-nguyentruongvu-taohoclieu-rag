"""Project-based RAG APIs for incremental section authoring."""

from __future__ import annotations

import asyncio
import json
import importlib
import logging
import re
import unicodedata
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import quote
import os
import uuid

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from fastapi.responses import PlainTextResponse, Response
from pydantic import BaseModel, Field

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - handled at runtime by endpoint guard.
    DocxDocument = None

from ..prompts.project_rag_system_prompts import (
    build_project_rag_system_prompt,
    build_project_rag_batch_system_prompt,
    get_batch_group_type,
)
from ..prompts.project_rag_user_prompts import (
    build_outline_user_prompt,
    build_project_rag_combined_prompt,
    build_section_user_prompt,
    build_project_rag_batch_user_prompt,
    build_section_edit_user_prompt,
)
from ..prompts.project_rag_section_profiles import (
    RETRIEVAL_MAP,
    get_retrieval_profile,
    get_section_user_intent_hint,
    normalize_section_profile_key,
)
from ..rag_pipeline import rag_pipeline
from ..security import enforce_rate_limit, get_current_user
from ..db.projects import (
    create_editor_project,
    list_editor_projects,
    get_editor_project_detail_for_user,
    update_editor_project_for_user,
    list_editor_sections,
    delete_project_for_user,
    get_project_for_user,
    create_editor_section,
    delete_editor_section,
    replace_editor_sections,
    get_editor_section_for_user,
    update_editor_section,
    add_editor_section_history,
    list_editor_section_history,
    get_editor_section_history_entry,
)
from ..db.documents import (
    list_documents,
    create_project_document,
    get_document_with_sections,
    set_document_sections,
)

router = APIRouter(tags=["project-rag"])
logger = logging.getLogger(__name__)


# ── Background evaluation helper ─────────────────────────────────────────────

def _evaluate_and_save_background(
    section_id: str,
    user_id: int,
    role: str,
    section_name: str,
    context_text: str,
    generated_content: str,
) -> None:
    """Run quality evaluation and persist result — called as a BackgroundTask."""
    try:
        evaluation, _ = _evaluate_section_quality(
            section_name=section_name,
            context_text=context_text,
            generated_content=generated_content,
        )
        update_editor_section(
            section_id=section_id,
            user_id=user_id,
            role=role,
            evaluation=evaluation,
        )
        logger.debug("Background evaluation saved for section %s", section_id)
    except Exception:
        logger.exception("Background evaluation failed for section %s", section_id)

class ProjectCreateRequest(BaseModel):
    title: str = Field(min_length=2, max_length=255)
    description: str = ""
    knowledge_base_ids: list[str] = Field(default_factory=list)
    level: str = "basic"
    format: str = "markdown"
    teaching_tone: str = ""  # "" | "academic" | "inspiring" | "practical"
    syllabus_doc_id: str | None = None


class ProjectCreateResponse(BaseModel):
    id: str
    project_id: str
    title: str
    description: str
    knowledge_base_ids: list[str]
    level: str
    format: str
    teaching_tone: str = ""
    syllabus_doc_id: str | None = None
    created_at: str
    updated_at: str


class ProjectUpdateRequest(BaseModel):
    title: str | None = Field(default=None, min_length=2, max_length=255)
    description: str | None = None
    knowledge_base_ids: list[str] | None = None
    level: str | None = None
    format: str | None = None
    teaching_tone: str | None = None  # "" | "academic" | "inspiring" | "practical"
    syllabus_doc_id: str | None = None


class SectionPayload(BaseModel):
    section_id: str
    title: str
    content: str = ""
    status: str = "empty"


class DocumentCreateRequest(BaseModel):
    project_id: str
    title: str = Field(min_length=2, max_length=255)
    source_document_ids: list[str] = Field(default_factory=list)


class DocumentResponse(BaseModel):
    doc_id: str
    project_id: str
    title: str
    created_at: str
    updated_at: str
    sections: list[SectionPayload]
    source_document_ids: list[str]


class OutlineRequest(BaseModel):
    doc_id: str
    prompt: str = Field(min_length=3)
    selected_documents: list[str] = Field(default_factory=list)


class OutlineResponse(BaseModel):
    doc_id: str
    sections: list[SectionPayload]


class GenerateSectionRequest(BaseModel):
    project_id: str
    section_id: str
    prompt: str = ""


class BatchGenerateRequest(BaseModel):
    project_id: str
    section_ids: list[str]
    prompt: str = ""


class UpdateSectionRequest(BaseModel):
    title: str | None = None
    content: str | None = None
    prompt: str | None = None
    order: int | None = None


class CreateSectionRequest(BaseModel):
    project_id: str
    title: str = Field(min_length=1, max_length=255)
    prompt: str = ""
    order: int | None = None


class GenerateProjectOutlineRequest(BaseModel):
    prompt: str = Field(min_length=3, max_length=3000)


def _strip_accents(text: str) -> str:
    normalized = unicodedata.normalize("NFD", text or "")
    stripped = "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")
    return stripped.replace("đ", "d").replace("Đ", "D")


def _normalize_heading_label(text: str) -> str:
    clean = _strip_accents((text or "").lower())
    clean = re.sub(r"[^a-z0-9\s]", " ", clean)
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean


def _build_download_filename(title: str) -> tuple[str, str]:
    raw = (title or "").strip() or "teaching_project"
    with_ext = f"{raw}.md"

    # ASCII fallback for legacy header field value.
    ascii_fallback = re.sub(r"[^A-Za-z0-9._-]+", "_", with_ext).strip("._") or "teaching_project.md"
    # RFC 5987 filename* for UTF-8 names.
    utf8_encoded = quote(with_ext, safe="")
    return ascii_fallback, utf8_encoded


def _build_download_filename_with_ext(title: str, extension: str) -> tuple[str, str]:
    ext = re.sub(r"[^A-Za-z0-9]+", "", (extension or "").lower()) or "md"
    raw = (title or "").strip() or "teaching_project"
    with_ext = f"{raw}.{ext}"

    ascii_fallback = (
        re.sub(r"[^A-Za-z0-9._-]+", "_", with_ext).strip("._") or f"teaching_project.{ext}"
    )
    utf8_encoded = quote(with_ext, safe="")
    return ascii_fallback, utf8_encoded


def _project_sections_sorted(project: dict[str, Any]) -> list[dict[str, Any]]:
    sections = project.get("sections", [])
    if not isinstance(sections, list):
        return []

    def _safe_order_index(item: dict[str, Any]) -> int:
        try:
            return int(item.get("order_index", 0) or 0)
        except (TypeError, ValueError):
            return 0

    return sorted(sections, key=_safe_order_index)


def _strip_source_citations_for_export(markdown_text: str) -> str:
    text = str(markdown_text or "")
    if not text.strip():
        return ""

    cleaned = text

    # Remove inline/standalone citation lines.
    cleaned = re.sub(
        r"(?im)^\s*(?:[-*•]\s+)?📚\s*Nguồn\s*:\s*.*$\n?",
        "",
        cleaned,
    )

    # Remove grouped citation blocks (heading + bullets).
    cleaned = re.sub(
        r"(?is)\n{0,2}📚\s*Nguồn\s*:\s*\n(?:\s*[-*•]\s+.*(?:\n|$))+",
        "\n",
        cleaned,
    )

    # Remove source tracing footer blocks.
    cleaned = re.sub(
        r"(?im)^\s*---\s*\*?\s*(nguồn|nguon)\s*:[^\n]*\*?\s*$",
        "",
        cleaned,
    )

    # Remove internal source anchor links if they remain.
    cleaned = re.sub(
        r"\[([^\]]+)\]\(#source:[^)]+\)",
        r"\1",
        cleaned,
        flags=re.IGNORECASE,
    )

    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _normalize_kb_ids(raw_ids: Any) -> list[str]:
    normalized: list[str] = []
    for item in raw_ids or []:
        value = str(item or "").strip()
        if value and value not in normalized:
            normalized.append(value)
    return normalized


def infer_level_from_title(title: str) -> int:
    normalized = (title or "").strip()
    if normalized.lower().startswith("chương") or normalized.lower().startswith("chuong"):
        return 1
    matched = re.match(r"^(\d+(?:\.\d+)*)", normalized)
    if not matched:
        return 1
    return max(1, len(matched.group(1).split(".")))


def _fetch_image_bytes(url: str) -> bytes | None:
    if not url.startswith("http"):
        return None
        
    import hashlib
    import os
    import httpx
    
    upload_dir = os.path.join(os.path.dirname(__file__), '../../uploads')
    cache_dir = os.path.join(upload_dir, "image_cache")
    os.makedirs(cache_dir, exist_ok=True)
    
    url_hash = hashlib.md5(url.encode("utf-8")).hexdigest()
    cache_path = os.path.join(cache_dir, url_hash)
    
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "rb") as f:
                return f.read()
        except Exception as e:
            logging.getLogger(__name__).error(f"Failed to read cached image: {e}")
            
    try:
        response = httpx.get(url, timeout=15.0)
        if response.status_code == 200:
            try:
                with open(cache_path, "wb") as f:
                    f.write(response.content)
            except Exception as e:
                logging.getLogger(__name__).error(f"Failed to write image cache: {e}")
            return response.content
    except Exception as e:
        logging.getLogger(__name__).error(f"Failed to fetch image from URL {url}: {e}")
    return None


def _latex_to_png(latex: str) -> bytes | None:
    # Strip delimiters
    clean_latex = latex.strip("$").strip().strip("\\(").strip("\\)").strip("\\[").strip("\\]").strip()
    if not clean_latex:
        return None
        
    import hashlib
    import os
    import httpx
    
    upload_dir = os.path.join(os.path.dirname(__file__), '../../uploads')
    cache_dir = os.path.join(upload_dir, "latex_cache")
    os.makedirs(cache_dir, exist_ok=True)
    
    latex_hash = hashlib.md5(clean_latex.encode("utf-8")).hexdigest()
    cache_path = os.path.join(cache_dir, f"{latex_hash}.png")
    
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "rb") as f:
                return f.read()
        except Exception as e:
            logging.getLogger(__name__).error(f"Failed to read cached latex PNG: {e}")
            
    # URL encode
    from urllib.parse import quote
    encoded = quote(clean_latex)
    # Using CodeCogs API with DPI (300) and white background for extreme sharpness
    url = f"https://latex.codecogs.com/png.latex?%5Cbg_white%20%5Cdpi%7B300%7D%20{encoded}"
    try:
        response = httpx.get(url, timeout=10.0)
        if response.status_code == 200:
            try:
                with open(cache_path, "wb") as f:
                    f.write(response.content)
            except Exception as e:
                logging.getLogger(__name__).error(f"Failed to write latex PNG cache: {e}")
            return response.content
    except Exception as e:
        logging.getLogger(__name__).error(f"Failed to render LaTeX {clean_latex} to PNG: {e}")
    return None


def _prefetch_resources(markdown: str) -> None:
    import re
    import hashlib
    import os
    from concurrent.futures import ThreadPoolExecutor
    
    # 1. Parse all image URLs (excluding placeholder:)
    img_urls = re.findall(r'!\[.*?\]\(\s*<?(http[s]?://[^)>]+)>?\s*\)', markdown)
    
    # 2. Parse all math formulas: $...$, $$...$$, \(...\), \[...\]
    math_segments = re.findall(r'(\$\$[^\n$]+\$\$|\$[^\n$]+\$|\\\[.*?\\\]|\\\(.*?\\\))', markdown)
    
    # Clean unique items
    unique_urls = list(set(img_urls))
    unique_maths = list(set(math_segments))
    
    # Filter uncached images
    upload_dir = os.path.join(os.path.dirname(__file__), '../../uploads')
    img_cache_dir = os.path.join(upload_dir, "image_cache")
    os.makedirs(img_cache_dir, exist_ok=True)
    
    uncached_urls = []
    for url in unique_urls:
        url_hash = hashlib.md5(url.encode("utf-8")).hexdigest()
        cache_path = os.path.join(img_cache_dir, url_hash)
        if not os.path.exists(cache_path):
            uncached_urls.append(url)
            
    # Filter uncached math
    latex_cache_dir = os.path.join(upload_dir, "latex_cache")
    os.makedirs(latex_cache_dir, exist_ok=True)
    
    uncached_maths = []
    for math in unique_maths:
        clean_latex = math.strip("$").strip().strip("\\(").strip("\\)").strip("\\[").strip("\\]").strip()
        if clean_latex:
            latex_hash = hashlib.md5(clean_latex.encode("utf-8")).hexdigest()
            cache_path = os.path.join(latex_cache_dir, f"{latex_hash}.png")
            if not os.path.exists(cache_path):
                uncached_maths.append(math)
                
    if not uncached_urls and not uncached_maths:
        return
        
    # Parallel fetch
    with ThreadPoolExecutor(max_workers=12) as executor:
        for url in uncached_urls:
            executor.submit(_fetch_image_bytes, url)
        for math in uncached_maths:
            executor.submit(_latex_to_png, math)


def _render_project_markdown(project: dict[str, Any]) -> str:
    sections = _project_sections_sorted(project)
    lines = [f"# {project['title']}", ""]
    if str(project.get("description") or "").strip():
        lines.extend([_strip_source_citations_for_export(str(project.get("description") or "")), ""])

    for section in sections:
        title = section.get('title') or ""
        level = infer_level_from_title(title)
        hashes = "#" * (level + 1)
        lines.append(f"{hashes} {title}")
        lines.append("")
        lines.append(_strip_source_citations_for_export(str(section.get("content_markdown") or "")))
        lines.append("")

    return "\n".join(lines).strip() + "\n"


_PROJECT_EXPORT_PDF_FONT = ""
_PROJECT_EXPORT_PDF_MONO_FONT = ""


def _load_reportlab_runtime() -> tuple[Any, Any, Any, Any] | None:
    try:
        pagesizes_module = importlib.import_module("reportlab.lib.pagesizes")
        pdfmetrics_module = importlib.import_module("reportlab.pdfbase.pdfmetrics")
        ttfonts_module = importlib.import_module("reportlab.pdfbase.ttfonts")
        canvas_module = importlib.import_module("reportlab.pdfgen.canvas")
    except Exception:
        return None

    return pagesizes_module.A4, pdfmetrics_module, ttfonts_module.TTFont, canvas_module


def _register_font_family(pdfmetrics_module: Any, ttfont_cls: Any) -> str:
    global _PROJECT_EXPORT_PDF_FONT
    if _PROJECT_EXPORT_PDF_FONT:
        return _PROJECT_EXPORT_PDF_FONT

    fallback_font = "Helvetica"

    font_candidates = [
        # Name, Regular, Bold, Italic, BoldItalic
        ("RAGTimes", "C:/Windows/Fonts/times.ttf", "C:/Windows/Fonts/timesbd.ttf", "C:/Windows/Fonts/timesi.ttf", "C:/Windows/Fonts/timesbi.ttf"),
        ("RAGArial", "C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf", "C:/Windows/Fonts/ariali.ttf", "C:/Windows/Fonts/arialbi.ttf"),
        ("RAGDejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf"),
        ("RAGNoto", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf", "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf", "/usr/share/fonts/truetype/noto/NotoSans-Italic.ttf", "/usr/share/fonts/truetype/noto/NotoSans-BoldItalic.ttf"),
        ("RAGLiberation", "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf"),
    ]

    for name, reg, bold, ital, bold_ital in font_candidates:
        if Path(reg).exists():
            try:
                pdfmetrics_module.registerFont(ttfont_cls(name, reg))
                
                bold_name = f"{name}-Bold"
                if Path(bold).exists():
                    pdfmetrics_module.registerFont(ttfont_cls(bold_name, bold))
                else:
                    bold_name = name
                    
                ital_name = f"{name}-Italic"
                if Path(ital).exists():
                    pdfmetrics_module.registerFont(ttfont_cls(ital_name, ital))
                else:
                    ital_name = name
                    
                bi_name = f"{name}-BoldItalic"
                if Path(bold_ital).exists():
                    pdfmetrics_module.registerFont(ttfont_cls(bi_name, bold_ital))
                else:
                    bi_name = bold_name
                
                pdfmetrics_module.registerFontFamily(
                    name,
                    normal=name,
                    bold=bold_name,
                    italic=ital_name,
                boldItalic=bi_name
                )
                _PROJECT_EXPORT_PDF_FONT = name
                return _PROJECT_EXPORT_PDF_FONT
            except Exception:
                continue

    _PROJECT_EXPORT_PDF_FONT = fallback_font
    return _PROJECT_EXPORT_PDF_FONT


def _register_mono_font_family(pdfmetrics_module: Any, ttfont_cls: Any) -> str:
    global _PROJECT_EXPORT_PDF_MONO_FONT
    if _PROJECT_EXPORT_PDF_MONO_FONT:
        return _PROJECT_EXPORT_PDF_MONO_FONT

    fallback_font = "Courier"

    font_candidates = [
        # Name, Regular, Bold, Italic, BoldItalic
        ("RAGCourierNew", "C:/Windows/Fonts/cour.ttf", "C:/Windows/Fonts/courbd.ttf", "C:/Windows/Fonts/couri.ttf", "C:/Windows/Fonts/courbi.ttf"),
        ("RAGDejaVuMono", "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Oblique.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSansMono-BoldOblique.ttf"),
        ("RAGLiberationMono", "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationMono-Bold.ttf", "/usr/share/fonts/truetype/liberation/LiberationMono-Italic.ttf", "/usr/share/fonts/truetype/liberation/LiberationMono-BoldItalic.ttf"),
    ]

    for name, reg, bold, ital, bold_ital in font_candidates:
        if Path(reg).exists():
            try:
                pdfmetrics_module.registerFont(ttfont_cls(name, reg))
                
                bold_name = f"{name}-Bold"
                if Path(bold).exists():
                    pdfmetrics_module.registerFont(ttfont_cls(bold_name, bold))
                else:
                    bold_name = name
                    
                ital_name = f"{name}-Italic"
                if Path(ital).exists():
                    pdfmetrics_module.registerFont(ttfont_cls(ital_name, ital))
                else:
                    ital_name = name
                    
                bi_name = f"{name}-BoldItalic"
                if Path(bold_ital).exists():
                    pdfmetrics_module.registerFont(ttfont_cls(bi_name, bold_ital))
                else:
                    bi_name = bold_name
                
                pdfmetrics_module.registerFontFamily(
                    name,
                    normal=name,
                    bold=bold_name,
                    italic=ital_name,
                    boldItalic=bi_name
                )
                _PROJECT_EXPORT_PDF_MONO_FONT = name
                return _PROJECT_EXPORT_PDF_MONO_FONT
            except Exception:
                continue

    global _PROJECT_EXPORT_PDF_FONT
    if _PROJECT_EXPORT_PDF_FONT and _PROJECT_EXPORT_PDF_FONT != "Helvetica":
        _PROJECT_EXPORT_PDF_MONO_FONT = _PROJECT_EXPORT_PDF_FONT
    else:
        _PROJECT_EXPORT_PDF_MONO_FONT = fallback_font

    return _PROJECT_EXPORT_PDF_MONO_FONT


def _get_mermaid_diagram_image(code: str) -> bytes | None:
    import urllib.request
    import urllib.error
    import base64
    import json
    
    try:
        data = {
            "code": code.strip(),
            "mermaid": {"theme": "default"}
        }
        json_str = json.dumps(data)
        b64 = base64.urlsafe_b64encode(json_str.encode('utf-8')).decode('utf-8').rstrip('=')
        url = f"https://mermaid.ink/img/{b64}"
        
        req = urllib.request.Request(
            url,
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        )
        with urllib.request.urlopen(req, timeout=5.0) as response:
            if response.status == 200:
                return response.read()
    except Exception:
        pass
    return None


def _wrap_pdf_text_line(
    text: str,
    font_name: str,
    font_size: int,
    max_width: float,
    pdfmetrics_module: Any,
) -> list[str]:
    normalized = (text or "").replace("\t", "    ").strip()
    if not normalized:
        return [""]

    words = normalized.split()
    if not words:
        return [""]

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if pdfmetrics_module.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word

    lines.append(current)
    return lines


def _render_project_pdf_bytes(project: dict[str, Any]) -> bytes:
    import html
    import re
    import importlib
    
    reportlab_runtime = _load_reportlab_runtime()
    if reportlab_runtime is None:
        raise RuntimeError("PDF export dependency is not available")

    a4_pagesize, pdfmetrics_module, ttfont_cls, canvas_module = reportlab_runtime
    font_name = _register_font_family(pdfmetrics_module, ttfont_cls)
    mono_font_name = _register_mono_font_family(pdfmetrics_module, ttfont_cls)

    try:
        platypus_module = importlib.import_module("reportlab.platypus")
        SimpleDocTemplate = platypus_module.SimpleDocTemplate
        Paragraph = platypus_module.Paragraph
        Spacer = platypus_module.Spacer
        Table = platypus_module.Table
        TableStyle = platypus_module.TableStyle
        PageBreak = platypus_module.PageBreak
        KeepTogether = platypus_module.KeepTogether
        RLImage = platypus_module.Image
        
        styles_module = importlib.import_module("reportlab.lib.styles")
        getSampleStyleSheet = styles_module.getSampleStyleSheet
        ParagraphStyle = styles_module.ParagraphStyle
        
        colors_module = importlib.import_module("reportlab.lib.colors")
        HexColor = colors_module.HexColor
    except Exception as exc:
        raise RuntimeError(f"Failed to load ReportLab Platypus components: {exc}")

    markdown = _render_project_markdown(project)
    _prefetch_resources(markdown)
    stream = BytesIO()
    
    page_width, page_height = a4_pagesize
    margin_left = 44
    margin_right = 44
    margin_top = 48
    margin_bottom = 44
    max_width = page_width - margin_left - margin_right
    
    styles = getSampleStyleSheet()
    
    body_style = ParagraphStyle(
        'PDFBodyText',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=14,
        textColor=HexColor("#334155"),
        spaceAfter=6
    )
    
    h1_style = ParagraphStyle(
        'PDFH1',
        parent=styles['Heading1'],
        fontName=f"{font_name}-Bold" if font_name != "Helvetica" else "Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=HexColor("#1e3a8a"),
        spaceBefore=14,
        spaceAfter=10,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'PDFH2',
        parent=styles['Heading2'],
        fontName=f"{font_name}-Bold" if font_name != "Helvetica" else "Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=HexColor("#0f766e"),
        spaceBefore=12,
        spaceAfter=8,
        keepWithNext=True
    )
    
    h3_style = ParagraphStyle(
        'PDFH3',
        parent=styles['Heading3'],
        fontName=f"{font_name}-Bold" if font_name != "Helvetica" else "Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=HexColor("#1e293b"),
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )
    
    h4_style = ParagraphStyle(
        'PDFH4',
        parent=styles['Heading3'],
        fontName=f"{font_name}-Bold" if font_name != "Helvetica" else "Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=HexColor("#334155"),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True
    )
    
    h5_style = ParagraphStyle(
        'PDFH5',
        parent=styles['Heading3'],
        fontName=f"{font_name}-Bold" if font_name != "Helvetica" else "Helvetica-Bold",
        fontSize=10,
        leading=13,
        textColor=HexColor("#475569"),
        spaceBefore=6,
        spaceAfter=3,
        keepWithNext=True
    )
    
    list_style = ParagraphStyle(
        'PDFListStyle',
        parent=body_style,
        leftIndent=18,
        firstLineIndent=-10,
        spaceAfter=4
    )
    
    def _inline_markdown_to_html(text: str) -> str:
        escaped = html.escape(text)
        escaped = re.sub(r'`(.*?)`', f'<font face="{mono_font_name}" color="#c7254e"><b>\\1</b></font>', escaped)
        escaped = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', escaped)
        escaped = re.sub(r'\*(.*?)\*', r'<i>\1</i>', escaped)
        return escaped

    temp_files = []

    def _process_math_in_text(html_text: str, text_style: ParagraphStyle) -> str:
        math_segments = re.findall(r'(\$\$[^\n$]+\$\$|\$[^\n$]+\$|\\\(.*?\\\)|\\\[.*?\\\])', html_text)
        for math in math_segments:
            png_bytes = _latex_to_png(math)
            if png_bytes:
                try:
                    from PIL import Image as PILImage
                    pil_img = PILImage.open(BytesIO(png_bytes))
                    w, h = pil_img.size
                    
                    target_h = max(10, text_style.fontSize * 1.1)
                    target_w = int(w * (target_h / h))
                    if target_w > max_width:
                        target_w = int(max_width)
                        target_h = int(h * (max_width / w))
                    
                    temp_dir = os.path.join(os.path.dirname(__file__), '../../uploads/temp_export')
                    os.makedirs(temp_dir, exist_ok=True)
                    temp_file = os.path.join(temp_dir, f"math_h_{uuid.uuid4().hex}.png")
                    with open(temp_file, "wb") as f_out:
                        f_out.write(png_bytes)
                    temp_files.append(temp_file)
                    
                    img_tag = f'<img src="{temp_file}" width="{target_w}" height="{target_h}" valign="middle"/>'
                    html_text = html_text.replace(math, img_tag)
                except Exception as e:
                    logging.getLogger(__name__).error(f"Failed to render math in text: {e}")
        return html_text

    def _process_pdf_paragraph(text: str, body_style: ParagraphStyle) -> list[Any]:
        stripped = text.strip()
        img_match = re.match(r"^!\[(.*?)\]\(\s*<?(placeholder:[^)>]*|http[s]?://[^)>]+)>?\s*\)$", stripped)
        if img_match:
            alt = img_match.group(1)
            path = img_match.group(2).strip()
            
            if path.startswith("placeholder:"):
                desc = path[len("placeholder:"):].strip()
                parts = desc.split("|")
                raw_vi = parts[0] or ""
                raw_en = parts[1] if len(parts) > 1 else raw_vi
                
                clean_vi = raw_vi.strip().replace("_", " ")
                clean_en = raw_en.strip().replace("_", " ")
                
                title_style = ParagraphStyle(
                    'PDFPlacholderTitle',
                    parent=body_style,
                    fontName=body_style.fontName + "-Bold" if "Bold" not in body_style.fontName else body_style.fontName,
                    textColor=HexColor("#164e63")
                )
                desc_style = ParagraphStyle(
                    'PDFPlaceholderDesc',
                    parent=body_style,
                    fontStyle="italic",
                    textColor=HexColor("#0891b2")
                )
                
                p_title = Paragraph(f"<b>📊 Khung hình minh họa gợi ý: {html.escape(alt)}</b>", title_style)
                p_vi = Paragraph(f"Mô tả gợi ý: {html.escape(clean_vi)}", desc_style)
                p_en = Paragraph(f"Prompt AI: {html.escape(clean_en)}", body_style)
                
                t_data = [[p_title], [p_vi], [p_en]]
                t = Table(t_data, colWidths=[max_width])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), HexColor("#ecfeff")),
                    ('PADDING', (0,0), (-1,-1), 12),
                    ('BOX', (0,0), (-1,-1), 1.5, HexColor("#0891b2")),
                    ('TOPPADDING', (0,0), (-1,-1), 4),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                ]))
                return [t, Spacer(1, 8)]
            else:
                img_bytes = _fetch_image_bytes(path)
                if img_bytes:
                    try:
                        from PIL import Image as PILImage
                        pil_img = PILImage.open(BytesIO(img_bytes))
                        w, h = pil_img.size
                        if w > max_width:
                            h = int(h * (max_width / w))
                            w = int(max_width)
                        
                        temp_dir = os.path.join(os.path.dirname(__file__), '../../uploads/temp_export')
                        os.makedirs(temp_dir, exist_ok=True)
                        temp_file = os.path.join(temp_dir, f"img_{uuid.uuid4().hex}.png")
                        with open(temp_file, "wb") as f_out:
                            f_out.write(img_bytes)
                        temp_files.append(temp_file)
                        
                        return [RLImage(temp_file, width=w, height=h), Spacer(1, 8)]
                    except Exception as e:
                        logging.getLogger(__name__).error(f"Failed to render block image: {e}")
                
                fallback_style = ParagraphStyle('PDFFallbackImage', parent=body_style, textColor=HexColor("#ef4444"))
                return [Paragraph(f"<b>[Lỗi tải hình ảnh: {html.escape(alt)} ({path})]</b>", fallback_style), Spacer(1, 4)]

        html_text = _inline_markdown_to_html(text)
        html_text = _process_math_in_text(html_text, body_style)
                    
        inline_images = re.findall(r'(!\[.*?\]\(\s*(?:<.*?>|.*?)\s*\))', html_text)
        for img_syntax in inline_images:
            match = re.match(r'!\[(.*?)\]\(\s*<?(placeholder:[^)>]*|http[s]?://[^)>]+)>?\s*\)', img_syntax)
            if match:
                alt = match.group(1)
                path = match.group(2).strip()
                
                if path.startswith("placeholder:"):
                    html_text = html_text.replace(img_syntax, f"<b>📊 [Khung ảnh gợi ý: {html.escape(alt)}]</b>")
                else:
                    img_bytes = _fetch_image_bytes(path)
                    if img_bytes:
                        try:
                            from PIL import Image as PILImage
                            pil_img = PILImage.open(BytesIO(img_bytes))
                            w, h = pil_img.size
                            
                            target_w = min(120, w)
                            target_h = int(h * (target_w / w))
                            
                            temp_dir = os.path.join(os.path.dirname(__file__), '../../uploads/temp_export')
                            os.makedirs(temp_dir, exist_ok=True)
                            temp_file = os.path.join(temp_dir, f"inline_img_{uuid.uuid4().hex}.png")
                            with open(temp_file, "wb") as f_out:
                                f_out.write(img_bytes)
                            temp_files.append(temp_file)
                            
                            img_tag = f'<img src="{temp_file}" width="{target_w}" height="{target_h}" valign="middle"/>'
                            html_text = html_text.replace(img_syntax, img_tag)
                        except Exception as e:
                            logging.getLogger(__name__).error(f"Failed to inline render image: {e}")
                            html_text = html_text.replace(img_syntax, f"<b>[Lỗi tải hình: {html.escape(alt)}]</b>")
                    else:
                        html_text = html_text.replace(img_syntax, f"<b>[Lỗi tải hình: {html.escape(alt)}]</b>")
                        
        return [Paragraph(html_text, body_style)]

    blocks = []
    lines = markdown.splitlines()
    in_code_block = False
    code_content = []
    current_block_type = None
    current_block_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.strip().startswith("```"):
            if in_code_block:
                if code_lang == "mermaid":
                    img_bytes = _get_mermaid_diagram_image("\n".join(code_content))
                    if img_bytes:
                        blocks.append({
                            "type": "mermaid_image",
                            "content": img_bytes
                        })
                    else:
                        blocks.append({
                            "type": "code",
                            "lang": "mermaid",
                            "content": "\n".join(code_content)
                        })
                else:
                    blocks.append({
                        "type": "code",
                        "lang": code_lang,
                        "content": "\n".join(code_content)
                    })
                code_content = []
                in_code_block = False
            else:
                if current_block_lines:
                    blocks.append({
                        "type": current_block_type or "paragraph",
                        "lines": current_block_lines
                    })
                    current_block_lines = []
                    current_block_type = None
                in_code_block = True
                code_lang = line.strip()[3:].strip().lower()
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            if current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
                current_block_type = None
            i += 1
            continue

        if stripped.startswith("#"):
            if current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            
            level = 0
            while level < len(stripped) and stripped[level] == '#':
                level += 1
            content = stripped[level:].strip()
            blocks.append({
                "type": f"h{level}",
                "content": content
            })
            current_block_type = None
            i += 1
            continue

        if stripped.startswith(">") or stripped.startswith("[!NOTE]"):
            if current_block_type not in {"blockquote", "note"} and current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            
            if stripped.startswith("[!NOTE]"):
                current_block_type = "note"
                content = stripped[7:].strip()
            elif stripped.startswith("> [!NOTE]"):
                current_block_type = "note"
                content = stripped[9:].strip()
            else:
                current_block_type = "blockquote"
                content = stripped[1:].strip()
            
            current_block_lines.append(content)
            i += 1
            continue

        if stripped.startswith("|"):
            if current_block_type != "table" and current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            current_block_type = "table"
            current_block_lines.append(stripped)
            i += 1
            continue

        is_bullet = stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ ")
        is_numbered = re.match(r"^\d+\.\s+", stripped) is not None

        if is_bullet or is_numbered:
            if current_block_type != "list" and current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            current_block_type = "list"
            current_block_lines.append(stripped)
            i += 1
            continue

        if current_block_type in {"h1", "h2", "h3", "table", "list", "blockquote", "note"} and current_block_lines:
            blocks.append({
                "type": current_block_type,
                "lines": current_block_lines
            })
            current_block_lines = []
            current_block_type = "paragraph"

        current_block_type = "paragraph"
        current_block_lines.append(line)
        i += 1

    if current_block_lines:
        blocks.append({
            "type": current_block_type or "paragraph",
            "lines": current_block_lines
        })
    elif in_code_block and code_content:
        blocks.append({
            "type": "code",
            "content": "\n".join(code_content)
        })

    flowables = []
    doc = SimpleDocTemplate(
        stream,
        pagesize=a4_pagesize,
        leftMargin=margin_left,
        rightMargin=margin_right,
        topMargin=margin_top,
        bottomMargin=margin_bottom
    )

    for block in blocks:
        b_type = block["type"]
        
        if b_type == "h1":
            processed_html = _process_math_in_text(_inline_markdown_to_html(block["content"]), h1_style)
            flowables.append(Paragraph(processed_html, h1_style))
        elif b_type == "h2":
            processed_html = _process_math_in_text(_inline_markdown_to_html(block["content"]), h2_style)
            flowables.append(Paragraph(processed_html, h2_style))
        elif b_type == "h3":
            processed_html = _process_math_in_text(_inline_markdown_to_html(block["content"]), h3_style)
            flowables.append(Paragraph(processed_html, h3_style))
        elif b_type == "h4":
            processed_html = _process_math_in_text(_inline_markdown_to_html(block["content"]), h4_style)
            flowables.append(Paragraph(processed_html, h4_style))
        elif b_type == "h5" or (b_type.startswith("h") and b_type[1:].isdigit()):
            processed_html = _process_math_in_text(_inline_markdown_to_html(block["content"]), h5_style)
            flowables.append(Paragraph(processed_html, h5_style))
            
        elif b_type == "paragraph":
            text = " ".join([l.strip() for l in block["lines"]])
            flowables.extend(_process_pdf_paragraph(text, body_style))
            
        elif b_type == "code":
            code_style = ParagraphStyle(
                'PDFCodeStyle',
                parent=body_style,
                fontName=mono_font_name,
                fontSize=8.5,
                leading=11,
                textColor=HexColor("#0f172a")
            )
            code_escaped = html.escape(block["content"]).replace("\n", "<br/>").replace(" ", "&nbsp;")
            p = Paragraph(code_escaped, code_style)
            tbl_style = TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), HexColor("#f1f5f9")),
                ('PADDING', (0,0), (-1,-1), 8),
                ('BOX', (0,0), (-1,-1), 0.5, HexColor("#cbd5e1")),
            ])
            t = Table([[p]], colWidths=[max_width])
            t.setStyle(tbl_style)
            flowables.append(t)
            flowables.append(Spacer(1, 8))
            
        elif b_type == "mermaid_image":
            try:
                from PIL import Image as PILImage
                from reportlab.platypus import Image as RLImage
                
                img_data = BytesIO(block["content"])
                pil_img = PILImage.open(img_data)
                w, h = pil_img.size
                
                if w > max_width:
                    h = int(h * (max_width / w))
                    w = int(max_width)
                
                img_data.seek(0)
                flowables.append(RLImage(img_data, width=w, height=h))
                flowables.append(Spacer(1, 8))
            except Exception:
                pass
            
        elif b_type in {"blockquote", "note"}:
            is_note = b_type == "note"
            note_style = ParagraphStyle(
                'PDFNoteStyle',
                parent=body_style,
                fontSize=9.5,
                leading=13,
                textColor=HexColor("#1e3a8a") if is_note else HexColor("#334155")
            )
            note_text = "<br/>".join([_process_math_in_text(_inline_markdown_to_html(line), note_style) for line in block["lines"]])
            p = Paragraph(note_text, note_style)
            tbl_style = TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), HexColor("#eff6ff") if is_note else HexColor("#f8fafc")),
                ('LEFTPADDING', (0,0), (-1,-1), 12),
                ('RIGHTPADDING', (0,0), (-1,-1), 12),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LINELEFT', (0,0), (0,-1), 3, HexColor("#2563eb") if is_note else HexColor("#94a3b8")),
                ('LINEBELOW', (0,0), (-1,-1), 0, colors_module.transparent),
                ('LINEABOVE', (0,0), (-1,-1), 0, colors_module.transparent),
                ('LINERIGHT', (0,0), (-1,-1), 0, colors_module.transparent),
            ])
            t = Table([[p]], colWidths=[max_width])
            t.setStyle(tbl_style)
            flowables.append(t)
            flowables.append(Spacer(1, 8))
            
        elif b_type == "list":
            for line in block["lines"]:
                stripped = line.strip()
                is_bullet = stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ ")
                
                if is_bullet:
                    prefix = "&bull; "
                    content = stripped[2:].strip()
                else:
                    match = re.match(r"^(\d+\.\s+)(.*)$", stripped)
                    if match:
                        prefix = match.group(1)
                        content = match.group(2).strip()
                    else:
                        prefix = "&bull; "
                        content = stripped
                        
                processed_content = _process_math_in_text(_inline_markdown_to_html(content), list_style)
                flowables.append(Paragraph(f"{prefix}{processed_content}", list_style))
            flowables.append(Spacer(1, 6))
            
        elif b_type == "table":
            table_data = []
            for line in block["lines"]:
                if re.match(r"^\|?\s*:?-+:?\s*(\|?\s*:?-+:?\s*)*\|?$", line):
                    continue
                parts = [p.strip() for p in line.split("|")]
                if parts and not parts[0]:
                    parts.pop(0)
                if parts and not parts[-1]:
                    parts.pop()
                if parts:
                    table_data.append(parts)
            
            if table_data:
                cell_style = ParagraphStyle(
                    'PDFTableCellStyle',
                    parent=body_style,
                    fontSize=9,
                    leading=12
                )
                header_style = ParagraphStyle(
                    'PDFTableHeaderCellStyle',
                    parent=body_style,
                    fontSize=9,
                    leading=12,
                    fontName=f"{font_name}-Bold" if font_name != "Helvetica" else "Helvetica-Bold",
                    textColor=HexColor("#1e293b")
                )

                formatted_data = []
                for row_idx, row in enumerate(table_data):
                    formatted_row = []
                    for cell in row:
                        cell_html = _inline_markdown_to_html(cell)
                        if row_idx == 0:
                            cell_html = _process_math_in_text(cell_html, header_style)
                            formatted_row.append(Paragraph(cell_html, header_style))
                        else:
                            cell_html = _process_math_in_text(cell_html, cell_style)
                            formatted_row.append(Paragraph(cell_html, cell_style))
                    formatted_data.append(formatted_row)
                
                num_cols = len(table_data[0])
                col_width = max_width / max(1, num_cols)
                col_widths = [col_width] * num_cols
                
                t_style = TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), HexColor("#f8fafc")),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 0.5, HexColor("#e2e8f0")),
                    ('TOPPADDING', (0,0), (-1,-1), 5),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                    ('LEFTPADDING', (0,0), (-1,-1), 6),
                    ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ])
                
                t = Table(formatted_data, colWidths=col_widths)
                t.setStyle(t_style)
                
                flowables.append(KeepTogether([t]))
                flowables.append(Spacer(1, 8))

    try:
        doc.build(flowables)
        stream.seek(0)
        return stream.getvalue()
    finally:
        for fpath in temp_files:
            try:
                if os.path.exists(fpath):
                    os.remove(fpath)
            except Exception as e:
                logging.getLogger(__name__).error(f"Failed to remove temp file {fpath}: {e}")


def _render_project_docx_bytes(project: dict[str, Any]) -> bytes:
    if DocxDocument is None:
        raise RuntimeError("DOCX export dependency is not available")
    
    import re
    from io import BytesIO
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    markdown = _render_project_markdown(project)
    _prefetch_resources(markdown)

    def add_formatted_text(p, text, force_bold=False, force_italic=False, force_color=None):
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                if force_italic:
                    run.italic = True
                if force_color:
                    run.font.color.rgb = force_color
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
                if force_bold:
                    run.bold = True
                if force_color:
                    run.font.color.rgb = force_color
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                if force_bold:
                    run.bold = True
                if force_italic:
                    run.italic = True
                if force_color:
                    run.font.color.rgb = force_color
            else:
                run = p.add_run(part)
                if force_bold:
                    run.bold = True
                if force_italic:
                    run.italic = True
                if force_color:
                    run.font.color.rgb = force_color

    def process_text_runs_in_paragraph(p, text, force_bold=False, force_italic=False, force_color=None):
        parts = re.split(r'(!\[.*?\]\(\s*(?:<.*?>|.*?)\s*\)|\$\$[^\n$]+\$\$|\$[^\n$]+\$|\\\[.*?\\\]|\\\(.*?\\\))', text)
        for part in parts:
            if not part:
                continue
            img_match = re.match(r'^!\[(.*?)\]\(\s*<?(placeholder:[^)>]*|http[s]?://[^)>]+)>?\s*\)$', part.strip())
            if img_match:
                alt = img_match.group(1)
                path = img_match.group(2).strip()
                if path.startswith("placeholder:"):
                    run = p.add_run(f"📊 [Khung ảnh gợi ý: {alt}]")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x08, 0x91, 0xb2)
                    run.font.name = 'Times New Roman'
                else:
                    img_bytes = _fetch_image_bytes(path)
                    if img_bytes:
                        try:
                            from PIL import Image as PILImage
                            pil_img = PILImage.open(BytesIO(img_bytes))
                            w, h = pil_img.size
                            target_h = Pt(24)
                            target_w = Pt(24 * (w / h))
                            run = p.add_run()
                            run.add_picture(BytesIO(img_bytes), width=target_w, height=target_h)
                        except Exception:
                            run = p.add_run(f"[Lỗi ảnh: {alt}]")
                            run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                            run.font.name = 'Times New Roman'
                    else:
                        run = p.add_run(f"[Lỗi ảnh: {alt}]")
                        run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                        run.font.name = 'Times New Roman'
            elif part.startswith('$') or part.startswith('\\(') or part.startswith('\\['):
                math_bytes = _latex_to_png(part)
                if math_bytes:
                    try:
                        from PIL import Image as PILImage
                        pil_img = PILImage.open(BytesIO(math_bytes))
                        w, h = pil_img.size
                        is_block = part.startswith('$$') or part.startswith('\\[')
                        target_h = Pt(20) if is_block else Pt(11.5)
                        target_w = Pt(target_h.pt * (w / h))
                        run = p.add_run()
                        run.add_picture(BytesIO(math_bytes), width=target_w, height=target_h)
                    except Exception:
                        run = p.add_run(part)
                        run.font.name = 'Courier New'
                else:
                    run = p.add_run(part)
                    run.font.name = 'Courier New'
            else:
                add_formatted_text(p, part, force_bold=force_bold, force_italic=force_italic, force_color=force_color)

    blocks = []
    lines = markdown.splitlines()
    in_code_block = False
    code_content = []
    current_block_type = None
    current_block_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code_block:
                if code_lang == "mermaid":
                    img_bytes = _get_mermaid_diagram_image("\n".join(code_content))
                    if img_bytes:
                        blocks.append({
                            "type": "mermaid_image",
                            "content": img_bytes
                        })
                    else:
                        blocks.append({
                            "type": "code",
                            "lang": "mermaid",
                            "content": "\n".join(code_content)
                        })
                else:
                    blocks.append({
                        "type": "code",
                        "lang": code_lang,
                        "content": "\n".join(code_content)
                    })
                code_content = []
                in_code_block = False
            else:
                if current_block_lines:
                    blocks.append({
                        "type": current_block_type or "paragraph",
                        "lines": current_block_lines
                    })
                    current_block_lines = []
                    current_block_type = None
                in_code_block = True
                code_lang = line.strip()[3:].strip().lower()
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            if current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
                current_block_type = None
            i += 1
            continue

        if stripped.startswith("#"):
            if current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            
            level = 0
            while level < len(stripped) and stripped[level] == '#':
                level += 1
            content = stripped[level:].strip()
            blocks.append({
                "type": f"h{level}",
                "content": content
            })
            current_block_type = None
            i += 1
            continue

        if stripped.startswith(">") or stripped.startswith("[!NOTE]"):
            if current_block_type not in {"blockquote", "note"} and current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            
            if stripped.startswith("[!NOTE]"):
                current_block_type = "note"
                content = stripped[7:].strip()
            elif stripped.startswith("> [!NOTE]"):
                current_block_type = "note"
                content = stripped[9:].strip()
            else:
                current_block_type = "blockquote"
                content = stripped[1:].strip()
            
            current_block_lines.append(content)
            i += 1
            continue

        if stripped.startswith("|"):
            if current_block_type != "table" and current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            current_block_type = "table"
            current_block_lines.append(stripped)
            i += 1
            continue

        is_bullet = stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ ")
        is_numbered = re.match(r"^\d+\.\s+", stripped) is not None

        if is_bullet or is_numbered:
            if current_block_type != "list" and current_block_lines:
                blocks.append({
                    "type": current_block_type or "paragraph",
                    "lines": current_block_lines
                })
                current_block_lines = []
            current_block_type = "list"
            current_block_lines.append(stripped)
            i += 1
            continue

        if current_block_type in {"h1", "h2", "h3", "table", "list", "blockquote", "note"} and current_block_lines:
            blocks.append({
                "type": current_block_type,
                "lines": current_block_lines
            })
            current_block_lines = []
            current_block_type = "paragraph"

        current_block_type = "paragraph"
        current_block_lines.append(line)
        i += 1

    if current_block_lines:
        blocks.append({
            "type": current_block_type or "paragraph",
            "lines": current_block_lines
        })
    elif in_code_block and code_content:
        blocks.append({
            "type": "code",
            "content": "\n".join(code_content)
        })

    document = DocxDocument()
    
    # Configure document margins
    for section in document.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure default text styles
    style_normal = document.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(0x00, 0x00, 0x00)
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.line_spacing = 1.25

    # Configure heading styles
    for h_name, size, color in [
        ('Heading 1', Pt(20), RGBColor(0x1e, 0x3a, 0x8a)),
        ('Heading 2', Pt(16), RGBColor(0x0f, 0x76, 0x6e)),
        ('Heading 3', Pt(13), RGBColor(0x1e, 0x29, 0x3b)),
    ]:
        if h_name in document.styles:
            h_style = document.styles[h_name]
            h_style.font.name = 'Times New Roman'
            h_style.font.size = size
            h_style.font.color.rgb = color
            h_style.font.bold = True
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
            h_style.paragraph_format.keep_with_next = True

    # Configure list styles
    for l_style_name in ['List Bullet', 'List Number']:
        if l_style_name in document.styles:
            l_style = document.styles[l_style_name]
            l_style.font.name = 'Times New Roman'
            l_style.font.size = Pt(12)
            l_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            l_style.paragraph_format.space_after = Pt(3)

    if 'Title' in document.styles:
        title_style = document.styles['Title']
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(26)
        title_style.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
        title_style.font.bold = True
        title_style.paragraph_format.space_after = Pt(12)
    
    project_title = str(project.get("title") or "Bài giảng").strip() or "Bài giảng"
    document.add_heading(project_title, level=0)
    
    project_description = str(project.get("description") or "").strip()
    if project_description:
        document.add_paragraph(project_description)
    
    for block in blocks:
        b_type = block["type"]
        if b_type.startswith("h") and b_type[1:].isdigit():
            level = int(b_type[1:])
            level = min(level, 3)
            p = document.add_heading(level=level)
            color = None
            if level == 1:
                color = RGBColor(0x1e, 0x3a, 0x8a)
            elif level == 2:
                color = RGBColor(0x0f, 0x76, 0x6e)
            elif level == 3:
                color = RGBColor(0x1e, 0x29, 0x3b)
            process_text_runs_in_paragraph(p, block["content"], force_bold=True, force_color=color)
        
        elif b_type == "paragraph":
            text = " ".join([l.strip() for l in block["lines"]])
            stripped = text.strip()
            
            # Check block image / placeholder
            img_match = re.match(r"^!\[(.*?)\]\(\s*<?(placeholder:[^)>]*|http[s]?://[^)>]+)>?\s*\)$", stripped)
            if img_match:
                alt = img_match.group(1)
                path = img_match.group(2).strip()
                if path.startswith("placeholder:"):
                    desc = path[len("placeholder:"):].strip()
                    from urllib.parse import unquote
                    try:
                        desc = unquote(desc)
                    except Exception:
                        pass
                    parts = desc.split("|")
                    raw_vi = parts[0] or ""
                    raw_en = parts[1] if len(parts) > 1 else raw_vi
                    clean_vi = raw_vi.strip().replace("_", " ")
                    clean_en = raw_en.strip().replace("_", " ")
                    
                    table = document.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    cell = table.cell(0, 0)
                    
                    # Style shading: #ECFEFF (light cyan)
                    shading_xml = f'<w:shd {nsdecls("w")} w:fill="ECFEFF"/>'
                    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
                    
                    # Style borders: left 3pt, others 1pt in #0891B2
                    borders_xml = f'''
                    <w:tcBorders {nsdecls("w")}>
                        <w:top w:val="single" w:sz="8" w:space="0" w:color="0891B2"/>
                        <w:left w:val="single" w:sz="24" w:space="0" w:color="0891B2"/>
                        <w:bottom w:val="single" w:sz="8" w:space="0" w:color="0891B2"/>
                        <w:right w:val="single" w:sz="8" w:space="0" w:color="0891B2"/>
                    </w:tcBorders>
                    '''
                    cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))
                    
                    # Style padding
                    margins_xml = f'''
                    <w:tcMar {nsdecls("w")}>
                        <w:top w:w="160" w:type="dxa"/>
                        <w:bottom w:w="160" w:type="dxa"/>
                        <w:left w:w="240" w:type="dxa"/>
                        <w:right w:w="240" w:type="dxa"/>
                    </w:tcMar>
                    '''
                    cell._tc.get_or_add_tcPr().append(parse_xml(margins_xml))
                    
                    p0 = cell.paragraphs[0]
                    p0.paragraph_format.space_before = Pt(0)
                    p0.paragraph_format.space_after = Pt(2)
                    run_title = p0.add_run(f"📊 Khung hình minh họa gợi ý: {alt}")
                    run_title.bold = True
                    run_title.font.color.rgb = RGBColor(0x16, 0x4e, 0x63)
                    run_title.font.name = 'Times New Roman'
                    run_title.font.size = Pt(11)
                    
                    p1 = cell.add_paragraph()
                    p1.paragraph_format.space_before = Pt(2)
                    p1.paragraph_format.space_after = Pt(2)
                    run_vi = p1.add_run(f"Mô tả gợi ý: {clean_vi}")
                    run_vi.italic = True
                    run_vi.font.color.rgb = RGBColor(0x08, 0x91, 0xb2)
                    run_vi.font.name = 'Times New Roman'
                    run_vi.font.size = Pt(10)
                    
                    p2 = cell.add_paragraph()
                    p2.paragraph_format.space_before = Pt(2)
                    p2.paragraph_format.space_after = Pt(0)
                    run_en = p2.add_run(f"Prompt AI: {clean_en}")
                    run_en.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                    run_en.font.name = 'Times New Roman'
                    run_en.font.size = Pt(10)
                    
                    document.add_paragraph()
                else:
                    img_bytes = _fetch_image_bytes(path)
                    if img_bytes:
                        try:
                            from PIL import Image as PILImage
                            pil_img = PILImage.open(BytesIO(img_bytes))
                            w, h = pil_img.size
                            max_w_inches = 5.5
                            if w / 96.0 < max_w_inches:
                                fit_w = Inches(w / 96.0)
                            else:
                                fit_w = Inches(max_w_inches)
                            
                            fit_h = Inches(fit_w.inches * (h / w))
                            document.add_picture(BytesIO(img_bytes), width=fit_w, height=fit_h)
                            p_img = document.paragraphs[-1]
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            if alt:
                                p_cap = document.add_paragraph()
                                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_cap = p_cap.add_run(f"Hình: {alt}")
                                run_cap.italic = True
                                run_cap.font.size = Pt(10)
                                run_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                        except Exception:
                            p_err = document.add_paragraph()
                            run_err = p_err.add_run(f"[Lỗi tải hình ảnh: {alt} ({path})]")
                            run_err.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                            run_err.bold = True
                    else:
                        p_err = document.add_paragraph()
                        run_err = p_err.add_run(f"[Lỗi tải hình ảnh: {alt} ({path})]")
                        run_err.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                        run_err.bold = True
            else:
                p = document.add_paragraph()
                process_text_runs_in_paragraph(p, text)
            
        elif b_type == "list":
            for line in block["lines"]:
                line_stripped = line.strip()
                if line_stripped.startswith("- ") or line_stripped.startswith("* ") or line_stripped.startswith("+ "):
                    text = line_stripped[2:].strip()
                    p = document.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(3)
                    process_text_runs_in_paragraph(p, text)
                else:
                    match = re.match(r"^\d+\.\s+", line_stripped)
                    if match:
                        text = line_stripped[len(match.group()):].strip()
                        p = document.add_paragraph(style='List Number')
                        p.paragraph_format.space_after = Pt(3)
                        process_text_runs_in_paragraph(p, text)
                    else:
                        p = document.add_paragraph()
                        p.paragraph_format.space_after = Pt(3)
                        process_text_runs_in_paragraph(p, line_stripped)
                        
        elif b_type in {"blockquote", "note"}:
            is_note = b_type == "note"
            bg_color = "EFF6FF" if is_note else "F8FAFC"
            border_color = "2563EB" if is_note else "94A3B8"
            text_color = RGBColor(0x1e, 0x3a, 0x8a) if is_note else RGBColor(0x33, 0x41, 0x55)
            
            table = document.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            
            # Set background shading
            shading_xml = f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>'
            cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
            
            # Set left border only
            borders_xml = f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
            '''
            cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))
            
            # Set padding
            margins_xml = f'''
            <w:tcMar {nsdecls("w")}>
                <w:top w:w="160" w:type="dxa"/>
                <w:bottom w:w="160" w:type="dxa"/>
                <w:left w:w="240" w:type="dxa"/>
                <w:right w:w="240" w:type="dxa"/>
            </w:tcMar>
            '''
            cell._tc.get_or_add_tcPr().append(parse_xml(margins_xml))
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            text = "\n".join(block["lines"])
            process_text_runs_in_paragraph(p, text, force_italic=True, force_color=text_color)
            document.add_paragraph()
            
        elif b_type == "code":
            table = document.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            
            shading_xml = f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>'
            cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
            
            borders_xml = f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
            </w:tcBorders>
            '''
            cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))
            
            margins_xml = f'''
            <w:tcMar {nsdecls("w")}>
                <w:top w:w="160" w:type="dxa"/>
                <w:bottom w:w="160" w:type="dxa"/>
                <w:left w:w="160" w:type="dxa"/>
                <w:right w:w="160" w:type="dxa"/>
            </w:tcMar>
            '''
            cell._tc.get_or_add_tcPr().append(parse_xml(margins_xml))
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            content = block["content"].strip()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            document.add_paragraph()
            
        elif b_type == "mermaid_image":
            try:
                from PIL import Image as PILImage
                img_data = BytesIO(block["content"])
                pil_img = PILImage.open(img_data)
                w, _ = pil_img.size
                
                # Scaled width in inches (max 6.0)
                docx_width = Inches(min(6.0, w / 96.0))
                
                document.add_picture(img_data, width=docx_width)
                document.add_paragraph()
            except Exception:
                pass
            
        elif b_type == "table":
            raw_rows = []
            for line in block["lines"]:
                parts = [cell.strip() for cell in line.split("|")]
                if len(parts) >= 2:
                    cells = parts[1:-1]
                    raw_rows.append(cells)
            
            if not raw_rows:
                continue
            
            is_separator = False
            if len(raw_rows) > 1:
                second_row = raw_rows[1]
                if all(re.match(r"^:?-+:?$", cell) for cell in second_row):
                    is_separator = True
            
            cleaned_rows = []
            for idx, r in enumerate(raw_rows):
                if is_separator and idx == 1:
                    continue
                cleaned_rows.append(r)
            
            if not cleaned_rows:
                continue
                
            num_cols = max(len(r) for r in cleaned_rows)
            num_rows = len(cleaned_rows)
            
            table = document.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            for r_idx, row_cells in enumerate(cleaned_rows):
                row = table.rows[r_idx]
                for c_idx, cell_value in enumerate(row_cells):
                    if c_idx < len(row.cells):
                        cell = row.cells[c_idx]
                        
                        borders_xml = f'''
                        <w:tcBorders {nsdecls("w")}>
                            <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                            <w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                            <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                        </w:tcBorders>
                        '''
                        cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        if r_idx == 0:
                            shading_xml = f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>'
                            cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
                            process_text_runs_in_paragraph(p, cell_value, force_bold=True, force_color=RGBColor(0x1e, 0x29, 0x3b))
                        else:
                            process_text_runs_in_paragraph(p, cell_value, force_color=RGBColor(0x33, 0x41, 0x55))
            document.add_paragraph()

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.getvalue()


def _strip_duplicate_section_heading(content: str, section_title: str) -> str:
    if normalize_section_profile_key(section_title) == "main_content":
        return (content or "").strip()

    lines = (content or "").splitlines()
    if not lines:
        return content

    target = _normalize_heading_label(section_title)
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    if idx >= len(lines):
        return content

    first = lines[idx].strip()
    first = re.sub(r"^#{1,6}\s+", "", first)
    first = re.sub(r"^[\-*+]\s+", "", first)
    if _normalize_heading_label(first) == target:
        lines.pop(idx)
        while idx < len(lines) and not lines[idx].strip():
            lines.pop(idx)

    return "\n".join(lines).strip()


def _evaluate_depth_distribution(markdown_text: str) -> tuple[bool, str]:
    lines = (markdown_text or "").splitlines()
    buckets = {
        "objectives": [],
        "intro": [],
        "main": [],
        "examples": [],
        "summary": [],
    }
    current_bucket = "main"

    for line in lines:
        heading = re.match(r"^#{1,6}\s+(.+)$", line.strip())
        if heading:
            label = _normalize_heading_label(heading.group(1))
            if any(k in label for k in ["learning objectives", "muc tieu hoc tap", "objectives"]):
                current_bucket = "objectives"
            elif any(k in label for k in ["introduction", "gioi thieu", "overview"]):
                current_bucket = "intro"
            elif any(k in label for k in ["main content", "noi dung chinh", "main concepts"]):
                current_bucket = "main"
            elif any(k in label for k in ["example", "vi du", "application", "ung dung"]):
                current_bucket = "examples"
            elif any(k in label for k in ["summary", "tom tat", "tong ket"]):
                current_bucket = "summary"
            continue

        buckets[current_bucket].append(line)

    def _word_count(block_lines: list[str]) -> int:
        text = " ".join(block_lines)
        return len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))

    objectives_words = _word_count(buckets["objectives"])
    intro_words = _word_count(buckets["intro"])
    main_words = _word_count(buckets["main"])
    example_words = _word_count(buckets["examples"])
    summary_words = _word_count(buckets["summary"])

    total = objectives_words + intro_words + main_words + example_words + summary_words
    if total <= 0:
        return False, "empty content"

    main_ratio = main_words / total
    example_ratio = example_words / total
    summary_ratio = summary_words / total
    intro_obj_ratio = (intro_words + objectives_words) / total

    if main_words < 180:
        return False, "main content too short"
    if main_ratio < 0.45:
        return False, "main content ratio too low"
    if example_ratio < 0.12:
        return False, "examples ratio too low"
    if summary_ratio < 0.06:
        return False, "summary ratio too low"
    if intro_obj_ratio > 0.30:
        return False, "intro/objectives too long"
    return True, "ok"


def _extract_verification_verdict(markdown_text: str) -> tuple[str, list[str]]:
    text = markdown_text or ""
    verdict = "UNKNOWN"

    verdict_match = re.search(
        r"(?im)^\s*(?:#{1,6}\s*)?(?:verdict|ket luan|kết luận)\s*:\s*(.+)$",
        text,
    )
    if verdict_match:
        verdict_value = verdict_match.group(1).strip().upper()
        if "INVALID" in verdict_value:
            verdict = "INVALID"
        elif "VALID" in verdict_value:
            verdict = "VALID"

    ooc_concepts: list[str] = []
    ooc_match = re.search(
        r"(?ims)^\s*(?:#{1,6}\s*)?(?:out-of-context concepts|khai niem vuot ngu canh|khái niệm vượt ngữ cảnh)\s*:\s*(.+?)(?:\n\s*(?:#{1,6}\s*)?(?:verdict|ket luan|kết luận)\s*:|\Z)",
        text,
    )
    if ooc_match:
        raw_ooc = ooc_match.group(1).strip()
        for item in re.split(r"\n|,|;", raw_ooc):
            clean = re.sub(r"^[\-*+]\s*", "", item).strip()
            if clean and clean.lower() not in {"none", "khong co", "không có", "n/a"}:
                ooc_concepts.append(clean)

    return verdict, ooc_concepts


def _strip_verification_block(markdown_text: str) -> str:
    text = markdown_text or ""
    marker = re.search(
        r"(?im)^\s*(?:#{1,6}\s*)?(?:"
        r"phase\s*[123]\b|"
        r"content type|main topic|key concepts|context quality|missing information|example needs|hallucination risk|suggested retrieval queries|"
        r"boundary check|out-of-context concepts|verdict|"
        r"loai noi dung|chu de chinh|khai niem cot loi|chat luong ngu canh|thong tin thieu|nhu cau vi du|rui ro ao giac|truy van truy xuat de xuat|"
        r"kiem tra bien|khai niem vuot ngu canh"
        r")\s*[:\-]?",
        text,
    )
    if not marker:
        return text.strip()
    return text[: marker.start()].rstrip()


def _filter_unbounded_derived_sections(markdown_text: str) -> str:
    lines = (markdown_text or "").splitlines()
    filtered: list[str] = []
    skip_block = False

    for line in lines:
        heading = re.match(r"^\s*#{1,6}\s+(.+)$", line)
        if heading:
            label = _normalize_heading_label(heading.group(1))
            skip_block = any(
                key in label
                for key in [
                    "example",
                    "vi du",
                    "application",
                    "ung dung",
                    "practice question",
                    "cau hoi on tap",
                    "cau hoi",
                    "exercise",
                    "bai tap",
                ]
            )
            if skip_block:
                continue

        if skip_block:
            continue
        filtered.append(line)

    output = "\n".join(filtered).strip()
    if not output:
        return "Nội dung bị giới hạn do vi phạm biên kiến thức của ngữ cảnh."

    output += (
        "\n\n## Giới hạn nội dung\n"
        "Một số phần Ví dụ/Câu hỏi/Bài tập đã được lược bỏ vì vượt quá biên kiến thức của ngữ cảnh."
    )
    return output


def _classify_section_kind(section_title: str) -> str:
    label = _normalize_heading_label(section_title)
    if any(key in label for key in ["main content", "noi dung chinh", "main concepts"]):
        return "main"
    if any(key in label for key in ["example", "vi du", "ung dung", "application"]):
        return "examples"
    if any(key in label for key in ["summary", "tom tat", "tong ket", "ket luan"]):
        return "summary"
    if any(
        key in label
        for key in [
            "practice question",
            "review question",
            "cau hoi",
            "on tap",
            "quiz",
            "trac nghiem",
            "exercise",
            "bai tap",
        ]
    ):
        return "questions"
    if any(key in label for key in ["learning objective", "muc tieu hoc tap", "muc tieu", "objective"]):
        return "objectives"
    if any(key in label for key in ["introduction", "gioi thieu", "overview", "mo dau", "dan nhap"]):
        return "intro"
    return "other"


def _section_disallowed_kinds(current_kind: str) -> set[str]:
    disallowed_map = {
        "main": {"examples", "summary", "questions", "objectives", "intro"},
        "examples": {"main", "summary", "questions", "objectives", "intro"},
        "summary": {"main", "examples", "questions", "objectives", "intro"},
        "questions": {"main", "examples", "summary", "objectives", "intro"},
        "objectives": {"main", "examples", "summary", "questions", "intro"},
        "intro": {"main", "examples", "summary", "questions", "objectives"},
    }
    return disallowed_map.get(current_kind, set())


def _detect_cross_section_leakage(section_title: str, markdown_text: str) -> list[str]:
    text = markdown_text or ""
    current_kind = _classify_section_kind(section_title)
    disallowed_kinds = _section_disallowed_kinds(current_kind)

    issues: list[str] = []
    for raw_line in text.splitlines():
        label = ""

        heading = re.match(r"^\s*#{1,6}\s+(.+?)\s*$", raw_line)
        if heading:
            label = heading.group(1).strip()
        else:
            numbered = re.match(r"^\s*\d+(?:\.\d+)*(?:[\.)-])?\s+(.+?)\s*$", raw_line)
            if numbered:
                label = numbered.group(1).strip()
            else:
                chapter_like = re.match(
                    r"^\s*(?:chapter|part|section|chuong|chương|phan|phần)\s*[:\-]?\s*\d*(?:\.\d+)*\s*[:\-]?\s*(.+?)\s*$",
                    raw_line,
                    flags=re.IGNORECASE,
                )
                if chapter_like:
                    label = chapter_like.group(1).strip()

        if not label:
            continue

        heading_kind = _classify_section_kind(label)
        if heading_kind in disallowed_kinds:
            issues.append(f'cross-section heading "{label}"')

    scaffold_pattern = re.compile(
        r"(?im)^\s*(?:#{1,6}\s*)?(?:phase\s*[123]|content type|main topic|key concepts|context quality|"
        r"missing information|example needs|hallucination risk|suggested retrieval queries|boundary check|"
        r"out-of-context concepts|verdict|kiểm tra biên|khái niệm vượt ngữ cảnh|kết luận)\s*[:\-]",
    )
    if scaffold_pattern.search(text):
        issues.append("audit/verification scaffold detected")

    return list(dict.fromkeys(issues))


def _strip_disallowed_section_blocks(section_title: str, markdown_text: str) -> str:
    current_kind = _classify_section_kind(section_title)
    disallowed_kinds = _section_disallowed_kinds(current_kind)
    if not disallowed_kinds:
        return (markdown_text or "").strip()

    kept_lines: list[str] = []
    skip_block = False

    for line in (markdown_text or "").splitlines():
        heading = re.match(r"^\s*#{1,6}\s+(.+?)\s*$", line)
        if heading:
            heading_kind = _classify_section_kind(heading.group(1).strip())
            skip_block = heading_kind in disallowed_kinds
            if skip_block:
                continue

        if skip_block:
            continue
        kept_lines.append(line)

    cleaned = "\n".join(kept_lines).strip()
    return cleaned or (markdown_text or "").strip()


def _truncate_words(text: str, max_words: int) -> str:
    words = re.findall(r"\S+", text or "", flags=re.UNICODE)
    if not words:
        return ""
    return " ".join(words[:max_words]).strip()


def _is_likely_cut_sentence(text: str) -> bool:
    sentence = re.sub(r"\s+", " ", (text or "")).strip()
    if not sentence:
        return True

    if re.search(r"[\(\[\{,;:/\-]\s*$", sentence):
        return True

    normalized = _normalize_heading_label(sentence)
    tokens = normalized.split()
    if not tokens:
        return True

    trailing_connectors = {
        "hoac",
        "va",
        "voi",
        "tu",
        "de",
        "khi",
        "la",
        "nhu",
        "gom",
        "bao",
        "including",
        "with",
        "from",
        "and",
        "or",
        "using",
        "use",
    }
    if tokens[-1] in trailing_connectors:
        return True

    # Sentence tails ending near conjunctions are often truncated chunk fragments.
    if any(token in {"hoac", "and", "or"} for token in tokens[-8:]):
        return True

    if re.search(
        r"\b(?:sử\s*dụng|su\s*dung|áp\s*dụng|ap\s*dung|bao\s*gồm|bao\s*gom|gồm|gom)\s+các\s*$",
        sentence,
        flags=re.IGNORECASE,
    ):
        return True

    return False


def _ensure_complete_sentence(text: str, fallback: str = "") -> str:
    sentence = re.sub(r"\s+", " ", (text or "")).strip()
    fallback_sentence = re.sub(r"\s+", " ", (fallback or "")).strip()

    if not sentence:
        sentence = fallback_sentence
    if not sentence:
        return ""

    sentence = re.sub(r"[\(\[\{,;:/\-]\s*$", "", sentence).strip()
    sentence = re.sub(
        r"\b(?:hoặc|hoac|và|va|với|voi|từ|tu|để|de|khi|là|la|như|nhu|gồm|gom|bao\s*gồm|bao\s*gom|including|with|from|and|or|using|use)\s*$",
        "",
        sentence,
        flags=re.IGNORECASE,
    ).strip()

    if _is_likely_cut_sentence(sentence):
        sentence = fallback_sentence or sentence

    if not sentence:
        return ""

    if not re.search(r"[.!?]$", sentence):
        if _is_likely_cut_sentence(sentence):
            sentence = f"{sentence}, phù hợp với ngữ cảnh của tài liệu."
        else:
            sentence = f"{sentence}."

    return sentence


def _extract_control_sentinel(text: str) -> str:
    match = re.search(r"\b(NOT_ENOUGH_CONTEXT|FAIL_COVERAGE)\b", text or "", flags=re.IGNORECASE)
    if not match:
        return ""
    return str(match.group(1)).upper()


def _parse_section_json_response(raw: str) -> tuple[str, str]:
    """Parse LLM JSON response for section generation.

    Returns:
        (content, sentinel) where:
        - content: cleaned Markdown string (may be empty if sentinel is set)
        - sentinel: "" | "NOT_ENOUGH_CONTEXT" | "FAIL_COVERAGE"

    Falls back to legacy raw-Markdown handling if LLM ignores the JSON contract.
    """
    text = (raw or "").strip()
    if not text:
        return "", "NOT_ENOUGH_CONTEXT"

    # Strip common markdown fences that some models prepend/append
    cleaned = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()

    # ── Try strict JSON parse ────────────────────────────────────────────────
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            content_raw = str(parsed.get("content") or "")
            sentinel_raw = str(parsed.get("sentinel") or "").strip().upper()
            sentinel = sentinel_raw if sentinel_raw in {"NOT_ENOUGH_CONTEXT", "FAIL_COVERAGE"} else ""
            return content_raw.strip(), sentinel
    except (json.JSONDecodeError, ValueError):
        pass

    # ── Try to extract a JSON object embedded in surrounding text ────────────
    obj_match = re.search(r"\{[\s\S]*\}", cleaned)
    if obj_match:
        try:
            parsed = json.loads(obj_match.group())
            if isinstance(parsed, dict):
                content_raw = str(parsed.get("content") or "")
                sentinel_raw = str(parsed.get("sentinel") or "").strip().upper()
                sentinel = sentinel_raw if sentinel_raw in {"NOT_ENOUGH_CONTEXT", "FAIL_COVERAGE"} else ""
                return content_raw.strip(), sentinel
        except (json.JSONDecodeError, ValueError):
            pass

    # ── Try to extract content via Regex if JSON is malformed (e.g. missing commas) ──
    content_match = re.search(r'"content"\s*:\s*"((?:[^"\\]|\\.)*)"', cleaned, re.DOTALL)
    if content_match:
        try:
            content_val = json.loads(f'"{content_match.group(1)}"')
            sentinel = ""
            sentinel_match = re.search(r'"sentinel"\s*:\s*"([^"]*)"', cleaned)
            if sentinel_match:
                sentinel_raw = sentinel_match.group(1).strip().upper()
                if sentinel_raw in {"NOT_ENOUGH_CONTEXT", "FAIL_COVERAGE"}:
                    sentinel = sentinel_raw
            return content_val.strip(), sentinel
        except Exception:
            pass

    # If the text itself looks like a raw JSON object string representation but failed parsing,
    # let's try to unescape it or clean up quotes.
    if text.startswith("{") and text.endswith("}"):
        # Check if it contains escaped quotes like \"content\"
        if '\\"' in text:
            try:
                # Replace escaped quotes and attempt parse
                unescaped = text.replace('\\"', '"')
                parsed = json.loads(unescaped)
                if isinstance(parsed, dict):
                    content_raw = str(parsed.get("content") or "")
                    sentinel_raw = str(parsed.get("sentinel") or "").strip().upper()
                    sentinel = sentinel_raw if sentinel_raw in {"NOT_ENOUGH_CONTEXT", "FAIL_COVERAGE"} else ""
                    return content_raw.strip(), sentinel
            except Exception:
                pass

    # ── Legacy fallback: LLM returned raw Markdown (pre-JSON contract) ───────
    logger.warning("_parse_section_json_response: LLM returned non-JSON; using raw Markdown fallback")
    sentinel_fallback = _extract_control_sentinel(text)
    return text if not sentinel_fallback else "", sentinel_fallback


def _map_control_sentinel_to_user_message(sentinel: str, section_title: str) -> str:
    normalized = str(sentinel or "").strip().upper()
    safe_title = (section_title or "section này").strip() or "section này"

    if normalized == "NOT_ENOUGH_CONTEXT":
        return (
            f"Chưa đủ ngữ cảnh để tạo nội dung cho phần '{safe_title}'. "
            "Hãy bổ sung tài liệu nguồn hoặc điều chỉnh prompt cụ thể hơn."
        )

    if normalized == "FAIL_COVERAGE":
        return (
            f"Ngữ cảnh hiện tại chưa bao phủ đủ ý chính cho phần '{safe_title}'. "
            "Hãy bổ sung tài liệu liên quan và tạo lại."
        )

    return ""


def _detect_structure_lock_violations(markdown_text: str) -> list[str]:
    issues: list[str] = []

    for raw_line in (markdown_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading_match = re.match(r"^#{1,6}\s+(.+)$", line)
        if heading_match:
            line = heading_match.group(1).strip()

        if re.match(r"^\d+(?:\.\d+)*(?:[\.)-])?\s+", line):
            issues.append("forbidden numeric heading prefix")
            continue

        if re.match(r"^(?:chapter|part|section|chuong|chương|phan|phần)\b", line, flags=re.IGNORECASE):
            issues.append("forbidden chapter/part heading prefix")

    return list(dict.fromkeys(issues))


def _apply_structure_lock(markdown_text: str) -> str:
    sentinel = _extract_control_sentinel(markdown_text)
    if sentinel:
        return sentinel

    locked_lines: list[str] = []
    for raw_line in (markdown_text or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            locked_lines.append(line)
            continue

        heading_match = re.match(r"^(\s*#{1,6}\s+)(.+)$", line)
        if heading_match:
            prefix = heading_match.group(1)
            heading_text = heading_match.group(2).strip()
            heading_text = re.sub(r"^\d+(?:\.\d+)*(?:[\.)-])?\s+", "", heading_text)
            heading_text = re.sub(
                r"^(?:chapter|part|section|chuong|chương|phan|phần)\s*[:\-]?\s*\d*(?:\.\d+)*\s*[:\-]?\s*",
                "",
                heading_text,
                flags=re.IGNORECASE,
            ).strip()
            if heading_text:
                locked_lines.append(f"{prefix}{heading_text}")
            continue

        numbered = re.match(r"^\s*\d+(?:\.\d+)*(?:[\.)-])?\s+(.+)$", line)
        if numbered:
            body = numbered.group(1).strip()
            if body:
                locked_lines.append(f"- {body}")
            continue

        chapter_like = re.match(
            r"^\s*(?:chapter|part|section|chuong|chương|phan|phần)\s*[:\-]?\s*\d*(?:\.\d+)*\s*[:\-]?\s*(.+)$",
            line,
            flags=re.IGNORECASE,
        )
        if chapter_like:
            body = chapter_like.group(1).strip()
            if body:
                locked_lines.append(f"- {body}")
            continue

        locked_lines.append(line)

    return "\n".join(locked_lines).strip()


def _clean_line_candidate(raw_line: str) -> str:
    line = (raw_line or "").strip()
    if not line:
        return ""

    line = re.sub(r"^#{1,6}\s+", "", line)
    line = re.sub(r"^[-*+]\s+", "", line)
    line = re.sub(r"^\d+[\.)]\s+", "", line)
    line = re.sub(r"^(?:Q\s*)?\d+\s*[:\)\.-]\s*", "", line, flags=re.IGNORECASE)
    line = re.sub(r"^[A-D][\.:\)]\s*", "", line, flags=re.IGNORECASE)
    line = re.sub(r"^(?:answer|dap an|đáp án)\s*[:\-]\s*", "", line, flags=re.IGNORECASE)
    line = re.sub(r"\s+", " ", line)
    return line.strip(" -:;,.")


def _dedupe_text_items(items: list[str]) -> list[str]:
    deduped: list[str] = []
    seen: set[str] = set()
    for item in items:
        normalized = _normalize_heading_label(item)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(item)
    return deduped


def _extract_line_candidates(markdown_text: str, min_len: int = 4) -> list[str]:
    candidates: list[str] = []
    scaffold_re = re.compile(
        r"^(?:phase\s*[123]|content type|main topic|key concepts|context quality|"
        r"missing information|example needs|hallucination risk|suggested retrieval queries|"
        r"boundary check|out-of-context concepts|verdict)\b",
        flags=re.IGNORECASE,
    )

    for raw_line in (markdown_text or "").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("```"):
            continue
        if scaffold_re.match(_clean_line_candidate(line)):
            continue
        cleaned = _clean_line_candidate(line)
        if len(cleaned) >= min_len:
            candidates.append(cleaned)

    if not candidates:
        for chunk in re.split(r"[\.\n;!?]+", markdown_text or ""):
            cleaned = _clean_line_candidate(chunk)
            if len(cleaned) >= min_len:
                candidates.append(cleaned)

    return _dedupe_text_items(candidates)


def _enforce_objective_hard_format(markdown_text: str) -> str:
    source = _extract_line_candidates(markdown_text)
    if not source:
        fallback = _clean_line_candidate(markdown_text)
        source = [fallback] if fallback else []
    if not source:
        return (markdown_text or "").strip()

    verbs = ["Hiểu", "Nắm được", "Phân biệt", "Áp dụng", "Thực hiện"]
    objective_phrases: list[str] = []
    for item in source:
        phrase = re.sub(
            r"^(?:(?:Hiểu|Nắm được|Phân biệt|Áp dụng|Thực hiện)\s+)+",
            "",
            item,
            flags=re.IGNORECASE,
        )
        phrase = phrase.strip(" .;,:")
        if phrase:
            objective_phrases.append(phrase)
        if len(objective_phrases) >= 5:
            break

    if not objective_phrases:
        objective_phrases = ["nội dung cốt lõi trong tài liệu"]

    target_count = min(5, max(3, len(objective_phrases)))
    while len(objective_phrases) < target_count:
        objective_phrases.append(objective_phrases[-1])

    bullets: list[str] = []
    for idx in range(target_count):
        phrase = objective_phrases[idx].strip(" .;,:")
        if not phrase:
            phrase = "nội dung cốt lõi trong tài liệu"
        phrase = _ensure_complete_sentence(phrase, fallback="Nội dung cốt lõi trong tài liệu.")
        phrase = re.sub(
            r"^(?:(?:Hiểu|Nắm được|Phân biệt|Áp dụng|Thực hiện)\s+)+",
            "",
            phrase,
            flags=re.IGNORECASE,
        )
        verb = verbs[idx % len(verbs)]
        bullets.append(f"- {verb} {phrase}")
    return "\n".join(bullets).strip()


def _enforce_summary_hard_format(markdown_text: str) -> str:
    min_points = 3
    sparse_min_points = 2
    max_points = 5
    source = _extract_line_candidates(markdown_text, min_len=10)
    if not source:
        fallback = _clean_line_candidate(markdown_text)
        source = [fallback] if fallback else []
    if not source:
        return (markdown_text or "").strip()

    def _to_sentence(text: str) -> str:
        sentence = re.sub(r"\s+", " ", (text or "").strip(" .;,:-"))
        if not sentence:
            return ""
        if len(re.findall(r"\S+", sentence, flags=re.UNICODE)) < 4:
            return ""
        if sentence and sentence[0].islower():
            sentence = sentence[0].upper() + sentence[1:]
        if not re.search(r"[.!?]$", sentence):
            sentence = f"{sentence}."
        return sentence

    summary_points: list[str] = []
    for item in source:
        if item.endswith("?"):
            continue
        phrase = re.sub(r"^(?:Hiểu|Nắm được|Phân biệt|Áp dụng)\s+", "", item, flags=re.IGNORECASE)
        phrase = _to_sentence(phrase)
        if phrase:
            summary_points.append(phrase)
        if len(summary_points) >= 12:
            break

    summary_points = _dedupe_text_items(summary_points)

    if len(summary_points) < min_points:
        sentence_candidates = re.split(r"(?<=[.!?])\s+|\n+", markdown_text or "")
        for item in sentence_candidates:
            phrase = _to_sentence(_clean_line_candidate(item))
            if phrase:
                summary_points.append(phrase)
            if len(summary_points) >= 12:
                break
        summary_points = _dedupe_text_items(summary_points)

    if not summary_points:
        return "NOT_ENOUGH_CONTEXT"

    if len(summary_points) < sparse_min_points:
        return "NOT_ENOUGH_CONTEXT"

    target_floor = min_points if len(summary_points) >= min_points else sparse_min_points
    target_count = min(max_points, max(target_floor, len(summary_points)))
    return "\n".join(f"- {item}" for item in summary_points[:target_count]).strip()


_QUIZ_CONCEPT_QUERY_MAP: list[tuple[str, tuple[str, ...]]] = [
    ("Concept Definition", ("định nghĩa", "khái niệm", "là gì", "definition", "concept")),
    ("Mechanism/Operation", ("hoạt động", "vận hành", "cách dùng", "how it works", "mechanism")),
    ("Constraints/Rules", ("quy tắc", "điều kiện", "ràng buộc", "rules", "constraints")),
    ("Application/Example", ("ví dụ", "thực tế", "tình huống", "example", "application")),
    ("Comparison", ("phân biệt", "so sánh", "khác nhau", "difference", "comparison")),
]


def _detect_quiz_focus_queries(text: str) -> list[str]:
    normalized = f" {_normalize_heading_label(text)} "
    queries: list[str] = []
    for query, aliases in _QUIZ_CONCEPT_QUERY_MAP:
        if any(f" {alias} " in normalized for alias in aliases):
            queries.append(query)
    return queries


def _build_quiz_source_records(chunks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    seen_chunk_ids: set[str] = set()

    for item in chunks or []:
        chunk_id = str(item.get("id") or item.get("chunk_id") or "").strip()
        if not chunk_id or chunk_id in seen_chunk_ids:
            continue

        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
        metadata = dict(metadata)
        if not metadata.get("file_name"):
            metadata["file_name"] = _metadata_source_label(metadata)

        source_text = _strip_metadata_heading_prefix(str(item.get("text") or ""), metadata).strip()
        file_name = str(metadata.get("file_name") or _metadata_source_label(metadata) or "").strip()
        chapter = _clean_structure_label(
            metadata.get("chapter_title") or metadata.get("chapter") or metadata.get("h1"),
            "chapter",
        ) or ""
        section = _clean_structure_label(
            metadata.get("section_title") or metadata.get("section") or metadata.get("h2"),
            "section",
        ) or ""
        subsection = _clean_structure_label(
            metadata.get("subsection_title") or metadata.get("subsection") or metadata.get("h3"),
            "subsection",
        ) or ""
        start_page, end_page = _metadata_page_range(metadata)
        score = float(item.get("rerank_score", item.get("hybrid_score", item.get("_score", 0.0))))

        payload = " ".join([file_name, chapter, section, subsection, source_text])
        normalized_payload = _normalize_heading_label(payload)
        tokens = {
            token
            for token in re.split(r"\W+", normalized_payload)
            if len(token) >= 2 and not token.isdigit()
        }

        records.append(
            {
                "chunk_id": chunk_id,
                "score": score,
                "file_name": file_name,
                "chapter": chapter,
                "section": section,
                "subsection": subsection,
                "start_page": start_page,
                "end_page": end_page,
                "tokens": tokens,
                "normalized_payload": normalized_payload,
            }
        )
        seen_chunk_ids.add(chunk_id)

    records.sort(key=lambda item: float(item.get("score", 0.0)), reverse=True)
    return records


def _merge_quiz_source_records(
    base_records: list[dict[str, Any]],
    extra_records: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    seen: set[str] = set()

    for record in [*base_records, *extra_records]:
        chunk_id = str(record.get("chunk_id") or "").strip()
        if not chunk_id or chunk_id in seen:
            continue
        merged.append(record)
        seen.add(chunk_id)

    merged.sort(key=lambda item: float(item.get("score", 0.0)), reverse=True)
    return merged


def _retrieve_quiz_focus_chunks_for_question(
    question_text: str,
    selected_source_docs: list[dict[str, Any]] | None,
) -> list[dict[str, Any]]:
    docs = selected_source_docs or []
    if not docs:
        return []

    focus_queries = _detect_quiz_focus_queries(question_text)
    if not focus_queries:
        focus_queries = [question_text]

    merged_by_id: dict[str, dict[str, Any]] = {}

    for query in focus_queries[:2]:
        focused = _retrieve_context_with_retry(
            query=query,
            selected_source_docs=docs,
            min_total_chars=120,
            max_chunks=3,
            top_k_levels=[3],
        )
        for item in focused:
            chunk_id = str(item.get("id") or item.get("chunk_id") or "").strip()
            if not chunk_id:
                continue
            score = float(item.get("rerank_score", item.get("hybrid_score", item.get("_score", 0.0))))
            prev = merged_by_id.get(chunk_id)
            if prev is None or float(prev.get("_score", -1.0)) < score:
                merged_by_id[chunk_id] = {**item, "_score": score}

    merged = list(merged_by_id.values())
    merged.sort(key=lambda item: float(item.get("_score", 0.0)), reverse=True)
    return merged[:3]


def _ensure_min_quiz_source_records(
    source_records: list[dict[str, Any]],
    required_count: int,
    selected_source_docs: list[dict[str, Any]] | None,
    question_pool: list[str],
) -> list[dict[str, Any]]:
    docs = selected_source_docs or []
    unique_count = len({str(item.get("chunk_id") or "") for item in source_records if item.get("chunk_id")})
    if unique_count >= required_count or not docs:
        return source_records

    query_candidates: list[str] = []
    for question in question_pool[:required_count]:
        focus_queries = _detect_quiz_focus_queries(question)
        query_candidates.extend(focus_queries)
        query_candidates.append(question)
    query_candidates.extend(query for query, _aliases in _QUIZ_CONCEPT_QUERY_MAP)

    dedup_queries: list[str] = []
    seen_queries: set[str] = set()
    for query in query_candidates:
        normalized = _normalize_heading_label(query)
        if not normalized or normalized in seen_queries:
            continue
        seen_queries.add(normalized)
        dedup_queries.append(query)

    records = source_records[:]
    for query in dedup_queries:
        unique_count = len({str(item.get("chunk_id") or "") for item in records if item.get("chunk_id")})
        if unique_count >= required_count:
            break

        focused_chunks = _retrieve_context_with_retry(
            query=query,
            selected_source_docs=docs,
            min_total_chars=120,
            max_chunks=6,
            top_k_levels=[3, 4, 5],
        )
        focused_records = _build_quiz_source_records(focused_chunks)
        records = _merge_quiz_source_records(records, focused_records)

    return records


def _format_source_citation_line(record: dict[str, Any] | None) -> str:
    if not record:
        return "📚 Nguồn: Tài liệu đã chọn"

    file_name = str(record.get("file_name") or "").strip() or "Tài liệu đã chọn"
    parts = [
        str(record.get("chapter") or "").strip(),
        str(record.get("section") or "").strip(),
        str(record.get("subsection") or "").strip(),
    ]
    parts = [part for part in parts if part]

    source_detail = file_name
    if parts:
        source_detail = f"{source_detail} – {', '.join(parts)}"

    start_page = _coerce_page_number(record.get("start_page"))
    end_page = _coerce_page_number(record.get("end_page"))
    page_label = _format_page_range_label(start_page, end_page)
    if page_label:
        page_label = re.sub(r"^tr\.\s*", "Trang ", page_label, flags=re.IGNORECASE)
        source_detail = f"{source_detail} ({page_label})"

    chunk_id = str(record.get("chunk_id") or "").strip()
    if chunk_id:
        safe_chunk_id = quote(chunk_id, safe="")
        safe_label = source_detail.replace("]", r"\]")
        return f"📚 Nguồn: [{safe_label}](#source:{safe_chunk_id})"

    return f"📚 Nguồn: {source_detail}"


def _select_quiz_source_record_for_question(
    question_text: str,
    explanation_text: str,
    source_records: list[dict[str, Any]],
    used_chunk_ids: set[str],
    selected_source_docs: list[dict[str, Any]] | None,
) -> tuple[dict[str, Any] | None, list[dict[str, Any]]]:
    combined = f"{question_text}\n{explanation_text}".strip()
    normalized_question = _normalize_heading_label(combined)
    question_tokens = {
        token
        for token in re.split(r"\W+", normalized_question)
        if len(token) >= 2 and not token.isdigit()
    }
    concept_queries = _detect_quiz_focus_queries(combined)
    concept_tokens = {
        _normalize_heading_label(query.replace("SQL", "").strip())
        for query in concept_queries
        if query
    }

    def _record_rank(record: dict[str, Any]) -> tuple[int, int, int, float]:
        chunk_id = str(record.get("chunk_id") or "")
        is_unused = 1 if chunk_id and chunk_id not in used_chunk_ids else 0
        payload_tokens = set(record.get("tokens") or set())
        overlap = len(question_tokens & payload_tokens)
        normalized_payload = str(record.get("normalized_payload") or "")
        concept_hit = 1 if any(concept in normalized_payload for concept in concept_tokens) else 0
        score = float(record.get("score", 0.0))
        return (is_unused, concept_hit, overlap, score)

    ranked = sorted(source_records, key=_record_rank, reverse=True)
    for record in ranked:
        chunk_id = str(record.get("chunk_id") or "")
        if not chunk_id or chunk_id in used_chunk_ids:
            continue
        if concept_tokens:
            if _record_rank(record)[1] > 0:
                return record, source_records
            continue
        if _record_rank(record)[2] > 0:
            return record, source_records

    focused_chunks = _retrieve_quiz_focus_chunks_for_question(
        question_text=combined,
        selected_source_docs=selected_source_docs,
    )
    if focused_chunks:
        focused_records = _build_quiz_source_records(focused_chunks)
        source_records = _merge_quiz_source_records(source_records, focused_records)
        ranked = sorted(source_records, key=_record_rank, reverse=True)
        for record in ranked:
            chunk_id = str(record.get("chunk_id") or "")
            if not chunk_id or chunk_id in used_chunk_ids:
                continue
            return record, source_records

    for record in ranked:
        chunk_id = str(record.get("chunk_id") or "")
        if chunk_id and chunk_id not in used_chunk_ids:
            return record, source_records

    if ranked:
        return ranked[0], source_records
    return None, source_records


def _enforce_quiz_hard_format(
    markdown_text: str,
    retrieved_chunks: list[dict[str, Any]] | None = None,
    selected_source_docs: list[dict[str, Any]] | None = None,
) -> str:
    cleaned = re.sub(r"^```(?:json)?\s*", "", (markdown_text or "").strip(), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()
    
    try:
        data = json.loads(cleaned)
    except Exception:
        # Fallback if the AI failed to produce valid JSON
        return markdown_text

    mcq_list = data.get("mcq", [])
    essay_list = data.get("essay", [])
    
    question_pool: list[str] = []
    for item in mcq_list:
        question_pool.append(str(item.get("question") or ""))
    for item in essay_list:
        question_pool.append(str(item.get("question") or ""))
        
    source_records = _build_quiz_source_records(retrieved_chunks or [])
    source_records = _ensure_min_quiz_source_records(
        source_records=source_records,
        required_count=len(question_pool) or 6,
        selected_source_docs=selected_source_docs,
        question_pool=question_pool,
    )
    
    used_chunk_ids: set[str] = set()
    lines: list[str] = ["## Câu hỏi trắc nghiệm", ""]
    
    for idx, item in enumerate(mcq_list):
        question_no = idx + 1
        question = str(item.get("question") or f"Câu hỏi {question_no}")
        bloom_level = str(item.get("bloom_level") or "Hiểu")
        options = item.get("options") or {}
        ans_A = str(options.get("A") or "")
        ans_B = str(options.get("B") or "")
        ans_C = str(options.get("C") or "")
        ans_D = str(options.get("D") or "")
        answer = str(item.get("answer") or "A")
        explanation = str(item.get("explanation") or "")
        wrong_analysis = item.get("wrong_analysis") or {}
        
        selected_record, source_records = _select_quiz_source_record_for_question(
            question_text=question,
            explanation_text=explanation,
            source_records=source_records,
            used_chunk_ids=used_chunk_ids,
            selected_source_docs=selected_source_docs,
        )
        chunk_id = str(selected_record.get("chunk_id") or "") if selected_record else ""
        if chunk_id:
            used_chunk_ids.add(chunk_id)
            
        lines.append(f"**Câu hỏi {question_no}:** {question}")
        lines.append(f"🎯 Mục tiêu: {bloom_level}")
        lines.append(f"A. {ans_A}")
        lines.append(f"B. {ans_B}")
        lines.append(f"C. {ans_C}")
        lines.append(f"D. {ans_D}")
        lines.append("")
        lines.append(f"**Đáp án:** {answer}")
        lines.append(f"**Giải thích:** {explanation}")
        
        if wrong_analysis:
            lines.append("**Phân tích đáp án sai:**")
            for letter, reason in wrong_analysis.items():
                lines.append(f"- {letter}: {reason}")
                
        lines.append("")
        lines.append(_format_source_citation_line(selected_record))
        lines.append("")
        lines.append("---")
        lines.append("")
        
    lines.extend(["## Câu hỏi tự luận", ""])
    
    for idx, item in enumerate(essay_list):
        question_no = len(mcq_list) + idx + 1
        question = str(item.get("question") or f"Câu hỏi {question_no}")
        bloom_level = str(item.get("bloom_level") or "Áp dụng")
        model_answer = str(item.get("model_answer") or "")
        guidance = str(item.get("guidance") or "")
        pedagogical_insight = str(item.get("pedagogical_insight") or "")
        
        explanation_for_search = f"{model_answer}\n{guidance}\n{pedagogical_insight}"
        
        selected_record, source_records = _select_quiz_source_record_for_question(
            question_text=question,
            explanation_text=explanation_for_search,
            source_records=source_records,
            used_chunk_ids=used_chunk_ids,
            selected_source_docs=selected_source_docs,
        )
        chunk_id = str(selected_record.get("chunk_id") or "") if selected_record else ""
        if chunk_id:
            used_chunk_ids.add(chunk_id)
            
        lines.append(f"**Câu hỏi {question_no}:** {question}")
        lines.append(f"🎯 Mục tiêu: {bloom_level}")
        lines.append(f"**Gợi ý đáp án:** {model_answer}")
        if guidance:
            lines.append(f"**Gợi ý hướng dẫn (dành cho giảng viên):** {guidance}")
        if pedagogical_insight:
            lines.append(f"> 🏫 **Nhận xét sư phạm:** {pedagogical_insight}")
            
        lines.append("")
        lines.append(_format_source_citation_line(selected_record))
        
        if idx < len(essay_list) - 1:
            lines.append("")
            lines.append("---")
            lines.append("")
            
    return "\n".join(lines).strip()


def _clamp_ten_score(value: Any, default: float = 0.0) -> float:
    try:
        score = float(value)
    except (TypeError, ValueError):
        score = default
    return round(max(0.0, min(10.0, score)), 1)


def _normalize_eval_list(value: Any, max_items: int = 8) -> list[str]:
    raw_items: list[Any]
    if isinstance(value, list):
        raw_items = value
    elif isinstance(value, str):
        raw_items = re.split(r"\n|;", value)
    else:
        raw_items = []

    cleaned: list[str] = []
    for item in raw_items:
        text = re.sub(r"^[\-*+•]\s*", "", str(item or "").strip())
        if not text:
            continue
        cleaned.append(text)
        if len(cleaned) >= max_items:
            break
    return cleaned


def _fallback_section_evaluation(section_name: str, generated_content: str) -> dict[str, Any]:
    sentinel = _extract_control_sentinel(generated_content)
    if sentinel == "NOT_ENOUGH_CONTEXT":
        return {
            "is_fallback": True,
            "scores": {
                "accuracy": 0.0,
                "coverage": 0.0,
                "structure": 3.0,
                "clarity": 3.0,
            },
            "strengths": ["Hệ thống đã phản hồi đúng trạng thái không đủ ngữ cảnh."],
            "weaknesses": ["[missing] Thiếu ngữ cảnh truy hồi để đánh giá nội dung section này."],
            "suggestions": [
                f"Bổ sung tài liệu nguồn liên quan đến section '{section_name}'.",
                "Tăng số chunk truy hồi trước khi tạo nội dung.",
            ],
        }

    if sentinel == "FAIL_COVERAGE":
        return {
            "is_fallback": True,
            "scores": {
                "accuracy": 7.0,
                "coverage": 3.0,
                "structure": 6.0,
                "clarity": 6.0,
            },
            "strengths": ["Nội dung có bám theo một phần context đã truy hồi."],
            "weaknesses": ["[missing] Thiếu các ý quan trọng trong context nên không đạt bao phủ."],
            "suggestions": [
                "Bổ sung các ý trọng tâm còn thiếu từ context vào section.",
                "Kiểm tra lại prompt để ép bao phủ đầy đủ theo section hiện tại.",
            ],
        }

    return {
        "is_fallback": True,
        "scores": {
            "accuracy": 7.0,
            "coverage": 7.0,
            "structure": 7.0,
            "clarity": 7.0,
        },
        "strengths": ["Nội dung có cấu trúc và bám context ở mức cơ bản."],
        "weaknesses": ["Cần thêm đánh giá chi tiết để phát hiện chính xác các điểm thiếu."],
        "suggestions": ["Sinh lại nội dung để lấy đánh giá tự động chi tiết hơn."],
    }


def _parse_eval_json(raw_eval: str) -> dict[str, Any]:
    cleaned = (raw_eval or "").strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    match = re.search(r"\{[\s\S]*\}", cleaned)
    if not match:
        return {}

    try:
        parsed = json.loads(match.group(0))
    except json.JSONDecodeError:
        return {}

    return parsed if isinstance(parsed, dict) else {}


def _evaluate_section_quality(section_name: str, context_text: str, generated_content: str) -> tuple[dict[str, Any], bool]:
    fallback = _fallback_section_evaluation(section_name=section_name, generated_content=generated_content)

    if not (context_text or "").strip():
        return fallback, False
    if not rag_pipeline.gemini_api_key:
        return fallback, False

    eval_prompt_template = (
        "Bạn là hệ thống đánh giá chất lượng nội dung bài giảng.\n"
        "INPUT:\n"
        "- Section: {section_name}\n"
        "- Context (retrieved chunks)\n"
        "- Generated content\n\n"
        "YÊU CẦU:\n"
        "1. Đánh giá theo các tiêu chí (0-10):\n"
        "- accuracy: có bám context không\n"
        "- coverage: có thiếu ý quan trọng không\n"
        "- structure: có đúng format section không\n"
        "- clarity: có rõ ràng, dễ học không\n"
        "2. Phát hiện lỗi:\n"
        "- Nếu có nội dung không có trong context, thêm mục weakness có tiền tố [hallucination].\n"
        "- Nếu thiếu nội dung quan trọng trong context, thêm mục weakness có tiền tố [missing].\n"
        "3. Chỉ trả về JSON hợp lệ, không markdown, không văn bản giải thích ngoài JSON.\n"
        "JSON schema:\n"
        "{\n"
        "  \"scores\": {\n"
        "    \"accuracy\": 0-10,\n"
        "    \"coverage\": 0-10,\n"
        "    \"structure\": 0-10,\n"
        "    \"clarity\": 0-10\n"
        "  },\n"
        "  \"strengths\": [\"...\"],\n"
        "  \"weaknesses\": [\"...\"],\n"
        "  \"suggestions\": [\"...\"]\n"
        "}\n\n"
        "Section: {section_name}\n"
        "Generated content:\n"
        "{generated_content}"
    )
    eval_prompt = (
        eval_prompt_template
        .replace("{section_name}", section_name)
        .replace("{generated_content}", generated_content)
    )

    try:
        raw_eval, gemini_real_call = rag_pipeline.generate_with_gemini_from_markdown(
            markdown=context_text,
            prompt=eval_prompt,
        )
    except Exception:
        return fallback, False

    parsed = _parse_eval_json(raw_eval)
    if not parsed:
        return fallback, gemini_real_call

    score_data = parsed.get("scores") if isinstance(parsed.get("scores"), dict) else {}
    fallback_scores = fallback.get("scores", {}) if isinstance(fallback.get("scores"), dict) else {}

    normalized = {
        "is_fallback": False,
        "scores": {
            "accuracy": _clamp_ten_score(score_data.get("accuracy"), float(fallback_scores.get("accuracy", 0.0))),
            "coverage": _clamp_ten_score(score_data.get("coverage"), float(fallback_scores.get("coverage", 0.0))),
            "structure": _clamp_ten_score(score_data.get("structure"), float(fallback_scores.get("structure", 0.0))),
            "clarity": _clamp_ten_score(score_data.get("clarity"), float(fallback_scores.get("clarity", 0.0))),
        },
        "strengths": _normalize_eval_list(parsed.get("strengths")) or fallback.get("strengths", []),
        "weaknesses": _normalize_eval_list(parsed.get("weaknesses")) or fallback.get("weaknesses", []),
        "suggestions": _normalize_eval_list(parsed.get("suggestions")) or fallback.get("suggestions", []),
    }

    return normalized, gemini_real_call


def _extract_code_candidates(text: str) -> list[str]:
    """Extract code-like statements (SQL, MongoDB, JSON, etc.) from text."""
    # Broader regex for technical keywords across various DBs and languages
    tech_re = re.compile(
        r"\b(select|insert|update|delete|create|alter|drop|truncate|join|where|"
        r"find|aggregate|db\.|collection|document|bson|json|key|value|schema)\b",
        flags=re.IGNORECASE,
    )
    candidates: list[str] = []

    def _looks_like_code_statement(candidate: str) -> bool:
        normalized = re.sub(r"\s+", " ", (candidate or "").strip().lower())
        if not normalized:
            return False

        # SQL checks
        if "select " in normalized and " from " in normalized: return True
        if "insert " in normalized and " into " in normalized: return True
        if "update " in normalized and " set " in normalized: return True
        if "delete " in normalized and " from " in normalized: return True
        
        # MongoDB/NoSQL checks
        if "db." in normalized and "(" in normalized and ")" in normalized: return True
        if ".find(" in normalized or ".aggregate(" in normalized: return True
        if normalized.startswith("{") and normalized.endswith("}") and ":" in normalized: return True
        
        # General technical structure
        if any(k in normalized for k in ("create table", "alter table", "drop table")): return True
        
        return False

    def _clean_code_candidate(raw: str) -> str:
        candidate = re.sub(r"\s+", " ", (raw or "").strip())
        if not candidate:
            return ""

        candidate = re.sub(r"^(?:vi\s*du|ví\s*dụ|example|sql|mongodb|code)\s*[:\-]\s*", "", candidate, flags=re.IGNORECASE)
        tick_match = re.search(r"`([^`]+)`", candidate)
        if tick_match:
            candidate = tick_match.group(1).strip()

        candidate = candidate.strip("` ")
        candidate = re.split(r"\s+-\s+", candidate, maxsplit=1)[0].strip()
        
        if not _looks_like_code_statement(candidate):
            return ""
        return candidate

    # Extract from markdown code blocks
    for block in re.findall(r"```(?:sql|json|javascript|js|mongodb)?\s*([\s\S]*?)```", text or "", flags=re.IGNORECASE):
        normalized = block.strip()
        if not normalized:
            continue
        # Split by potential statement separators
        statements = [item.strip() for item in re.split(r";\n|\n\n", normalized) if item.strip()]
        for stmt in statements:
            cleaned_stmt = _clean_code_candidate(stmt)
            if cleaned_stmt:
                candidates.append(cleaned_stmt)

    # Extract from single lines
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("```"):
            continue
        cleaned = re.sub(r"^[-*+]\s*", "", line)
        cleaned = _clean_code_candidate(cleaned)
        if cleaned:
            candidates.append(cleaned)

    return _dedupe_text_items(candidates)


_MAIN_CONTENT_LABEL_PREFIX_RE = re.compile(
    r"^(?:[-*+]\s*|\d+[.)]\s*)?(?:\*{1,2}\s*)?"
    r"(?:Định\s*nghĩa|Dinh\s*nghia|Giải\s*thích|Giai\s*thich|"
    r"Cách\s*dùng|Cach\s*dung|Khi\s*nào\s*dùng|Khi\s*nao\s*dung|"
    r"Ví\s*dụ\s*thực\s*tế|Vi\s*du\s*thuc\s*te|Ví\s*dụ|Vi\s*du|"
    r"Tình\s*huống|Tinh\s*huong|Yêu\s*cầu|Yeu\s*cau|"
    r"Thực\s*hiện|Thuc\s*hien|Nội\s*dung\s*thực\s*hiện|Noi\s*dung\s*thuc\s*hien|"
    r"Chú\s*thích|Chu\s*thich|Explanation)"
    r"(?:\s*\*{1,2})?(?:\s*[:\-]\s*|\s+)",
    flags=re.IGNORECASE,
)


def _strip_known_main_content_label_prefix(text: str) -> str:
    value = (text or "").strip()
    if not value:
        return ""

    for _ in range(2):
        cleaned = _MAIN_CONTENT_LABEL_PREFIX_RE.sub("", value, count=1).strip()
        if cleaned == value:
            break
        value = cleaned

    return value


def _extract_labeled_value(body_lines: list[str], labels: tuple[str, ...]) -> str:
    label_pattern = "|".join(re.escape(label) for label in labels)
    regex = re.compile(
        rf"^(?:(?:[-*+]\s*)|(?:\d+[.)]\s*))?(?:\*{{1,2}}\s*)?(?:{label_pattern})(?:\s*[:\-])?(?:\s*\*{{1,2}})?\s*[:\-]?\s*(.+)$",
        flags=re.IGNORECASE,
    )
    for line in body_lines:
        match = regex.match(line.strip())
        if match:
            cleaned = _strip_known_main_content_label_prefix(match.group(1))
            return re.sub(r"\s+", " ", cleaned).strip()
    return ""


def _extract_plain_sentences(text: str) -> list[str]:
    without_code = re.sub(r"```[\s\S]*?```", " ", text or "")
    normalized = without_code.replace("\r\n", "\n")

    cleaned_lines: list[str] = []
    for raw_line in normalized.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if re.match(r"^#{1,6}\s+", line):
            continue

        # Skip markdown table rows/separators in fallback prose extraction.
        if re.match(r"^\|.*\|$", line):
            continue
        if re.match(r"^\|?[\s:\-]+\|[\s|:\-]*$", line):
            continue

        line = re.sub(r"^(?:[-*+]\s+|\d+[.)]\s+)", "", line)
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"(?<![A-Za-z0-9])\*([^*]+)\*(?![A-Za-z0-9])", r"\1", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        line = _strip_known_main_content_label_prefix(line)
        line = re.sub(r"\s+", " ", line).strip()

        if line:
            cleaned_lines.append(line)

    cleaned = re.sub(r"\s+", " ", " ".join(cleaned_lines)).strip()
    if not cleaned:
        return []
    split_pattern = r"(?<!\b[Tt][Rr])(?<!\b[Tt][Rr][Gg])(?<!\b[Pp])(?<!\b[Pp][Pp])(?<!\b[Ee][Gg])(?<!\b[Ii][Ee])(?<=[.!?])\s+"
    return [item.strip(" .;:") for item in re.split(split_pattern, cleaned) if item.strip()]


def _parse_main_content_subsections(markdown_text: str) -> list[tuple[str, list[str], int]]:
    subsections: list[tuple[str, list[str], int]] = []
    current_title = ""
    current_level = 3
    current_lines: list[str] = []

    def _flush() -> None:
        nonlocal current_title, current_lines, current_level
        title = current_title.strip()
        if title:
            subsections.append((title, current_lines[:], current_level))
        current_title = ""
        current_level = 3
        current_lines = []

    for raw_line in (markdown_text or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            if current_title:
                current_lines.append("")
            continue

        heading3 = re.match(r"^###\s+(.+)$", stripped)
        if heading3:
            _flush()
            current_title = heading3.group(1).strip()
            current_level = 3
            continue

        heading4 = re.match(r"^####\s+(.+)$", stripped)
        if heading4:
            _flush()
            current_title = heading4.group(1).strip()
            current_level = 4
            continue

        heading2 = re.match(r"^##\s+(.+)$", stripped)
        if heading2:
            label = heading2.group(1).strip()
            normalized_label = _normalize_heading_label(label)
            if "noi dung chinh" in normalized_label:
                continue
            _flush()
            current_title = label
            current_level = 3
            continue

        if re.match(r"^#{1,6}\s+", stripped):
            continue

        if current_title:
            current_lines.append(stripped)

    _flush()

    if subsections:
        return subsections

    fallback_lines = [line.strip() for line in (markdown_text or "").splitlines() if line.strip() and not line.strip().startswith("#")]
    if fallback_lines:
        return [("Khái niệm chính", fallback_lines, 3)]
    return []


def _parse_body_blocks(body_text: str) -> list[str]:
    blocks = []
    lines = body_text.splitlines()
    in_code_block = False
    current_block_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith("```"):
            if in_code_block:
                current_block_lines.append(line)
                in_code_block = False
                blocks.append("\n".join(current_block_lines))
                current_block_lines = []
            else:
                if current_block_lines:
                    blocks.append("\n".join(current_block_lines))
                current_block_lines = [line]
                in_code_block = True
            continue
            
        if in_code_block:
            current_block_lines.append(line)
            continue
            
        if not stripped:
            if current_block_lines:
                blocks.append("\n".join(current_block_lines))
                current_block_lines = []
            continue
            
        current_block_lines.append(line)
        
    if current_block_lines:
        blocks.append("\n".join(current_block_lines))
        
    return [b.strip() for b in blocks if b.strip()]


def _clean_single_paragraph(p: str, banned_phrase_patterns, seen_sentence_norms) -> str:
    p = p.strip()
    if not p:
        return ""
    
    if p.startswith("|"):
        return p
        
    if p.startswith("```"):
        return p
        
    if p.startswith(">") or p.startswith("[!") or any(p.startswith(emoji) for emoji in ["💡", "📝", "🤔", "📖", "👉", "🔥", "⚠️"]):
        match = re.match(r"^([>\[!💡📝🤔📖👉🔥⚠️\s\-]+)(.*)$", p, flags=re.UNICODE)
        if match:
            prefix = match.group(1)
            content = match.group(2).strip()
            cleaned_content = _ensure_complete_sentence(content, fallback=content)
            return f"{prefix}{cleaned_content}"
        return p
        
    lines = p.splitlines()
    is_list = False
    for line in lines:
        stripped_line = line.strip()
        if stripped_line.startswith("-") or stripped_line.startswith("*") or stripped_line.startswith("+") or re.match(r"^\d+\.\s+", stripped_line):
            is_list = True
            break
            
    if is_list:
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                cleaned_lines.append("")
                continue
            list_match = re.match(r"^(\s*[-*+]\s*|\s*\d+\.\s*)(.*)$", line)
            if list_match:
                prefix = list_match.group(1)
                content = list_match.group(2).strip()
                
                bold_match = re.match(r"^(\*\*[^*]+\*\*\s*[:\-]?\s*)(.*)$", content)
                if bold_match:
                    bold_prefix = bold_match.group(1)
                    item_content = bold_match.group(2).strip()
                    cleaned_item = _ensure_complete_sentence(item_content, fallback=item_content)
                    cleaned_lines.append(f"{prefix}{bold_prefix}{cleaned_item}")
                else:
                    cleaned_item = _ensure_complete_sentence(content, fallback=content)
                    cleaned_lines.append(f"{prefix}{cleaned_item}")
            else:
                cleaned_lines.append(line)
        return "\n".join(cleaned_lines)
        
    split_pattern = r"(?<!\b[Tt][Rr])(?<!\b[Tt][Rr][Gg])(?<!\b[Pp])(?<!\b[Pp][Pp])(?<!\b[Ee][Gg])(?<!\b[Ii][Ee])(?<=[.!?])\s+"
    raw_sentences = [s.strip() for s in re.split(split_pattern, p) if s.strip()]
    
    cleaned_sentences = []
    for sentence in raw_sentences:
        cleaned = _strip_known_main_content_label_prefix(sentence)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" -;:")
        if not cleaned:
            continue
            
        normalized = _normalize_heading_label(cleaned)
        if not normalized:
            continue
            
        if any(re.search(pattern, cleaned, flags=re.IGNORECASE) for pattern in banned_phrase_patterns):
            continue
        if re.search(r"\b(rag|chunk|chunks)\b", normalized):
            continue
        if normalized.startswith("nguon ") or normalized.startswith("source "):
            continue
        if normalized in seen_sentence_norms:
            continue
            
        seen_sentence_norms.add(normalized)
        cleaned_sentences.append(_ensure_complete_sentence(cleaned, fallback=cleaned))
        
    return " ".join(cleaned_sentences)


def _enforce_main_content_hard_format(
    markdown_text: str,
    retrieved_chunks: list[dict[str, Any]] | None = None,
) -> str:
    sentinel = _extract_control_sentinel(markdown_text)
    if sentinel:
        return sentinel

    subsections = _parse_main_content_subsections(markdown_text)
    if not subsections:
        return "FAIL_COVERAGE"

    source_records = _build_quiz_source_records(retrieved_chunks or [])

    banned_phrase_patterns = (
        r"nội\s+dung\s+trọng\s+tâm\s+được\s+trích\s+xuất",
        r"trích\s+xuất\s+trực\s+tiếp\s+từ\s+tài\s+liệu",
        r"theo\s+ngữ\s+cảnh\s+tài\s+liệu",
    )

    output_lines: list[str] = []

    for idx, (raw_title, body_lines, _heading_level) in enumerate(subsections, start=1):
        title_seed = re.sub(r"^\d+(?:\.\d+)*(?:[\.)-])?\s*", "", (raw_title or "").strip())
        title_seed = title_seed.strip(" -")

        clean_title = title_seed.strip(" :")
        if not clean_title:
            clean_title = f"Mục {idx}"

        body_text = "\n".join(body_lines).strip()
        
        blocks = _parse_body_blocks(body_text)
        
        cleaned_paragraphs = []
        seen_sentence_norms: set[str] = set()
        for block in blocks:
            cleaned_p = _clean_single_paragraph(block, banned_phrase_patterns, seen_sentence_norms)
            if cleaned_p:
                cleaned_paragraphs.append(cleaned_p)
                
        if not cleaned_paragraphs:
            cleaned_paragraphs.append(
                f"{clean_title} là nội dung quan trọng cần nắm để hiểu đúng trọng tâm của bài học."
            )
            
        output_lines.append(f"### {clean_title}")
        for p in cleaned_paragraphs:
            output_lines.append(p)
            output_lines.append("")
            
        if source_records:
            # Skip duplicate footer if the text already contains a manual citation
            if not re.search(r"📚\s*Nguồn\s*:", body_text, flags=re.IGNORECASE):
                record_idx = (idx - 1) % len(source_records)
                citation = _format_source_citation_line(source_records[record_idx])
                output_lines.append(citation)
                output_lines.append("")

    return _compact_markdown_spacing("\n".join(output_lines).strip(), max_blank_lines=1)


def _enforce_hard_section_format(
    section_title: str,
    markdown_text: str,
    retrieved_chunks: list[dict[str, Any]] | None = None,
    selected_source_docs: list[dict[str, Any]] | None = None,
) -> str:
    raw = (markdown_text or "").strip()
    if not raw:
        return raw

    sentinel = _extract_control_sentinel(raw)
    if sentinel:
        return sentinel

    section_key = normalize_section_profile_key(section_title)
    if section_key == "main_content":
        return _enforce_main_content_hard_format(raw, retrieved_chunks=retrieved_chunks)
    if section_key == "objective":
        return _enforce_objective_hard_format(raw)
    if section_key == "summary":
        return _enforce_summary_hard_format(raw)
    if section_key == "quiz":
        return _enforce_quiz_hard_format(
            raw,
            retrieved_chunks=retrieved_chunks,
            selected_source_docs=selected_source_docs,
        )

    # For all other sections, inject a global citation footer if not already present.
    # Main Content and Quiz have their own per-item injection logic.
    if section_key not in {"main_content", "quiz"} and retrieved_chunks:
        # Check if it already has a citation line to avoid duplicates
        if not re.search(r"📚\s*Nguồn\s*:", raw, flags=re.IGNORECASE):
            source_records = _build_quiz_source_records(retrieved_chunks)
            if source_records:
                citation = _format_source_citation_line(source_records[0])
                raw = f"{raw.strip()}\n\n{citation}"

    return raw


def _parse_outline_to_sections(outline_markdown: str) -> list[dict[str, Any]]:
    def _clean_outline_title(text: str) -> str:
        cleaned = (text or "").strip()
        # Only strip leading markdown list markers like -, *, +, but keep numbers and "Chương/Mục" prefixes
        cleaned = re.sub(r"^[-*+•\s]+", "", cleaned)
        return cleaned.strip()

    sections: list[dict[str, Any]] = []
    raw_lines = (outline_markdown or "").splitlines()

    for raw in raw_lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = max(1, len(heading.group(1))) # keep absolute heading levels: # is level 1, ## is level 2, ### is level 3
            title = _clean_outline_title(heading.group(2))
            if title:
                sections.append({"title": title, "level": level})
            continue

        numbered = re.match(r"^(\d+(?:\.\d+)*)[\.)]?\s+(.+)$", stripped)
        if numbered:
            order_token = numbered.group(1)
            level = max(1, order_token.count(".") + 1)
            title = _clean_outline_title(stripped) # Keep the full numbered prefix line
            sections.append({"title": title, "level": level})
            continue

        bullet = re.match(r"^[-*+•]\s+(.+)$", stripped)
        if bullet:
            indent_spaces = len(line) - len(line.lstrip(" "))
            level = max(2, (indent_spaces // 2) + 2) # Indented bullets are sub-sections
            title = _clean_outline_title(bullet.group(1))
            if title:
                # If bullet starts with a hierarchy number, infer level from it
                sub_num = re.match(r"^(\d+(?:\.\d+)*)", title)
                if sub_num:
                    level = max(1, sub_num.group(1).count(".") + 1)
                sections.append({"title": title, "level": level})
            continue

        # Fallback for plain lines
        level = 1
        sub_num = re.match(r"^(\d+(?:\.\d+)*)", stripped)
        if sub_num:
            level = max(1, sub_num.group(1).count(".") + 1)
        sections.append({"title": stripped, "level": level})

    normalized: list[dict[str, Any]] = []
    for idx, item in enumerate(sections):
        level = max(1, min(6, int(item.get("level") or 1)))
        title = str(item.get("title") or "").strip()
        if not title:
            continue
        normalized.append({"title": title, "level": level, "order_index": idx})
    return normalized


def _normalize_teaching_outline_sections(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    import uuid
    normalized: list[dict[str, Any]] = []
    for order_idx, s in enumerate(sections):
        normalized.append(
            {
                "section_id": str(uuid.uuid4()),
                "title": str(s.get("title", "")).strip(),
                "level": int(s.get("level", 1)),
                "order_index": order_idx,
            }
        )

    return normalized


def _metadata_source_label(metadata: dict[str, Any]) -> str:
    if not isinstance(metadata, dict):
        return ""
    return str(
        metadata.get("file_name")
        or metadata.get("source")
        or metadata.get("source_file")
        or metadata.get("filename")
        or metadata.get("title")
        or ""
    ).strip()


def _metadata_page_range(metadata: dict[str, Any]) -> tuple[int | None, int | None]:
    if not isinstance(metadata, dict):
        return (None, None)

    start = _coerce_page_number(metadata.get("start_page", metadata.get("page_start")))
    end = _coerce_page_number(metadata.get("end_page", metadata.get("page_end")))

    if start is None and end is None:
        fallback = _coerce_page_number(metadata.get("page_number", metadata.get("page")))
        return (fallback, fallback)

    if start is None:
        start = end
    if end is None:
        end = start

    if start is not None and end is not None and end < start:
        start, end = end, start

    return (start, end)


def _format_page_range_label(start_page: int | None, end_page: int | None) -> str:
    if start_page is None and end_page is None:
        return ""
    if start_page is not None and end_page is not None and start_page != end_page:
        return f"tr. {start_page}-{end_page}"
    page = start_page if start_page is not None else end_page
    return f"tr. {page}" if page is not None else ""


def _strip_metadata_heading_prefix(text: str, metadata: dict[str, Any]) -> str:
    """
    Remove leading heading trail that duplicates metadata breadcrumb/chapter.

    This keeps retrieved chunk previews compact and avoids repeated patterns like:
    "chapter > section > subsection" + "# chapter\n## section\n### subsection".
    """
    raw = (text or "").strip()
    if not raw:
        return ""

    lines = raw.splitlines()
    cursor = 0

    def _seek_next_non_empty(index: int) -> int:
        while index < len(lines) and not lines[index].strip():
            index += 1
        return index

    def _norm(value: str) -> str:
        return _normalize_heading_label(value)

    cursor = _seek_next_non_empty(cursor)
    removed_any = False

    prefix_candidates = [
        str(metadata.get("heading_path") or "").strip(),
        str(metadata.get("breadcrumb") or "").strip(),
    ]
    prefix_norms = {norm for norm in (_norm(item) for item in prefix_candidates) if norm}
    if cursor < len(lines) and prefix_norms:
        first_norm = _norm(lines[cursor].strip())
        if first_norm in prefix_norms:
            removed_any = True
            cursor = _seek_next_non_empty(cursor + 1)

    expected_headings = [
        str(metadata.get("chapter") or metadata.get("h1") or "").strip(),
        str(metadata.get("section") or metadata.get("h2") or "").strip(),
        str(metadata.get("subsection") or metadata.get("h3") or "").strip(),
    ]
    expected_norms = [_norm(item) for item in expected_headings if _norm(item)]

    heading_index = 0
    while cursor < len(lines):
        stripped = lines[cursor].strip()
        if not stripped:
            cursor += 1
            continue

        heading_match = re.match(r"^#{1,6}\s+(.+)$", stripped)
        if not heading_match:
            break
        if heading_index >= len(expected_norms):
            break

        heading_norm = _norm(heading_match.group(1))
        if heading_norm != expected_norms[heading_index]:
            break

        removed_any = True
        heading_index += 1
        cursor += 1

    if not removed_any:
        return raw

    trimmed = "\n".join(lines[cursor:]).strip()
    return trimmed or raw


def _build_context_text(results: list[dict[str, Any]]) -> str:
    blocks: list[str] = []
    for idx, item in enumerate(results, 1):
        metadata = item.get("metadata", {}) if isinstance(item.get("metadata"), dict) else {}
        source = _metadata_source_label(metadata)
        title = str(metadata.get("title", ""))
        breadcrumb = str(metadata.get("breadcrumb") or metadata.get("heading_path") or "").strip()
        start_page, end_page = _metadata_page_range(metadata)
        page_label = _format_page_range_label(start_page, end_page)
        text = _strip_metadata_heading_prefix(
            str(item.get("text", "")),
            metadata,
        )
        if not text:
            continue
        header = f"[CONTEXT {idx}] source={source}; title={title}"
        if breadcrumb:
            header += f"; breadcrumb={breadcrumb}"
        if page_label:
            header += f"; page={page_label}"
        blocks.append(f"{header}\n{text}")
    return "\n\n".join(blocks).strip()


def _sanitize_heading_candidate(raw_text: str) -> str:
    candidate = (raw_text or "").strip()
    if not candidate:
        return ""

    candidate = re.sub(r"^#{1,6}\s+", "", candidate)
    candidate = re.sub(r"^[\-*+]\s+", "", candidate)
    candidate = re.sub(r"^\d+(?:\.\d+)*(?:[\.)-])?\s+", "", candidate)
    candidate = re.sub(r"\s+", " ", candidate).strip(" -:;,.")
    if len(candidate) < 3 or len(candidate) > 120:
        return ""

    normalized = _normalize_heading_label(candidate)
    if not normalized:
        return ""

    blocked_labels = {
        "noi dung chinh",
        "main content",
        "giai thich",
        "khi nao dung",
        "vi du",
    }
    if normalized in blocked_labels:
        return ""

    # Keep short one-word tokens only when they look like acronyms (DDL, DML, TCL...).
    if len(normalized.split()) == 1 and len(normalized) < 4 and not re.fullmatch(r"[A-Z0-9_]{2,8}", candidate):
        return ""

    return candidate


def _extract_heading_hints_from_retrieved(
    retrieved: list[dict[str, Any]],
    max_items: int = 20,
) -> list[str]:
    metadata_keys = (
        "heading",
        "heading_1",
        "heading_2",
        "heading_3",
        "heading_4",
        "h1",
        "h2",
        "h3",
        "h4",
        "section",
        "section_title",
        "chapter",
        "chapter_title",
    )
    heading_hints: list[str] = []
    seen_labels: set[str] = set()

    def _append_heading(raw_value: str) -> None:
        candidate = _sanitize_heading_candidate(raw_value)
        if not candidate:
            return

        normalized = _normalize_heading_label(candidate)
        if not normalized or normalized in seen_labels:
            return

        seen_labels.add(normalized)
        heading_hints.append(candidate)

    for item in retrieved or []:
        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
        for key in metadata_keys:
            value = str(metadata.get(key) or "").strip()
            if value:
                _append_heading(value)
            if len(heading_hints) >= max_items:
                return heading_hints[:max_items]

        text = str(item.get("text") or "")
        for raw_line in text.splitlines()[:24]:
            stripped = raw_line.strip()
            if not stripped:
                continue
            if re.match(r"^#{1,6}\s+.+$", stripped) or re.match(r"^\d+(?:\.\d+)*(?:[\.)-])\s+.+$", stripped):
                _append_heading(stripped)
            if len(heading_hints) >= max_items:
                return heading_hints[:max_items]

    return heading_hints[:max_items]


def _build_source_info_from_retrieved(
    retrieved: list[dict[str, Any]],
    max_items: int = 4,
) -> str:
    if not retrieved:
        return "Tài liệu đã chọn"

    source_items: list[str] = []
    seen: set[str] = set()

    for item in retrieved:
        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
        title = str(metadata.get("title") or "").strip()
        source = _metadata_source_label(metadata)
        breadcrumb = str(metadata.get("breadcrumb") or metadata.get("heading_path") or "").strip()
        display = source or title
        if not display:
            continue

        start_page, end_page = _metadata_page_range(metadata)
        page_label = _format_page_range_label(start_page, end_page)
        if page_label:
            display = f"{display} ({page_label})"
        if breadcrumb:
            display = f"{display} [{breadcrumb}]"

        norm = _normalize_heading_label(display)
        if not norm or norm in seen:
            continue
        seen.add(norm)
        source_items.append(display)
        if len(source_items) >= max_items:
            break

    return "; ".join(source_items) if source_items else "Tài liệu đã chọn"


def _append_source_tracing_footer(markdown_text: str, source_info: str) -> str:
    content = (markdown_text or "").strip()
    if not content:
        return content

    clean_source = (source_info or "").strip() or "Tài liệu đã chọn"
    footer = f"--- *Nguồn: {clean_source}*"

    if re.search(r"(?im)^---\s*\*?\s*(nguồn|nguon)\s*:", content):
        content = re.sub(
            r"(?im)^---\s*\*?\s*(nguồn|nguon)\s*:[^\n]*\*?\s*$",
            footer,
            content,
        )
        return content.strip()

    return f"{content}\n\n{footer}".strip()


def _compact_markdown_spacing(markdown_text: str, max_blank_lines: int = 1) -> str:
    text = (markdown_text or "").replace("\r\n", "\n")
    if not text.strip():
        return ""

    compact_lines: list[str] = []
    blank_count = 0
    in_code_block = False

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            blank_count = 0
            compact_lines.append(line.rstrip())
            continue

        if in_code_block:
            compact_lines.append(line.rstrip())
            continue

        if not stripped:
            blank_count += 1
            if blank_count <= max_blank_lines:
                compact_lines.append("")
            continue

        blank_count = 0
        compact_lines.append(line.rstrip())

    return "\n".join(compact_lines).strip()


def _build_main_content_heading_anchor_prompt(heading_hints: list[str], source_info: str) -> str:
    base = (
        "MAIN CONTENT HEADING LOCK (EDUCATOR REFINER MODE)\n"
        "- Anchor strictly to source heading/sub-heading order from context.\n"
        "- SCANNING RULE: bao phủ mọi sub-heading H2/H3/H4 nhận diện được trong context.\n"
        "- NO OMISSION: không gộp hoặc bỏ sót mục con.\n"
        "- TEACHING STYLE: viết như giảng viên, tiếng Việt tự nhiên, giải thích rõ là gì/khi nào dùng/hoạt động ra sao.\n"
        "- STRUCTURE: tổ chức thành các đoạn ngắn theo chủ đề, chuyển ý mượt, không sao chép máy móc.\n"
        "- CONTENT QUALITY: loại bỏ câu lặp/noise và cụm từ máy móc.\n"
        "- TECHNICAL ACCURACY: nếu có code/công thức/lệnh, giữ cú pháp đúng và ví dụ ngắn gọn, thực tế.\n"
        "- FORBIDDEN: không nhắc RAG/chunk/source hoặc câu truy vết.\n"
        f"- INTERNAL CONTEXT REFERENCE: {source_info or 'Tài liệu đã chọn'} (chỉ dùng để bám ngữ cảnh, không in ra output).\n"
        "- Mọi câu phải hoàn chỉnh, không cắt cụt.\n"
        "- Tránh dãn dòng/dãn đoạn lớn; tối đa 1 dòng trống giữa các khối nội dung.\n"
    )

    if not heading_hints:
        return (
            f"{base}"
            "- Nếu không nhận diện được heading đủ tin cậy, return exactly: FAIL_COVERAGE."
        )

    hint_lines = "\n".join(f"- {heading}" for heading in heading_hints)
    return (
        f"{base}"
        "Heading/sub-heading signals extracted from context:\n"
        f"{hint_lines}\n"
        "- Dùng danh sách này làm checklist coverage trước khi xuất kết quả cuối."
    )


def _normalize_context_text_for_llm(text: str) -> str:
    raw = (text or "").replace("\r\n", "\n")
    if not raw.strip():
        return ""

    # Merge hyphen-broken words split by line breaks.
    raw = re.sub(r"(?<=\w)-\n(?=\w)", "", raw)

    normalized_lines: list[str] = []
    buffer: list[str] = []

    def _flush_buffer() -> None:
        if not buffer:
            return
        merged = re.sub(r"\s+", " ", " ".join(buffer)).strip()
        if merged:
            normalized_lines.append(merged)
        buffer.clear()

    for raw_line in raw.splitlines():
        line = raw_line.strip()
        if not line:
            _flush_buffer()
            continue

        is_anchor_line = bool(
            line.startswith("[CONTEXT ")
            or re.match(r"^#{1,6}\s+", line)
            or re.match(r"^[-*+]\s+", line)
            or re.match(r"^\d+[\.)]\s+", line)
            or line.startswith("|")
        )
        if is_anchor_line:
            _flush_buffer()
            normalized_lines.append(line)
            continue

        buffer.append(line)

    _flush_buffer()
    return "\n".join(normalized_lines).strip()


def _retrieval_item_score(item: dict[str, Any]) -> float:
    return float(item.get("rerank_score", item.get("_score", item.get("hybrid_score", 0.0))))


def _coerce_page_number(value: Any) -> int | None:
    if value is None:
        return None
    if isinstance(value, int):
        return value if value > 0 else None

    text = str(value).strip()
    if not text:
        return None
    match = re.search(r"\d+", text)
    if not match:
        return None
    try:
        page = int(match.group(0))
        return page if page > 0 else None
    except ValueError:
        return None


def _optional_text(value: Any) -> str | None:
    text = str(value or "").strip()
    return text if text else None


_HEADING_TECH_NOISE_RE = re.compile(
    r"\b(SELECT|FROM|WHERE|INSERT|UPDATE|DELETE|JOIN|FIND|AGGREGATE|DB\.|COLLECTION|DOCUMENT|BSON|JSON)\b",
    re.IGNORECASE,
)


def _clean_structure_label(value: Any, level: str) -> str | None:
    text = str(value or "").strip()
    if not text:
        return None

    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"^[+\-*•]+\s*", "", text).strip()
    text = text.splitlines()[0].strip()
    if not text:
        return None

    noise_match = _HEADING_TECH_NOISE_RE.search(text)
    if noise_match and noise_match.start() > 0:
        text = text[: noise_match.start()].rstrip(" -,:;")
    if not text:
        return None

    # Reject overly long sentence-like values that are likely chunk content leaks.
    if len(text) > 120:
        return None

    if level == "chapter":
        if re.match(r"^\d+$", text):
            return text
        if re.match(r"^(chương|chuong|chapter)\s+\d+", text, flags=re.IGNORECASE):
            return text[:90].strip()
        if re.match(r"^\d+\s*[:.)-]", text):
            return text[:90].strip()
        if re.match(r"^(neu|nếu|khi|ta|dung|dùng|su\s+dung|sử\s+dụng)\b", text, flags=re.IGNORECASE):
            return None
        if re.match(r"^[^\w\d]+", text):
            return None
        if len(text.split()) > 8:
            return None
        return text

    if level in {"section", "subsection"}:
        numbered = re.match(r"^(\d+(?:\.\d+){1,4})\s*[:\-.)]?\s*(.*)$", text)
        if numbered:
            number = numbered.group(1)
            title = re.sub(r"\s+", " ", numbered.group(2).strip())
            if title:
                title = re.split(r"\s{2,}|[;|]", title, maxsplit=1)[0].strip()
                if len(title) > 80:
                    title = re.split(r"[.!?,]", title, maxsplit=1)[0].strip()
            if not title:
                return number
            return f"{number}. {title}"[:90].strip()

        if re.match(r"^[+\-*•]", text):
            return None
        if len(text) > 90:
            return None
        return text

    return text


def _infer_structure_from_chunk_text(text: str) -> tuple[str | None, str | None, str | None]:
    raw = (text or "").strip()
    if not raw:
        return None, None, None

    chapter: str | None = None
    section: str | None = None
    subsection: str | None = None

    for line in raw.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        h3 = re.match(r"^###\s+(.+)$", stripped)
        if h3 and not subsection:
            subsection = _clean_structure_label(h3.group(1), "subsection")
            continue
        h2 = re.match(r"^##\s+(.+)$", stripped)
        if h2 and not section:
            section = _clean_structure_label(h2.group(1), "section")
            continue
        h1 = re.match(r"^#\s+(.+)$", stripped)
        if h1 and not chapter:
            chapter = _clean_structure_label(h1.group(1), "chapter")
            continue

    if not chapter:
        chapter_match = re.search(r"(?im)^\s*((?:chương|chuong|chapter)\s+\d+[^\n]*)$", raw)
        if chapter_match:
            chapter = _clean_structure_label(chapter_match.group(1), "chapter")

    if not subsection:
        subsection_match = re.search(r"(?m)^\s*(\d+\.\d+\.\d+(?:\.\d+)?)\s*[:\-.)]?\s*(.+)$", raw)
        if subsection_match:
            subsection = _clean_structure_label(
                f"{subsection_match.group(1)}. {subsection_match.group(2)}",
                "subsection",
            )

    if not section:
        section_match = re.search(r"(?m)^\s*(\d+\.\d+(?:\.\d+)?)\s*[:\-.)]?\s*(.+)$", raw)
        if section_match:
            section = _clean_structure_label(
                f"{section_match.group(1)}. {section_match.group(2)}",
                "section",
            )

    return chapter, section, subsection


def _build_breadcrumb_from_structure(
    chapter: str | None,
    section: str | None,
    subsection: str | None,
) -> str | None:
    parts = [part for part in [chapter, section, subsection] if part]
    if not parts:
        return None
    return " > ".join(parts)


def _serialize_retrieved_chunk(item: dict[str, Any]) -> dict[str, Any]:
    metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
    title = str(metadata.get("title") or "").strip()
    source = _metadata_source_label(metadata)
    start_page, end_page = _metadata_page_range(metadata)
    page_number = start_page

    raw_text = str(item.get("text") or "")
    inferred_chapter, inferred_section, inferred_subsection = _infer_structure_from_chunk_text(raw_text)

    chapter = _clean_structure_label(
        metadata.get("chapter") or metadata.get("h1") or inferred_chapter,
        "chapter",
    )
    section = _clean_structure_label(
        metadata.get("section") or metadata.get("h2") or inferred_section,
        "section",
    )
    subsection = _clean_structure_label(
        metadata.get("subsection") or metadata.get("h3") or inferred_subsection,
        "subsection",
    )
    chapter_title = _clean_structure_label(metadata.get("chapter_title") or metadata.get("h1"), "chapter")
    section_title = _clean_structure_label(metadata.get("section_title") or metadata.get("h2"), "section")
    subsection_title = _clean_structure_label(
        metadata.get("subsection_title") or metadata.get("h3"),
        "subsection",
    )
    breadcrumb = _optional_text(metadata.get("breadcrumb") or metadata.get("heading_path"))
    if not breadcrumb:
        breadcrumb = _build_breadcrumb_from_structure(chapter, section, subsection)

    metadata_for_strip = dict(metadata)
    if chapter and not metadata_for_strip.get("chapter"):
        metadata_for_strip["chapter"] = chapter
    if section and not metadata_for_strip.get("section"):
        metadata_for_strip["section"] = section
    if subsection and not metadata_for_strip.get("subsection"):
        metadata_for_strip["subsection"] = subsection
    if breadcrumb and not metadata_for_strip.get("breadcrumb"):
        metadata_for_strip["breadcrumb"] = breadcrumb

    normalized_metadata = {
        "doc_id": str(metadata.get("doc_id") or metadata.get("document_id") or metadata.get("source") or "").strip(),
        "file_name": _optional_text(metadata.get("file_name") or source),
        "chapter": chapter,
        "section": section,
        "subsection": subsection,
        "chapter_title": chapter_title,
        "section_title": section_title,
        "subsection_title": subsection_title,
        "breadcrumb": breadcrumb,
        "start_page": start_page,
        "end_page": end_page,
    }

    cleaned_text = _strip_metadata_heading_prefix(
        raw_text,
        metadata_for_strip,
    )

    return {
        "id": str(item.get("id") or ""),
        "text": cleaned_text,
        "score": float(item.get("rerank_score", item.get("hybrid_score", item.get("_score", 0.0)))),
        "source": source,
        "title": title,
        "page_number": page_number,
        "start_page": start_page,
        "end_page": end_page,
        "metadata": normalized_metadata,
    }


def _extract_summary_group_key(metadata: dict[str, Any]) -> str:
    if not isinstance(metadata, dict):
        return ""

    priority_keys = (
        "breadcrumb",
        "chapter_id",
        "chapter",
        "chapter_title",
        "h1",
        "heading_1",
        "section_id",
        "section",
        "h2",
        "doc_id",
        "document_id",
        "file_name",
        "title",
    )
    for key in priority_keys:
        value = str(metadata.get(key) or "").strip()
        if value:
            return _normalize_heading_label(value)

    source = _metadata_source_label(metadata)
    if source:
        return f"source::{_normalize_heading_label(source)}"
    return ""


def _prioritize_summary_group_coverage(chunks: list[dict[str, Any]], max_chunks: int) -> list[dict[str, Any]]:
    if not chunks:
        return []

    grouped: dict[str, list[dict[str, Any]]] = {}
    ungrouped: list[dict[str, Any]] = []

    for item in chunks:
        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
        group_key = _extract_summary_group_key(metadata)
        if group_key:
            grouped.setdefault(group_key, []).append(item)
        else:
            ungrouped.append(item)

    for values in grouped.values():
        values.sort(key=_retrieval_item_score, reverse=True)
    ungrouped.sort(key=_retrieval_item_score, reverse=True)

    selected: list[dict[str, Any]] = []
    seen_ids: set[str] = set()

    if grouped:
        group_keys = sorted(grouped.keys())
        while len(selected) < max_chunks:
            progressed = False
            for key in group_keys:
                bucket = grouped.get(key, [])
                if not bucket:
                    continue
                candidate = bucket.pop(0)
                chunk_id = str(candidate.get("id") or "")
                if not chunk_id or chunk_id in seen_ids:
                    continue
                seen_ids.add(chunk_id)
                selected.append(candidate)
                progressed = True
                if len(selected) >= max_chunks:
                    break
            if not progressed:
                break

    leftovers: list[dict[str, Any]] = []
    for values in grouped.values():
        leftovers.extend(values)
    leftovers.extend(ungrouped)
    leftovers.sort(key=_retrieval_item_score, reverse=True)

    for candidate in leftovers:
        if len(selected) >= max_chunks:
            break
        chunk_id = str(candidate.get("id") or "")
        if not chunk_id or chunk_id in seen_ids:
            continue
        seen_ids.add(chunk_id)
        selected.append(candidate)

    return selected[:max_chunks]


def _interleave_by_source(
    chunks_by_source: dict[str, list[dict[str, Any]]],
    max_chunks: int,
    min_per_source: int = 1,
) -> list[dict[str, Any]]:
    """Round-robin interleave chunks across sources to prevent document starvation.

    Strategy:
    1.  Each source gets a guaranteed minimum slot (``min_per_source``).
    2.  After filling minimums, remaining slots are filled round-robin by score.
    3.  If a source runs out of chunks early its turns are skipped gracefully.

    Args:
        chunks_by_source: Mapping of source_key -> list of chunks sorted by score DESC.
        max_chunks: Total number of chunks to return.
        min_per_source: Minimum chunks guaranteed per source (default 1).

    Returns:
        Interleaved list of at most ``max_chunks`` chunks.
    """
    if not chunks_by_source or max_chunks <= 0:
        return []

    # Work on mutable copies so we can pop without mutating the originals.
    buckets: dict[str, list[dict[str, Any]]] = {
        src: list(items) for src, items in chunks_by_source.items() if items
    }
    source_keys = list(buckets.keys())
    if not source_keys:
        return []

    selected: list[dict[str, Any]] = []
    seen_ids: set[str] = set()

    def _pick(src: str) -> bool:
        """Pop the top unseen chunk from bucket[src]. Returns True if successful."""
        bucket = buckets.get(src, [])
        while bucket:
            candidate = bucket.pop(0)
            cid = str(candidate.get("id") or "").strip()
            if cid and cid not in seen_ids:
                seen_ids.add(cid)
                selected.append(candidate)
                return True
        return False

    # Phase 1: guarantee minimum slots per source.
    for _ in range(min_per_source):
        for src in source_keys:
            if len(selected) >= max_chunks:
                break
            _pick(src)
        if len(selected) >= max_chunks:
            break

    # Phase 2: fill remaining slots round-robin until all buckets exhausted.
    while len(selected) < max_chunks:
        any_picked = False
        for src in source_keys:
            if len(selected) >= max_chunks:
                break
            if _pick(src):
                any_picked = True
        if not any_picked:
            break  # All sources exhausted.

    return selected[:max_chunks]


def _retrieve_context_with_retry(
    query: str,
    selected_source_docs: list[dict[str, Any]],
    min_total_chars: int = 1200,
    max_chunks: int = 5,
    top_k_levels: list[int] | tuple[int, ...] | None = None,
) -> list[dict[str, Any]]:
    """Retrieve context chunks with adaptive top-k and multi-doc fairness.

    When multiple source docs are selected, chunks are interleaved round-robin
    so every document has representation in the final context window — preventing
    a single high-scoring document from monopolising all slots.
    """
    raw_levels = top_k_levels or [3, 4, 5]
    search_levels = sorted({max(1, int(level)) for level in raw_levels})

    # --- Per-source tracking (key = source_tag or doc id) ---
    # chunks_by_source[src_key][chunk_id] = best scored item
    chunks_by_source: dict[str, dict[str, dict[str, Any]]] = {}
    for src_doc in selected_source_docs:
        src_key = str(src_doc.get("source_tag") or src_doc.get("id") or "")
        chunks_by_source.setdefault(src_key, {})

    multi_doc = len(selected_source_docs) > 1

    for top_k in search_levels:
        for src_doc in selected_source_docs:
            src_key = str(src_doc.get("source_tag") or src_doc.get("id") or "")
            retrieved, _ = rag_pipeline.search_knowledge_base(
                query=query,
                collection_name=str(src_doc.get("collection_name") or "") or None,
                top_k=top_k,
                vector_weight=0.65,
                keyword_weight=0.35,
                source_filter=str(src_doc.get("source_tag") or "") or None,
                use_rerank=True,
            )
            for item in retrieved:
                chunk_id = str(item.get("id") or "")
                if not chunk_id:
                    continue
                score = float(item.get("rerank_score", item.get("hybrid_score", 0.0)))
                bucket = chunks_by_source[src_key]
                prev = bucket.get(chunk_id)
                if prev is None or float(prev.get("_score", -1.0)) < score:
                    bucket[chunk_id] = {**item, "_score": score}

        # --- Source-Parity Enforcement: Ensure every document has at least one chunk ---
        if multi_doc:
            starved_docs = [d for d in selected_source_docs if not chunks_by_source[str(d.get("source_tag") or d.get("id") or "")]]
            if starved_docs and top_k == search_levels[-1]:
                # Final attempt: search specifically for starved docs with more breadth
                for src_doc in starved_docs:
                    src_key = str(src_doc.get("source_tag") or src_doc.get("id") or "")
                    retrieved, _ = rag_pipeline.search_knowledge_base(
                        query=query,
                        collection_name=str(src_doc.get("collection_name") or "") or None,
                        top_k=max(top_k, 8), # Broaden search
                        vector_weight=0.4,   # More keyword focus for starvation recovery
                        keyword_weight=0.6,
                        source_filter=str(src_doc.get("source_tag") or "") or None,
                        use_rerank=False,    # Skip rerank for faster recovery
                    )
                    for item in retrieved:
                        chunk_id = str(item.get("id") or "")
                        if not chunk_id: continue
                        score = float(item.get("hybrid_score", 0.0))
                        chunks_by_source[src_key][chunk_id] = {**item, "_score": score}

        # Build sorted lists per source, then interleave or sort globally.
        if multi_doc:
            sorted_by_src: dict[str, list[dict[str, Any]]] = {}
            for src_key, bucket in chunks_by_source.items():
                sorted_items = sorted(
                    bucket.values(),
                    key=lambda x: float(x.get("_score", 0.0)),
                    reverse=True,
                )
                if sorted_items:
                    sorted_by_src[src_key] = sorted_items

            interleaved = _interleave_by_source(sorted_by_src, max_chunks=max_chunks)
            total_chars = sum(len(str(c.get("text", ""))) for c in interleaved)
            # If we have interleaved results, check if they are enough
            if len(interleaved) >= min(2, len(selected_source_docs)) and total_chars >= min_total_chars:
                return interleaved
        else:
            # Single-doc path — original global sort (no change in behaviour).
            all_chunks = list(next(iter(chunks_by_source.values()), {}).values())
            all_chunks.sort(key=lambda x: float(x.get("_score", 0.0)), reverse=True)
            merged = all_chunks[:max_chunks]
            total_chars = sum(len(str(c.get("text", ""))) for c in merged)
            if len(merged) >= 2 and total_chars >= min_total_chars:
                return merged

    # Final fallback — return best available result using same strategy.
    if multi_doc:
        sorted_by_src = {}
        for src_key, bucket in chunks_by_source.items():
            sorted_items = sorted(
                bucket.values(),
                key=lambda x: float(x.get("_score", 0.0)),
                reverse=True,
            )
            if sorted_items:
                sorted_by_src[src_key] = sorted_items
        return _interleave_by_source(sorted_by_src, max_chunks=max_chunks)
    else:
        all_chunks = list(next(iter(chunks_by_source.values()), {}).values())
        all_chunks.sort(key=lambda x: float(x.get("_score", 0.0)), reverse=True)
        return all_chunks[:max_chunks]



def _score_full_section_candidate(query: str, text: str) -> float:
    tokens = {token for token in re.split(r"\W+", _normalize_heading_label(query)) if len(token) >= 2}
    if not tokens:
        return 0.0
    normalized_text = _normalize_heading_label(text)
    matched = sum(1 for token in tokens if token in normalized_text)
    return matched / max(1, len(tokens))


def _retrieve_full_section_context_with_rerank(
    query: str,
    selected_source_docs: list[dict[str, Any]],
    max_chunks: int = 120,
) -> list[dict[str, Any]]:
    """Fetch all chunks for selected sources, rerank, then interleave across docs."""
    # candidate_by_src_id[src_key][chunk_id] = chunk — keeps best score per chunk per source.
    candidate_by_src: dict[str, dict[str, dict[str, Any]]] = {}
    per_source_limit = max(50, int(max_chunks or 120))
    multi_doc = len(selected_source_docs) > 1

    for src_doc in selected_source_docs:
        source_tag = str(src_doc.get("source_tag") or "")
        src_key = source_tag or str(src_doc.get("id") or "")
        collection_name = str(src_doc.get("collection_name") or "") or None
        raw_chunks = rag_pipeline.get_chunks_by_source(
            source_tag=source_tag,
            collection_name=collection_name,
            limit=per_source_limit,
        )

        src_bucket: dict[str, dict[str, Any]] = {}
        for idx, chunk in enumerate(raw_chunks):
            text = str(chunk.get("text") or "").strip()
            if not text:
                continue

            chunk_id = str(chunk.get("chunk_id") or chunk.get("id") or "").strip()
            if not chunk_id:
                chunk_id = f"{source_tag or 'source'}:{idx + 1}"

            metadata = chunk.get("metadata") if isinstance(chunk.get("metadata"), dict) else {}
            metadata = dict(metadata)
            if source_tag and not metadata.get("source"):
                metadata["source"] = source_tag
            if source_tag and not metadata.get("doc_id"):
                metadata["doc_id"] = source_tag
            if not metadata.get("file_name"):
                metadata["file_name"] = str(src_doc.get("original_filename") or source_tag)
            if not metadata.get("source_file"):
                metadata["source_file"] = metadata.get("file_name")
            if not metadata.get("filename"):
                metadata["filename"] = metadata.get("file_name")
            if not metadata.get("title"):
                metadata["title"] = str(src_doc.get("original_filename") or "")
            page_fallback = _coerce_page_number(metadata.get("page_number", metadata.get("page")))
            if metadata.get("start_page") in (None, ""):
                metadata["start_page"] = page_fallback if page_fallback is not None else -1
            if metadata.get("end_page") in (None, ""):
                metadata["end_page"] = metadata.get("start_page")
            if not metadata.get("breadcrumb"):
                metadata["breadcrumb"] = str(metadata.get("heading_path") or "").strip()

            score = _score_full_section_candidate(query=query, text=text)
            candidate = {"id": chunk_id, "text": text, "metadata": metadata, "hybrid_score": score}

            prev = src_bucket.get(chunk_id)
            if prev is None or float(prev.get("hybrid_score", 0.0)) < score:
                src_bucket[chunk_id] = candidate

        if src_bucket:
            candidate_by_src[src_key] = src_bucket

    if not candidate_by_src:
        return []

    # Rerank per-source, then interleave for multi-doc fairness.
    reranked_by_src: dict[str, list[dict[str, Any]]] = {}
    for src_key, src_bucket in candidate_by_src.items():
        candidates = sorted(
            src_bucket.values(),
            key=lambda x: float(x.get("hybrid_score", 0.0)),
            reverse=True,
        )
        # Rerank this source's candidates individually.
        reranked, _ = rag_pipeline.rerank(
            query=query,
            candidates=candidates,
            top_k=len(candidates),
            use_rerank=True,
        )
        reranked_by_src[src_key] = reranked or candidates

    if multi_doc:
        return _interleave_by_source(reranked_by_src, max_chunks=max_chunks)

    # Single-doc: just return the reranked list.
    single = next(iter(reranked_by_src.values()), [])
    return single[:max_chunks] if max_chunks > 0 else single



def _retrieve_summary_context(
    query: str,
    selected_source_docs: list[dict[str, Any]],
    min_total_chars: int = 1800,
    max_chunks: int = 24,
    top_k_levels: list[int] | tuple[int, ...] | None = None,
) -> list[dict[str, Any]]:
    levels = top_k_levels or [8, 10, 12, 14, 16, 18]

    top_k_results = _retrieve_context_with_retry(
        query=query,
        selected_source_docs=selected_source_docs,
        min_total_chars=min_total_chars,
        max_chunks=max(max_chunks * 2, 24),
        top_k_levels=levels,
    )
    rebalanced_top_k = _prioritize_summary_group_coverage(top_k_results, max_chunks=max_chunks)
    total_chars = sum(len(str(item.get("text", "") or "")) for item in rebalanced_top_k)
    if len(rebalanced_top_k) >= 3 and total_chars >= max(900, min_total_chars // 2):
        return rebalanced_top_k

    # Fallback: broaden with full-section retrieval then rebalance by chapter/group metadata.
    full_section = _retrieve_full_section_context_with_rerank(
        query=query,
        selected_source_docs=selected_source_docs,
        max_chunks=max(80, max_chunks * 5),
    )
    merged = _merge_retrieval_results(
        primary=rebalanced_top_k,
        secondary=full_section,
        max_chunks=max(120, max_chunks * 5),
    )
    rebalanced_full = _prioritize_summary_group_coverage(merged, max_chunks=max_chunks)
    return rebalanced_full or rebalanced_top_k


def _merge_retrieval_results(
    primary: list[dict[str, Any]],
    secondary: list[dict[str, Any]],
    max_chunks: int,
) -> list[dict[str, Any]]:
    merged_by_id: dict[str, dict[str, Any]] = {}

    for item in [*(primary or []), *(secondary or [])]:
        chunk_id = str(item.get("id") or "").strip()
        if not chunk_id:
            metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
            text = str(item.get("text") or "").strip()
            source = _metadata_source_label(metadata)
            chunk_id = f"{source}:{_normalize_heading_label(text)[:80]}"
        if not chunk_id:
            continue

        score = _retrieval_item_score(item)
        existing = merged_by_id.get(chunk_id)
        if existing is None or _retrieval_item_score(existing) < score:
            merged_by_id[chunk_id] = {**item, "id": chunk_id}

    merged = list(merged_by_id.values())
    merged.sort(key=_retrieval_item_score, reverse=True)
    if max_chunks > 0:
        return merged[:max_chunks]
    return merged


